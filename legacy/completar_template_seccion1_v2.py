"""
Script para completar el template Word de la Sección 1
Agrega todas las secciones faltantes (1.1 a 1.8) de forma segura
"""
from docx import Document
from docx.shared import Pt
from pathlib import Path
import config

def buscar_seccion(doc, texto_buscar):
    """Busca un párrafo que contenga el texto"""
    for i, para in enumerate(doc.paragraphs):
        if texto_buscar.upper() in para.text.upper():
            return i
    return None

def insertar_despues(doc, indice, texto, es_titulo=False, tamanio=11):
    """Inserta un párrafo después del índice dado"""
    if indice is None or indice >= len(doc.paragraphs):
        # Insertar al final
        p = doc.add_paragraph()
    else:
        p = doc.paragraphs[indice].insert_paragraph_before()
    
    run = p.add_run(texto)
    run.font.size = Pt(tamanio)
    if es_titulo:
        run.bold = True
    p.space_after = Pt(6 if es_titulo else 12)
    return p

def completar_template_seccion1():
    """Completa el template con todas las secciones faltantes"""
    template_path = config.TEMPLATES_DIR / "seccion_1_info_general.docx"
    
    if not template_path.exists():
        print(f"[ERROR] No existe: {template_path}")
        return
    
    doc = Document(str(template_path))
    
    cambios = []
    
    # Verificar y agregar 1.1 OBJETO
    indice_1_1 = buscar_seccion(doc, "1.1")
    if indice_1_1 is None:
        # Buscar después de Tabla 1
        indice_tabla1 = buscar_seccion(doc, "Tabla 1")
        if indice_tabla1 is not None:
            insertar_despues(doc, indice_tabla1 + 3, "1.1. OBJETO CONTRATO SCJ-1809-2024", es_titulo=True, tamanio=12)
            insertar_despues(doc, indice_tabla1 + 4, "{{ objeto_contrato }}", tamanio=11)
            insertar_despues(doc, indice_tabla1 + 5, "MANTENIMIENTO PREVENTIVO, MANTENIMIENTO CORRECTIVO Y SOPORTE AL SISTEMA DE VIDEOVIGILANCIA DE BOGOTÁ D.C., CON DISPONIBILIDAD DE BOLSA DE REPUESTOS, en las mejores condiciones técnicas y financieras y en aplicación de los principios de colaboración entre entidades públicas, de eficiencia y economía, resulta necesario adelantar un contrato interadministrativo de prestación de servicios con la EMPRESA DE TELECOMUNICACIONES DE BOGOTÁ SA ESP - ETB S.A. E.S.P., para el desarrollo del objeto contractual requerido.", tamanio=11)
            cambios.append("1.1 OBJETO")
    
    # Verificar y agregar 1.2 ALCANCE
    indice_1_2 = buscar_seccion(doc, "1.2")
    if indice_1_2 is None:
        indice_1_1 = buscar_seccion(doc, "1.1")
        if indice_1_1 is not None:
            insertar_despues(doc, indice_1_1 + 3, "1.2. ALCANCE", es_titulo=True, tamanio=12)
            insertar_despues(doc, indice_1_1 + 4, "{{ alcance }}", tamanio=11)
            cambios.append("1.2 ALCANCE")
    
    # Verificar 1.3 - ya debería estar
    indice_1_3 = buscar_seccion(doc, "1.3")
    if indice_1_3 is not None:
        cambios.append("1.3 INFRAESTRUCTURA (ya existe)")
    
    # Verificar y agregar 1.4 GLOSARIO (solo título si falta)
    indice_1_4 = buscar_seccion(doc, "1.4")
    if indice_1_4 is None:
        # Buscar tabla de glosario
        for i, tabla in enumerate(doc.tables):
            if tabla.rows and len(tabla.rows[0].cells) >= 2:
                primera_celda = tabla.rows[0].cells[0].text.upper()
                if "TÉRMINO" in primera_celda or "DEFINICIÓN" in primera_celda:
                    # Encontrar párrafo antes de esta tabla
                    # Buscar en párrafos cercanos
                    for j, para in enumerate(doc.paragraphs):
                        if j < len(doc.paragraphs) - 5:
                            # Insertar antes
                            insertar_despues(doc, j, "1.4. GLOSARIO", es_titulo=True, tamanio=12)
                            cambios.append("1.4 GLOSARIO")
                            break
                    break
    
    # Verificar y agregar 1.5 OBLIGACIONES
    indice_1_5 = buscar_seccion(doc, "1.5")
    if indice_1_5 is None:
        indice_1_4 = buscar_seccion(doc, "1.4")
        if indice_1_4 is None:
            indice_1_4 = len(doc.paragraphs) - 10
        
        if indice_1_4 is not None:
            base_idx = indice_1_4 + 15  # Después de tabla de glosario
            insertar_despues(doc, base_idx, "1.5. OBLIGACIONES", es_titulo=True, tamanio=12)
            insertar_despues(doc, base_idx + 1, "1.5.1. OBLIGACIONES GENERALES", es_titulo=True, tamanio=11)
            insertar_despues(doc, base_idx + 2, "{{ obligaciones_generales }}", tamanio=11)
            insertar_despues(doc, base_idx + 3, "1.5.2. OBLIGACIONES ESPECÍFICAS DEL CONTRATISTA", es_titulo=True, tamanio=11)
            insertar_despues(doc, base_idx + 4, "{{ obligaciones_especificas }}", tamanio=11)
            insertar_despues(doc, base_idx + 5, "1.5.3. OBLIGACIONES ESPECÍFICAS EN MATERIA AMBIENTAL", es_titulo=True, tamanio=11)
            insertar_despues(doc, base_idx + 6, "{{ obligaciones_ambientales }}", tamanio=11)
            insertar_despues(doc, base_idx + 7, "1.5.4. OBLIGACIONES ANEXOS", es_titulo=True, tamanio=11)
            insertar_despues(doc, base_idx + 8, "{{ obligaciones_anexos }}", tamanio=11)
            cambios.append("1.5 OBLIGACIONES")
    
    # Verificar y agregar 1.6 COMUNICADOS
    indice_1_6 = buscar_seccion(doc, "1.6")
    if indice_1_6 is None:
        indice_1_5 = buscar_seccion(doc, "1.5.4")
        if indice_1_5 is None:
            indice_1_5 = len(doc.paragraphs) - 5
        
        if indice_1_5 is not None:
            base_idx = indice_1_5 + 3
            insertar_despues(doc, base_idx, "1.6. COMUNICADOS CONTRATO SCJ-1809-2024", es_titulo=True, tamanio=12)
            insertar_despues(doc, base_idx + 1, "1.6.1. EMITIDOS CONTRATO SCJ-1809-2024", es_titulo=True, tamanio=11)
            insertar_despues(doc, base_idx + 2, "Total comunicados emitidos: {{ total_comunicados_emitidos }}", tamanio=11)
            # La tabla ya debería existir
            insertar_despues(doc, base_idx + 4, "1.6.2. RECIBIDOS CONTRATO SCJ-1809-2024", es_titulo=True, tamanio=11)
            insertar_despues(doc, base_idx + 5, "Total comunicados recibidos: {{ total_comunicados_recibidos }}", tamanio=11)
            cambios.append("1.6 COMUNICADOS")
    
    # Verificar y agregar 1.7 y 1.8 PERSONAL
    indice_1_7 = buscar_seccion(doc, "1.7")
    if indice_1_7 is None:
        indice_1_6 = buscar_seccion(doc, "1.6.2")
        if indice_1_6 is None:
            indice_1_6 = len(doc.paragraphs) - 5
        
        if indice_1_6 is not None:
            base_idx = indice_1_6 + 3
            insertar_despues(doc, base_idx, "1.7. PERSONAL MÍNIMO REQUERIDO", es_titulo=True, tamanio=12)
            # La tabla ya debería existir
            insertar_despues(doc, base_idx + 2, "1.8. PERSONAL DE APOYO", es_titulo=True, tamanio=12)
            cambios.append("1.7 y 1.8 PERSONAL")
    
    # Guardar
    if cambios:
        doc.save(str(template_path))
        print(f"[OK] Template actualizado: {template_path}")
        print(f"[INFO] Secciones agregadas: {', '.join(cambios)}")
    else:
        print("[INFO] No se requirieron cambios - todas las secciones ya existen")

if __name__ == "__main__":
    completar_template_seccion1()





