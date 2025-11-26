"""Script para verificar que la tabla de obligaciones se haya reemplazado correctamente"""
from docx import Document
from pathlib import Path

archivo = Path('output/2025/09_Septiembre/seccion_1_info_general.docx')

if not archivo.exists():
    print(f"Archivo no encontrado: {archivo}")
    exit(1)

doc = Document(str(archivo))
print(f"Total tablas en el documento: {len(doc.tables)}")

# Buscar tabla de obligaciones
tabla_oblig = None
for i, tabla in enumerate(doc.tables):
    if len(tabla.rows) > 0:
        primera_fila = tabla.rows[0]
        encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
        tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
        tiene_obligacion = any('OBLIGACION' in h or 'OBLIGACIÓN' in h for h in encabezados)
        
        if tiene_item and tiene_obligacion:
            tabla_oblig = tabla
            print(f"\nTabla de obligaciones encontrada (índice {i}):")
            print(f"  - Columnas: {len(tabla.columns)}")
            print(f"  - Filas: {len(tabla.rows)}")
            print(f"  - Encabezados: {encabezados}")
            
            # Mostrar primera fila de datos
            if len(tabla.rows) > 1:
                primera_datos = tabla.rows[1]
                print(f"\nPrimera fila de datos:")
                for j, celda in enumerate(primera_datos.cells):
                    texto = celda.text[:50] if len(celda.text) > 50 else celda.text
                    print(f"  Columna {j+1}: {texto}")
            break

if not tabla_oblig:
    print("\n[ERROR] No se encontró la tabla de obligaciones generales")
    print("\nTablas disponibles:")
    for i, tabla in enumerate(doc.tables):
        if len(tabla.rows) > 0:
            encabezados = [celda.text.strip()[:30] for celda in tabla.rows[0].cells]
            print(f"  Tabla {i}: {len(tabla.columns)} columnas, {len(tabla.rows)} filas - {encabezados[0] if encabezados else 'Sin encabezados'}")

