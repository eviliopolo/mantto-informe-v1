"""
Generador Sección 1: Información General del Contrato
Tipo: 🟦 CONTENIDO FIJO (mayoría) + 🟩 EXTRACCIÓN (comunicados, personal)
"""
from io import BytesIO
from pathlib import Path
from typing import Dict, Any, List, Optional
import json
import os
from .base import GeneradorSeccion
from src.utils.formato_moneda import formato_moneda_cop
from src.ia.extractor_observaciones import get_extractor_observaciones
from src.utils.informes_aprobados import obtener_contexto_informes_aprobados
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import config

class GeneradorSeccion1(GeneradorSeccion):
    """Genera la sección 1: Información General del Contrato"""
    
    @property
    def nombre_seccion(self) -> str:
        return "1. INFORMACIÓN GENERAL DEL CONTRATO"
    
    @property
    def template_file(self) -> str:
        return "seccion_1_info_general.docx"
    
    def __init__(self, anio: int, mes: int, usar_llm_observaciones: bool = True, cargar_desde_mongodb: bool = False):
        super().__init__(anio, mes)
        self.comunicados_emitidos: List[Dict] = []
        self.comunicados_recibidos: List[Dict] = []
        self.personal_minimo: List[Dict] = []
        self.personal_apoyo: List[Dict] = []
        self.obligaciones_generales_raw: List[Dict] = []
        self.obligaciones_especificas_raw: List[Dict] = []
        self.obligaciones_ambientales_raw: List[Dict] = []
        self.obligaciones_anexos_raw: List[Dict] = []
        self.usar_llm_observaciones = usar_llm_observaciones
        self.cargar_desde_mongodb = cargar_desde_mongodb
        self.extractor_observaciones = None
        if usar_llm_observaciones:
            try:
                # Obtener credenciales de SharePoint desde config (que ya carga del .env)
                sharepoint_site_url = getattr(config, 'SHAREPOINT_SITE_URL', None) or os.getenv("SHAREPOINT_SITE_URL")
                sharepoint_client_id = getattr(config, 'SHAREPOINT_CLIENT_ID', None) or os.getenv("SHAREPOINT_CLIENT_ID")
                sharepoint_client_secret = getattr(config, 'SHAREPOINT_CLIENT_SECRET', None) or os.getenv("SHAREPOINT_CLIENT_SECRET")
                sharepoint_base_path = getattr(config, 'SHAREPOINT_BASE_PATH', None) or os.getenv("SHAREPOINT_BASE_PATH")
                
                self.extractor_observaciones = get_extractor_observaciones(
                    sharepoint_site_url=sharepoint_site_url,
                    sharepoint_client_id=sharepoint_client_id,
                    sharepoint_client_secret=sharepoint_client_secret,
                    sharepoint_base_path=sharepoint_base_path
                )
            except Exception as e:
                print(f"[WARNING] No se pudo inicializar extractor de observaciones: {e}")
                self.usar_llm_observaciones = False
    
    def cargar_datos(self) -> None:
        """Carga datos fijos y variables de la sección 1"""
        # 1.1 - 1.5: Contenido fijo (ya está en config.CONTRATO)
        
        # 1.5: Obligaciones (EXTRACCIÓN + LLM para observaciones)
        self._cargar_obligaciones()
        
        # 1.6: Comunicados (EXTRACCIÓN)
        self._cargar_comunicados()
        
        # 1.7 - 1.8: Personal (FIJO + EXTRACCIÓN)
        self._cargar_personal()
    
    def _cargar_comunicados(self) -> None:
        """Carga comunicados emitidos y recibidos del mes desde MongoDB o archivo JSON"""
        if self.cargar_desde_mongodb:
            # Cargar desde MongoDB (se hace de forma asíncrona, este método se llama desde el servicio)
            # Los datos ya deben estar cargados en self.comunicados_emitidos y self.comunicados_recibidos
            return
        
        # Cargar desde archivo JSON (comportamiento original)
        archivo_comunicados = config.FUENTES_DIR / f"comunicados_{self.mes}_{self.anio}.json"
        
        if archivo_comunicados.exists():
            with open(archivo_comunicados, 'r', encoding='utf-8') as f:
                data = json.load(f)
                self.comunicados_emitidos = data.get("emitidos", [])
                self.comunicados_recibidos = data.get("recibidos", [])
        else:
            # Datos de ejemplo para desarrollo
            self.comunicados_emitidos = [
                {
                    "numero": "GSC-7444-2025",
                    "fecha": "23/09/2025",
                    "asunto": "INGRESOS ELEMENTOS ALMACÉN SEPTIEMBRE 2025",
                    "adjuntos": "Anexo_1.pdf"
                },
                {
                    "numero": "GSC-7445-2025", 
                    "fecha": "25/09/2025",
                    "asunto": "INFORME SEMANAL SEMANA 38",
                    "adjuntos": "Informe_S38.pdf"
                }
            ]
            self.comunicados_recibidos = [
                {
                    "numero": "ETB-2024-0892",
                    "fecha": "15/09/2025",
                    "asunto": "SOLICITUD INFORMACIÓN ADICIONAL",
                    "adjuntos": "-"
                }
            ]
    
    def _cargar_obligaciones(self) -> None:
        """Carga obligaciones desde MongoDB o JSON y genera observaciones dinámicas con LLM"""
        if self.cargar_desde_mongodb:
            # Cargar desde MongoDB (se hace de forma asíncrona, este método se llama desde el servicio)
            # Los datos ya deben estar cargados en las listas raw
            return
        
        # Cargar desde archivo JSON (comportamiento original)
        archivo_obligaciones = config.FUENTES_DIR / f"obligaciones_{self.mes}_{self.anio}.json"
        if not archivo_obligaciones.exists():
            archivo_obligaciones = config.FUENTES_DIR / f"obligaciones_{config.MESES[self.mes].lower()}_{self.anio}.json"
        
        if archivo_obligaciones.exists():
            try:
                with open(archivo_obligaciones, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.obligaciones_generales_raw = data.get("obligaciones_generales", [])
                    self.obligaciones_especificas_raw = data.get("obligaciones_especificas", [])
                    self.obligaciones_ambientales_raw = data.get("obligaciones_ambientales", [])
                    self.obligaciones_anexos_raw = data.get("obligaciones_anexos", [])
                    
                    # Generar observaciones dinámicas si está habilitado
                    if self.usar_llm_observaciones and self.extractor_observaciones:
                        print("[INFO] Generando observaciones dinámicas desde anexos usando LLM...")
                        
                        # Obtener contexto de los últimos 3 informes aprobados
                        print("[INFO] Obteniendo contexto de informes aprobados anteriores...")
                        contexto_informes = obtener_contexto_informes_aprobados(cantidad=3)
                        
                        # Procesar obligaciones con contexto de informes aprobados
                        self.obligaciones_generales_raw = [
                            self.extractor_observaciones.procesar_obligacion(obl, contexto_informes)
                            for obl in self.obligaciones_generales_raw
                        ]
                        self.obligaciones_especificas_raw = [
                            self.extractor_observaciones.procesar_obligacion(obl, contexto_informes)
                            for obl in self.obligaciones_especificas_raw
                        ]
                        self.obligaciones_ambientales_raw = [
                            self.extractor_observaciones.procesar_obligacion(obl, contexto_informes)
                            for obl in self.obligaciones_ambientales_raw
                        ]
                        self.obligaciones_anexos_raw = [
                            self.extractor_observaciones.procesar_obligacion(obl, contexto_informes)
                            for obl in self.obligaciones_anexos_raw
                        ]
            except Exception as e:
                print(f"[WARNING] Error al cargar obligaciones desde {archivo_obligaciones}: {e}")
        else:
            print(f"[INFO] Archivo de obligaciones no encontrado: {archivo_obligaciones}")
            # Las listas quedan vacías - se usarán datos fijos del texto
    
    def _cargar_personal(self) -> None:
        """Carga información del personal del contrato"""
        archivo_personal = config.FIJOS_DIR / "personal_requerido.json"
        
        if archivo_personal.exists():
            with open(archivo_personal, 'r', encoding='utf-8') as f:
                data = json.load(f)
                self.personal_minimo = data.get("minimo", [])
                self.personal_apoyo = data.get("apoyo", [])
        else:
            # Estructura de ejemplo
            self.personal_minimo = [
                {"cargo": "Director de Proyecto", "cantidad": 1, "nombre": "Por definir"},
                {"cargo": "Coordinador Técnico", "cantidad": 1, "nombre": "Por definir"},
                {"cargo": "Ingeniero de Soporte", "cantidad": 2, "nombre": "Por definir"},
                {"cargo": "Técnico de Campo", "cantidad": 8, "nombre": "Por definir"},
            ]
            self.personal_apoyo = [
                {"cargo": "Técnico de Laboratorio", "cantidad": 2, "nombre": "Por definir"},
                {"cargo": "Auxiliar Administrativo", "cantidad": 1, "nombre": "Por definir"},
            ]
    
    def _cargar_contenido_fijo(self, archivo: str) -> str:
        """Carga contenido fijo desde archivo de texto"""
        ruta = config.FIJOS_DIR / archivo
        if ruta.exists():
            with open(ruta, 'r', encoding='utf-8') as f:
                return f.read()
        return ""
    
    def procesar(self) -> Dict[str, Any]:
        """Procesa y retorna el contexto para el template"""
        # Texto introductorio oficial
        texto_intro = (
            f"Se celebra el número de proceso {config.CONTRATO['numero_proceso']} bajo número de contrato "
            f"{config.CONTRATO['numero']} con vigencia de doce (12) meses luego de suscripción de acta de inicio "
            f"suscrita el {self._formatear_fecha(config.CONTRATO['fecha_inicio'])}, fecha a partir de la cual el "
            f"sistema de video vigilancia de Bogotá D.C. queda con contrato de mantenimiento de videovigilancia. "
            f"Se detalla la información general del contrato."
        )
        
        return {
            # Texto introductorio
            "texto_intro": texto_intro,
            
            # TABLA 1: Información General del Contrato (formato lista para docxtpl)
            "tabla_1_filas": self._formatear_tabla_1(),
            "tabla_1_info_general": {
                "nit": config.CONTRATO["nit_entidad"],
                "razon_social": config.CONTRATO["razon_social"],
                "ciudad": config.CONTRATO["ciudad"],
                "direccion": config.CONTRATO["direccion"],
                "telefono": config.CONTRATO["telefono"],
                "numero_contrato": config.CONTRATO["numero"],
                "fecha_inicio": self._formatear_fecha(config.CONTRATO["fecha_inicio"]),
                "plazo_ejecucion": config.CONTRATO["plazo_ejecucion"],
                "fecha_terminacion": self._formatear_fecha(config.CONTRATO["fecha_fin"]),
                "valor_inicial": formato_moneda_cop(config.CONTRATO["valor_inicial"]),
                "adicion_1": formato_moneda_cop(config.CONTRATO["adicion_1"]),
                "valor_total": formato_moneda_cop(config.CONTRATO["valor_total"]),
                "objeto": config.CONTRATO["objeto"],
                "fecha_firma_acta": self._formatear_fecha(config.CONTRATO["fecha_inicio"]),
                "fecha_suscripcion": self._formatear_fecha(config.CONTRATO["fecha_suscripcion"]),
                "vigencia_poliza_inicial": f"{self._formatear_fecha(config.CONTRATO['vigencia_poliza_inicial_inicio'])} {self._formatear_fecha(config.CONTRATO['vigencia_poliza_inicial_fin'])}",
                "vigencia_poliza_acta": f"{self._formatear_fecha(config.CONTRATO['vigencia_poliza_acta_inicio'])} {self._formatear_fecha(config.CONTRATO['vigencia_poliza_acta_fin'])}",
            },
            
            # Variables para textos de anexos (opcionales)
            "ruta_acta_inicio": self._obtener_ruta_acta_inicio(),
            "numero_adicion": self._obtener_numero_adicion(),
            "ruta_poliza": self._obtener_ruta_poliza(),
            
            # 1.1 Objeto del contrato (FIJO)
            "objeto_contrato": config.CONTRATO["objeto_corto"],
            
            # 1.2 Alcance (FIJO)
            "alcance": self._cargar_contenido_fijo("alcance.txt"),
            
            # 1.3 Descripción infraestructura (FIJO + TABLAS)
            "descripcion_infraestructura": self._cargar_contenido_fijo("infraestructura.txt"),
            "subsistemas": config.SUBSISTEMAS,
            "componentes": self._cargar_tabla_componentes(),  # Renombrado para consistencia
            "centros": self._cargar_tabla_centros_monitoreo(),  # Renombrado para consistencia
            "forma_pago": self._cargar_tabla_forma_pago(),  # Renombrado para consistencia
            "tabla_componentes": self._cargar_tabla_componentes(),  # Mantener compatibilidad
            "tabla_centros_monitoreo": self._cargar_tabla_centros_monitoreo(),  # Mantener compatibilidad
            "tabla_forma_pago": self._cargar_tabla_forma_pago(),  # Mantener compatibilidad
            "nota_infraestructura": self._obtener_nota_infraestructura(),
            
            # 1.4 Glosario (FIJO)
            "glosario": self._cargar_glosario(),
            "glosario_tablas": self._formatear_glosario_tablas(),
            
            # 1.5 Obligaciones (FIJO texto + DINÁMICO tablas)
            "obligaciones_generales": self._cargar_contenido_fijo("obligaciones_generales.txt"),
            "obligaciones_especificas": self._cargar_contenido_fijo("obligaciones_especificas.txt"),
            "obligaciones_ambientales": self._cargar_contenido_fijo("obligaciones_ambientales.txt"),
            "obligaciones_anexos": self._cargar_contenido_fijo("obligaciones_anexos.txt"),
            
            # Tablas de obligaciones (DINÁMICAS - con cumplimiento)
            "tabla_obligaciones_generales": self._formatear_obligaciones_generales(),
            "tabla_obligaciones_especificas": self._formatear_obligaciones_especificas(),
            "tabla_obligaciones_ambientales": self._formatear_obligaciones_ambientales(),
            "tabla_obligaciones_anexos": self._formatear_obligaciones_anexos(),
            
            # 1.6 Comunicados (EXTRACCIÓN)
            "comunicados_emitidos": self.comunicados_emitidos,
            "comunicados_recibidos": self.comunicados_recibidos,
            "total_comunicados_emitidos": len(self.comunicados_emitidos),
            "total_comunicados_recibidos": len(self.comunicados_recibidos),
            "tabla_comunicados_emitidos": self._formatear_comunicados_emitidos(),
            "tabla_comunicados_recibidos": self._formatear_comunicados_recibidos(),
            
            # 1.7 - 1.8 Personal (FIJO estructura + EXTRACCIÓN datos)
            "personal_minimo": self.personal_minimo,
            "personal_apoyo": self.personal_apoyo,
            "tabla_personal_minimo": self._formatear_personal_minimo(),
            "tabla_personal_apoyo": self._formatear_personal_apoyo(),
            
            # Marcadores para identificación de tablas (renderizan como [[TABLA_XXX]])
            "TABLA_MARKER_OBLIGACIONES_GENERALES": "[[TABLA_OBLIGACIONES_GENERALES]]",
            "TABLA_MARKER_OBLIGACIONES_ESPECIFICAS": "[[TABLA_OBLIGACIONES_ESPECIFICAS]]",
            "TABLA_MARKER_OBLIGACIONES_AMBIENTALES": "[[TABLA_OBLIGACIONES_AMBIENTALES]]",
            "TABLA_MARKER_OBLIGACIONES_ANEXOS": "[[TABLA_OBLIGACIONES_ANEXOS]]",
            "TABLA_MARKER_COMUNICADOS_EMITIDOS": "[[TABLA_COMUNICADOS_EMITIDOS]]",
            "TABLA_MARKER_COMUNICADOS_RECIBIDOS": "[[TABLA_COMUNICADOS_RECIBIDOS]]",
        }
    
    def _cargar_glosario(self) -> List[Dict[str, str]]:
        """Carga el glosario de términos"""
        archivo = config.FIJOS_DIR / "glosario.json"
        if archivo.exists():
            with open(archivo, 'r', encoding='utf-8') as f:
                return json.load(f)
        
        # Glosario por defecto
        return [
            {"termino": "ANS", "definicion": "Acuerdo de Nivel de Servicio"},
            {"termino": "CCTV", "definicion": "Circuito Cerrado de Televisión"},
            {"termino": "DVR", "definicion": "Digital Video Recorder"},
            {"termino": "NVR", "definicion": "Network Video Recorder"},
            {"termino": "GLPI", "definicion": "Gestionnaire Libre de Parc Informatique"},
            {"termino": "NUSE", "definicion": "Número Único de Seguridad y Emergencias"},
        ]
    
    def _formatear_fecha(self, fecha_str: str) -> str:
        """Formatea fecha YYYY-MM-DD a formato DD DE MES DE YYYY"""
        from datetime import datetime
        try:
            fecha = datetime.strptime(fecha_str, "%Y-%m-%d")
            meses = {
                1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
                5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
                9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
            }
            return f"{fecha.day} de {meses[fecha.month]} de {fecha.year}"
        except:
            return fecha_str
    
    def _cargar_tabla_componentes(self) -> List[Dict]:
        """Carga tabla de componentes por subsistema"""
        # Datos según el informe oficial de Septiembre 2025
        return [
            {
                "numero": 1,
                "sistema": "CIUDADANA",
                "ubicaciones": 4451,
                "puntos_camara": 4451,
                "centros_monitoreo_c4": 4451,
                "visualizadas_localmente": 0
            },
            {
                "numero": 2,
                "sistema": "COLEGIOS",
                "ubicaciones": 98,
                "puntos_camara": 235,
                "centros_monitoreo_c4": 235,
                "visualizadas_localmente": 0
            },
            {
                "numero": 3,
                "sistema": "TRANSMILENIO",
                "ubicaciones": 71,
                "puntos_camara": 164,
                "centros_monitoreo_c4": 164,
                "visualizadas_localmente": 0
            },
            {
                "numero": 4,
                "sistema": "CAI",
                "ubicaciones": 157,
                "puntos_camara": 510,
                "centros_monitoreo_c4": 89,
                "visualizadas_localmente": 421
            },
            {
                "numero": 5,
                "sistema": "ESTADIO EL CAMPIN",
                "ubicaciones": 1,
                "puntos_camara": 58,
                "centros_monitoreo_c4": 0,
                "visualizadas_localmente": 58
            },
            {
                "numero": 6,
                "sistema": "CTP",
                "ubicaciones": 1,
                "puntos_camara": 104,
                "centros_monitoreo_c4": 0,
                "visualizadas_localmente": 104
            },
            {
                "numero": 7,
                "sistema": "ESTACIONES DE POLICÍA",
                "ubicaciones": 24,
                "puntos_camara": 302,
                "centros_monitoreo_c4": 0,
                "visualizadas_localmente": 302
            },
            {
                "numero": 8,
                "sistema": "TOTAL",
                "ubicaciones": 4803,
                "puntos_camara": 5824,
                "centros_monitoreo_c4": 4939,
                "visualizadas_localmente": 885
            }
        ]
    
    def _cargar_tabla_centros_monitoreo(self) -> List[Dict]:
        """Carga tabla de centros de monitoreo"""
        return [
            {"numero": 1, "nombre": "CENTRO DE COMANDO, CONTROL, CÓMPUTO Y COMUNICACIONES - C4", "direccion": "CALLE 20 NO 68 A 06", "localidad": "PUENTE ARANDA"},
            {"numero": 2, "nombre": "CENTRO DE MONITOREO ENGATIVÁ", "direccion": "KR 78A NO. 70 – 54", "localidad": "ENGATIVÁ"},
            {"numero": 3, "nombre": "CENTRO DE MONITOREO BARRIOS UNIDOS", "direccion": "ESTACIÓN POLICÍA CALLE 72 # 62-81", "localidad": "BARRIOS UNIDOS"},
            {"numero": 4, "nombre": "CENTRO DE MONITOREO TEUSAQUILLO", "direccion": "ESTACIÓN POLICÍA CRA 13 # 39-86", "localidad": "TEUSAQUILLO"},
            {"numero": 5, "nombre": "CENTRO DE MONITOREO KENNEDY", "direccion": "TRANSVERSAL 78 K CON CALLE 41 D SUR", "localidad": "KENNEDY"},
            {"numero": 6, "nombre": "CENTRO DE MONITOREO CHAPINERO", "direccion": "KR 1 CALLE 57-00", "localidad": "CHAPINERO"},
            {"numero": 7, "nombre": "CENTRO DE MONITOREO CIUDAD BOLÍVAR", "direccion": "DIAGONAL 70 SUR CON TRANSVERSAL 54", "localidad": "CIUDAD BOLÍVAR"},
            {"numero": 8, "nombre": "CENTRO DE MONITOREO PUENTE ARANDA", "direccion": "CRA 39 CON CALLE 10", "localidad": "PUENTE ARANDA"},
            {"numero": 9, "nombre": "CENTRO DE MONITOREO USAQUÉN", "direccion": "CL. 165 #8A-99", "localidad": "USAQUÉN"},
            {"numero": 10, "nombre": "CENTRO DE MONITOREO RAFAEL URIBE", "direccion": "Calle 27 Sur #24-39", "localidad": "RAFAEL URIBE URIBE"},
            {"numero": 11, "nombre": "CENTRO DE MONITOREO SANTA FE", "direccion": "Carrera 5 # 29-11", "localidad": "SANTA FE"},
        ]
    
    def _cargar_tabla_forma_pago(self) -> List[Dict]:
        """Carga tabla de forma de pago"""
        return [
            {
                "numero": 1,
                "descripcion": "Mantenimientos preventivos por UBICACIÓN, aprobados mediante cronograma con interventoría / supervisión.",
                "tipo_servicio": "Por Demanda"
            },
            {
                "numero": 2,
                "descripcion": "Servicio de mantenimiento correctivo y soporte al sistema de video vigilancia de Bogotá",
                "tipo_servicio": "Mensualidad"
            },
            {
                "numero": 3,
                "descripcion": "Bolsa de repuestos, elementos aprobados por interventoría / supervisión.",
                "tipo_servicio": "Por Demanda"
            }
        ]
    
    def _formatear_tabla_1(self) -> List[Dict[str, str]]:
        """Formatea la tabla 1 como lista de filas (Campo | Valor)"""
        return [
            {"campo": "NIT", "valor": config.CONTRATO["nit_entidad"]},
            {"campo": "RAZÓN SOCIAL", "valor": config.CONTRATO["razon_social"]},
            {"campo": "CIUDAD", "valor": config.CONTRATO["ciudad"]},
            {"campo": "DIRECCIÓN", "valor": config.CONTRATO["direccion"]},
            {"campo": "TELÉFONO", "valor": config.CONTRATO["telefono"]},
            {"campo": "NÚMERO DE CONTRATO", "valor": config.CONTRATO["numero"]},
            {"campo": "FECHA DE INICIO", "valor": self._formatear_fecha(config.CONTRATO["fecha_inicio"])},
            {"campo": "PLAZO DE EJECUCIÓN", "valor": config.CONTRATO["plazo_ejecucion"]},
            {"campo": "FECHA DE TERMINACIÓN", "valor": self._formatear_fecha(config.CONTRATO["fecha_fin"])},
            {"campo": "VALOR INICIAL", "valor": formato_moneda_cop(config.CONTRATO["valor_inicial"])},
            {"campo": "ADICIÓN N° 01", "valor": formato_moneda_cop(config.CONTRATO["adicion_1"])},
            {"campo": "VALOR TOTAL", "valor": formato_moneda_cop(config.CONTRATO["valor_total"])},
            {"campo": "OBJETO", "valor": config.CONTRATO["objeto"]},
            {"campo": "FECHA FIRMA ACTA DE INICIO", "valor": self._formatear_fecha(config.CONTRATO["fecha_inicio"])},
            {"campo": "FECHA DE SUSCRIPCIÓN", "valor": self._formatear_fecha(config.CONTRATO["fecha_suscripcion"])},
            {"campo": "VIGENCIA PÓLIZA INICIAL", "valor": f"{self._formatear_fecha(config.CONTRATO['vigencia_poliza_inicial_inicio'])} - {self._formatear_fecha(config.CONTRATO['vigencia_poliza_inicial_fin'])}"},
            {"campo": "VIGENCIA PÓLIZA ACTA DE INICIO", "valor": f"{self._formatear_fecha(config.CONTRATO['vigencia_poliza_acta_inicio'])} - {self._formatear_fecha(config.CONTRATO['vigencia_poliza_acta_fin'])}"},
        ]
    
    def _formatear_comunicados_emitidos(self) -> List[Dict]:
        """Formatea comunicados emitidos para tabla (ÍTEM, FECHA, CONSECUTIVO ETB, DESCRIPCIÓN)"""
        return [
            {
                "item": com.get("item", i+1),
                "fecha": com.get("fecha", ""),
                "consecutivo": com.get("numero", com.get("radicado", "")),
                "descripcion": com.get("asunto", "")
            }
            for i, com in enumerate(self.comunicados_emitidos)
        ]
    
    def _formatear_comunicados_recibidos(self) -> List[Dict]:
        """Formatea comunicados recibidos para tabla (ÍTEM, FECHA, CONSECUTIVO ETB, DESCRIPCIÓN)"""
        return [
            {
                "item": com.get("item", i+1),
                "fecha": com.get("fecha", ""),
                "consecutivo": com.get("numero", com.get("radicado", "")),
                "descripcion": com.get("asunto", ""),
                "asunto": com.get("asunto", "")  # Alias para compatibilidad
            }
            for i, com in enumerate(self.comunicados_recibidos)
        ]
    
    def _formatear_personal_minimo(self) -> List[Dict]:
        """Formatea personal mínimo para tabla"""
        return [
            {
                "cargo": p.get("cargo", ""),
                "cantidad": p.get("cantidad", 0),
                "nombre": p.get("nombre", "Por definir")
            }
            for p in self.personal_minimo
        ]
    
    def _formatear_personal_apoyo(self) -> List[Dict]:
        """Formatea personal de apoyo para tabla"""
        return [
            {
                "cargo": p.get("cargo", ""),
                "cantidad": p.get("cantidad", 0),
                "nombre": p.get("nombre", "Por definir")
            }
            for p in self.personal_apoyo
        ]
    
    def _formatear_glosario_tablas(self) -> List[Dict]:
        """Formatea glosario para múltiples tablas (se divide en grupos)"""
        glosario = self._cargar_glosario()
        # Retornar lista completa - el template dividirá en tablas si es necesario
        return glosario
    
    def _formatear_obligaciones_generales(self) -> List[Dict]:
        """Formatea obligaciones generales para tabla con cumplimiento"""
        # Retornar obligaciones cargadas (ya procesadas con observaciones)
        return self.obligaciones_generales_raw
    
    def _formatear_obligaciones_especificas(self) -> List[Dict]:
        """Formatea obligaciones específicas para tabla con cumplimiento"""
        return self.obligaciones_especificas_raw
    
    def _formatear_obligaciones_ambientales(self) -> List[Dict]:
        """Formatea obligaciones ambientales para tabla con cumplimiento"""
        return self.obligaciones_ambientales_raw
    
    def _formatear_obligaciones_anexos(self) -> List[Dict]:
        """Formatea obligaciones anexos para tabla con cumplimiento"""
        return self.obligaciones_anexos_raw
    
    def _obtener_ruta_acta_inicio(self) -> str:
        """Obtiene la ruta del acta de inicio (dinámico según mes)"""
        mes_nombre = config.MESES[self.mes]
        return f"01{mes_nombre[:3].upper()} - 30{mes_nombre[:3].upper()}/ 01 OBLIGACIONES GENERALES/ OBLIGACIÓN 2,5,6,9,13/ ANEXOS OTROS/ SCJ-1809-2024 ACTA DE INICIO.PDF"
    
    def _obtener_numero_adicion(self) -> str:
        """Obtiene el número de adición si aplica (dinámico)"""
        # TODO: Cargar desde fuente de datos o config
        return "01"  # Por defecto
    
    def _obtener_ruta_poliza(self) -> str:
        """Obtiene la ruta de modificación de póliza si aplica (dinámico)"""
        mes_nombre = config.MESES[self.mes]
        return f"01{mes_nombre[:3].upper()} - 30{mes_nombre[:3].upper()} / 01 OBLIGACIONES GENERALES/ OBLIGACIÓN 4/ ANEXOS OTROS/ MODIFICACIÓN PÓLIZA.PDF"
    
    def _obtener_nota_infraestructura(self) -> str:
        """Obtiene nota adicional sobre infraestructura si aplica (dinámico)"""
        # Ejemplo: "Se aclara que se desmonta La Estación de Transmilenio Calle 26 con 12 cámaras por obras del metro"
        # TODO: Cargar desde fuente de datos
        return ""
    
    def generar(self):
        """
        Genera la sección completa usando enfoque híbrido con marcadores únicos:
        - Jinja2 para variables simples (textos, números, etc.) - FUERA de tablas
        - Reemplazo programático para tablas usando marcadores únicos en el template
        
        IMPORTANTE: En el template Word, cada tabla debe tener un marcador único
        en la primera celda de la primera fila de datos (después del encabezado):
        - {{ TABLA_MARKER_OBLIGACIONES_GENERALES }} para 1.5.1
        - {{ TABLA_MARKER_OBLIGACIONES_ESPECIFICAS }} para 1.5.2
        - {{ TABLA_MARKER_OBLIGACIONES_AMBIENTALES }} para 1.5.3
        - {{ TABLA_MARKER_OBLIGACIONES_ANEXOS }} para 1.5.4
        - {{ TABLA_MARKER_COMUNICADOS_EMITIDOS }} para 1.6.1
        - {{ TABLA_MARKER_COMUNICADOS_RECIBIDOS }} para 1.6.2
        
        Estos marcadores son variables de Jinja2 que se renderizan con un texto único
        que luego se busca en el documento para identificar cada tabla.
        """
        # Generar el documento base usando el método de la clase padre (docxtpl)
        doc_template = super().generar()
        
        # Guardar el documento renderizado en memoria y reabrirlo con python-docx
        # Esto evita corromper el XML al manipular tablas después del render
        buffer = BytesIO()
        doc_template.save(buffer)
        buffer.seek(0)
        doc = Document(buffer)
        
        # Llenar las tablas usando reemplazo programático con marcadores únicos
        # Esto evita problemas de corrupción del XML y permite identificación precisa
        
        print("[INFO] Llenando tablas usando marcadores únicos...")
        print(f"[DEBUG] obligaciones_generales_raw: {len(self.obligaciones_generales_raw)} elementos")
        print(f"[DEBUG] obligaciones_especificas_raw: {len(self.obligaciones_especificas_raw)} elementos")
        print(f"[DEBUG] obligaciones_ambientales_raw: {len(self.obligaciones_ambientales_raw)} elementos")
        print(f"[DEBUG] obligaciones_anexos_raw: {len(self.obligaciones_anexos_raw)} elementos")
        print(f"[DEBUG] comunicados_emitidos: {len(self.comunicados_emitidos)} elementos")
        print(f"[DEBUG] comunicados_recibidos: {len(self.comunicados_recibidos)} elementos")
        
        # Procesar cada tabla usando su marcador único
        self._reemplazar_tabla_por_marcador(doc, "TABLA_OBLIGACIONES_GENERALES", 
                                           self._formatear_obligaciones_generales(),
                                           self._crear_tabla_obligaciones_generales)
        
        self._reemplazar_tabla_por_marcador(doc, "TABLA_OBLIGACIONES_ESPECIFICAS",
                                           self._formatear_obligaciones_especificas(),
                                           self._crear_tabla_obligaciones_especificas)
        
        self._reemplazar_tabla_por_marcador(doc, "TABLA_OBLIGACIONES_AMBIENTALES",
                                           self._formatear_obligaciones_ambientales(),
                                           self._crear_tabla_obligaciones_ambientales)
        
        self._reemplazar_tabla_por_marcador(doc, "TABLA_OBLIGACIONES_ANEXOS",
                                           self._formatear_obligaciones_anexos(),
                                           self._crear_tabla_obligaciones_anexos)
        
        self._reemplazar_tabla_por_marcador(doc, "TABLA_COMUNICADOS_EMITIDOS",
                                           self._formatear_comunicados_emitidos(),
                                           self._crear_tabla_comunicados_emitidos)
        
        self._reemplazar_tabla_por_marcador(doc, "TABLA_COMUNICADOS_RECIBIDOS",
                                           self._formatear_comunicados_recibidos(),
                                           self._crear_tabla_comunicados_recibidos)
        
        return doc
    
    def _reemplazar_tabla_por_marcador(self, doc: Document, marcador: str, datos: list, metodo_creacion) -> None:
        """
        Busca una tabla en el documento usando un marcador único y la reemplaza con datos.
        
        Args:
            doc: Documento de python-docx
            marcador: Nombre del marcador (ej: "TABLA_OBLIGACIONES_GENERALES")
            datos: Lista de datos para llenar la tabla
            metodo_creacion: Método que crea/llena la tabla (ej: self._crear_tabla_obligaciones_generales)
        """
        # El marcador se renderiza como [[TABLA_XXX]] después de procesar Jinja2
        marcador_renderizado = f"[[{marcador}]]"
        
        # Variaciones del marcador que pueden aparecer
        marcador_variaciones = [
            marcador_renderizado,  # Después de procesar Jinja2: [[TABLA_XXX]]
            f"{{{{ TABLA_MARKER_{marcador.replace('TABLA_', '')} }}}}",  # En el template original
            f"TABLA_MARKER_{marcador.replace('TABLA_', '')}",  # Variable de Jinja2 sin renderizar
            marcador.upper(),  # Nombre del marcador en mayúsculas
        ]
        
        # Normalizar el marcador para búsqueda
        marcador_busqueda = marcador.upper().replace("TABLA_", "")
        print(f"[INFO] Buscando tabla con marcador: {marcador_renderizado} (variaciones: {marcador_variaciones})")
        
        tabla_encontrada = None
        tabla_idx = None
        celda_con_marcador = None
        
        # Buscar el marcador en todas las tablas del documento
        for idx, tabla in enumerate(doc.tables):
            # Buscar en todas las celdas de la tabla
            for fila_idx, fila in enumerate(tabla.rows):
                for celda_idx, celda in enumerate(fila.cells):
                    # Obtener todo el texto de la celda (incluyendo párrafos)
                    texto_celda = ""
                    for parrafo in celda.paragraphs:
                        texto_celda += parrafo.text
                    
                    texto_celda_upper = texto_celda.upper()
                    
                    # Verificar si contiene el marcador renderizado [[TABLA_XXX]]
                    encontro_marcador = marcador_renderizado.upper() in texto_celda_upper
                    
                    # Si no se encuentra, buscar por variaciones
                    if not encontro_marcador:
                        for variacion in marcador_variaciones:
                            if variacion.upper() in texto_celda_upper:
                                encontro_marcador = True
                                break
                    
                    # También buscar por el nombre del marcador sin prefijos
                    if not encontro_marcador and marcador_busqueda in texto_celda_upper:
                        encontro_marcador = True
                    
                    if encontro_marcador:
                        tabla_encontrada = tabla
                        tabla_idx = idx
                        celda_con_marcador = (fila_idx, celda_idx)
                        print(f"[INFO] Marcador encontrado en tabla {idx}, fila {fila_idx}, celda {celda_idx}")
                        print(f"[DEBUG] Texto de la celda: '{texto_celda[:100]}...'")
                        break
                
                if tabla_encontrada:
                    break
            
            if tabla_encontrada:
                break
        
        if tabla_encontrada and tabla_idx is not None:
            # No necesitamos limpiar el marcador aquí porque el método _crear_tabla_*
            # ya limpia todas las filas excepto el encabezado antes de agregar los datos
            # Solo limpiamos el marcador si está en el encabezado (fila 0)
            if celda_con_marcador:
                fila_idx, celda_idx = celda_con_marcador
                if fila_idx == 0:
                    # Si el marcador está en el encabezado, limpiarlo
                    celda = tabla_encontrada.rows[0].cells[celda_idx]
                    for parrafo in celda.paragraphs:
                        parrafo.clear()
            
            # Llamar al método de creación para reemplazar la tabla
            # Este método ya limpia todas las filas excepto el encabezado
            print(f"[INFO] Reemplazando tabla {tabla_idx} con {len(datos)} elementos")
            try:
                metodo_creacion(doc, tabla_encontrada)
                print(f"[OK] Tabla '{marcador}' procesada correctamente")
            except Exception as e:
                print(f"[ERROR] Error al procesar tabla '{marcador}': {str(e)}")
                import traceback
                traceback.print_exc()
        else:
            print(f"[WARNING] No se encontró tabla con marcador '{marcador}'")
            print(f"[INFO] Asegúrate de agregar '{{{{ TABLA_MARKER_{marcador.replace('TABLA_', '')} }}}}' en la primera celda de datos de la tabla en el template")
            print(f"[DEBUG] Marcador buscado: {marcador_variaciones}")
    
    def guardar(self, output_path: Path) -> None:
        """
        Genera y guarda la sección, asegurando que los cambios en las tablas se guarden correctamente
        """
        doc = self.generar()
        
        # Guardar usando python-docx después de reemplazar las tablas
        doc.save(str(output_path))
        print(f"[OK] {self.nombre_seccion} guardada en: {output_path}")
    
    def _reemplazar_tabla_obligaciones_generales(self, doc: Document, tablas_procesadas: set = None) -> None:
        """
        Busca y reemplaza la tabla de obligaciones generales con datos dinámicos
        """
        if tablas_procesadas is None:
            tablas_procesadas = set()
        
        if not self.obligaciones_generales_raw:
            print("[WARNING] No hay obligaciones generales para reemplazar en la tabla")
            return
        
        # Buscar la tabla que está después del texto "1.5.1" o "OBLIGACIONES GENERALES"
        tabla_encontrada = None
        tabla_idx = None
        
        # Estrategia 1: Buscar párrafos que mencionen "1.5.1" o "OBLIGACIONES GENERALES"
        # y luego buscar la siguiente tabla
        encontro_titulo = False
        for parrafo in doc.paragraphs:
            texto = parrafo.text.upper()
            if '1.5.1' in texto or ('OBLIGACIONES' in texto and 'GENERALES' in texto):
                encontro_titulo = True
                # Buscar la siguiente tabla después de este párrafo
                # Necesitamos encontrar el elemento XML del párrafo y buscar el siguiente elemento tabla
                break
        
        # Si encontramos el título, buscar la tabla siguiente
        if encontro_titulo:
            # Buscar en el XML del documento
            elementos = doc.element.body
            tabla_count_before = 0
            encontro_titulo_xml = False
            
            for i, elemento in enumerate(elementos):
                # Si encontramos una tabla antes del título, incrementar contador
                if hasattr(elemento, 'tag') and elemento.tag.endswith('}tbl'):
                    if not encontro_titulo_xml:
                        tabla_count_before += 1
                
                # Verificar si es un párrafo con el texto buscado
                if hasattr(elemento, 'text') and elemento.text:
                    texto = elemento.text.upper()
                    if '1.5.1' in texto or ('OBLIGACIONES' in texto and 'GENERALES' in texto):
                        encontro_titulo_xml = True
                        print(f"[INFO] Título encontrado en elemento {i}, tablas antes: {tabla_count_before}")
                        # Buscar la siguiente tabla después del título
                        for j in range(i + 1, len(elementos)):
                            siguiente = elementos[j]
                            if hasattr(siguiente, 'tag') and siguiente.tag.endswith('}tbl'):
                                # Esta es la tabla que buscamos
                                tabla_index = tabla_count_before
                                if tabla_index < len(doc.tables):
                                    tabla_candidata = doc.tables[tabla_index]
                                    # Verificar que sea la tabla correcta por sus encabezados
                                    if len(tabla_candidata.rows) > 0:
                                        primera_fila = tabla_candidata.rows[0]
                                        encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                                        tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                                        tiene_obligacion = any('OBLIGACION' in h or 'OBLIGACIÓN' in h for h in encabezados)
                                        if tiene_item and tiene_obligacion:
                                            tabla_encontrada = tabla_candidata
                                            tabla_idx = tabla_index
                                            print(f"[INFO] Tabla de obligaciones encontrada (índice {tabla_index})")
                                            break
                                tabla_count_before += 1
                        break
        
        # Estrategia 2: Si no encontramos por título, buscar tabla con formato correcto
        if not tabla_encontrada:
            print("[INFO] Buscando tabla por formato de encabezados...")
            for idx, tabla in enumerate(doc.tables):
                if idx not in tablas_procesadas and len(tabla.rows) > 0:
                    primera_fila = tabla.rows[0]
                    encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                    print(f"[DEBUG] Tabla {idx}: {len(tabla.columns)} columnas, encabezados: {encabezados[:3]}...")
                    
                    # Buscar palabras clave en los encabezados
                    tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                    tiene_obligacion = any('OBLIGACION' in h or 'OBLIGACIÓN' in h for h in encabezados)
                    tiene_periodicidad = any('PERIODICIDAD' in h for h in encabezados)
                    tiene_observaciones = any('OBSERVACION' in h or 'OBSERVACIÓN' in h for h in encabezados)
                    tiene_anexo = any('ANEXO' in h for h in encabezados)
                    
                    # Priorizar tablas con más columnas y más palabras clave
                    puntuacion = sum([tiene_item, tiene_obligacion, tiene_periodicidad, tiene_observaciones, tiene_anexo])
                    
                    if tiene_item and tiene_obligacion and tiene_periodicidad:
                        if not tabla_encontrada or (puntuacion > 3 and len(tabla.columns) >= 5):
                            tabla_encontrada = tabla
                            tabla_idx = idx
                            print(f"[INFO] Tabla candidata encontrada (índice {idx}): {len(tabla.columns)} columnas, puntuación: {puntuacion}")
                            if puntuacion >= 4 and len(tabla.columns) >= 5:
                                print(f"[INFO] Tabla seleccionada: {len(tabla.columns)} columnas")
                                break
        
        if not tabla_encontrada:
            print("[WARNING] No se encontró la tabla de obligaciones generales en el template")
            print("[INFO] Listando todas las tablas disponibles...")
            for idx, tabla in enumerate(doc.tables):
                primera_fila = tabla.rows[0] if tabla.rows else None
                encabezados = [celda.text.strip().upper() for celda in primera_fila.cells] if primera_fila else []
                print(f"[INFO] Tabla {idx}: {len(tabla.columns)} columnas, {len(tabla.rows)} filas, encabezados: {encabezados}")
            
            print("[INFO] Intentando usar la primera tabla con más de 5 columnas que no haya sido procesada...")
            # Último recurso: usar la primera tabla grande
            for idx, tabla in enumerate(doc.tables):
                if idx not in tablas_procesadas and len(tabla.columns) >= 5:
                    tabla_encontrada = tabla
                    tabla_idx = idx
                    print(f"[INFO] Usando tabla con {len(tabla.columns)} columnas y {len(tabla.rows)} filas (índice {idx})")
                    break
        
        if tabla_encontrada:
            # Verificar que sea la tabla correcta antes de reemplazar
            if len(tabla_encontrada.rows) > 0:
                primera_fila = tabla_encontrada.rows[0]
                encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                tiene_obligacion = any('OBLIGACION' in h or 'OBLIGACIÓN' in h for h in encabezados)
                
                if not (tiene_item and tiene_obligacion):
                    print(f"[WARNING] La tabla encontrada no parece ser la tabla de obligaciones generales")
                    print(f"[WARNING] Encabezados: {encabezados}")
                    print(f"[WARNING] Buscando tabla alternativa...")
                    tabla_encontrada = None
                    tabla_idx = None
                    # Buscar por formato específico
                    for k, tabla in enumerate(doc.tables):
                        if k not in tablas_procesadas and len(tabla.rows) > 0:
                            primera_fila = tabla.rows[0]
                            encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                            tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                            tiene_obligacion = any('OBLIGACION' in h or 'OBLIGACIÓN' in h for h in encabezados)
                            tiene_periodicidad = any('PERIODICIDAD' in h for h in encabezados)
                            # Asegurar que NO es la tabla de anexos (que tiene menos columnas o formato diferente)
                            tiene_anexo_solo = any('ANEXO' in h and len(tabla.columns) < 6 for h in encabezados)
                            if tiene_item and tiene_obligacion and tiene_periodicidad and not tiene_anexo_solo:
                                tabla_encontrada = tabla
                                tabla_idx = k
                                print(f"[INFO] Tabla correcta encontrada por formato (índice {k})")
                                break
        
        if tabla_encontrada and tabla_idx is not None:
            # Verificar que esta tabla no haya sido procesada antes
            if tabla_idx in tablas_procesadas:
                print(f"[WARNING] La tabla {tabla_idx} ya fue procesada, buscando otra...")
                # Buscar otra tabla que no haya sido procesada
                tabla_encontrada = None
                tabla_idx = None
                for k, doc_table in enumerate(doc.tables):
                    if k not in tablas_procesadas and len(doc_table.columns) == 6:
                        primera_fila = doc_table.rows[0] if doc_table.rows else None
                        if primera_fila:
                            encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                            tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                            tiene_obligacion = any('OBLIGACION' in h or 'OBLIGACIÓN' in h for h in encabezados)
                            tiene_periodicidad = any('PERIODICIDAD' in h for h in encabezados)
                            if tiene_item and tiene_obligacion and tiene_periodicidad:
                                tabla_encontrada = doc_table
                                tabla_idx = k
                                print(f"[INFO] Usando tabla alternativa (índice {k}) para obligaciones generales")
                                break
            
            if tabla_idx is not None and tabla_idx not in tablas_procesadas:
                # Reemplazar la tabla encontrada
                self._crear_tabla_obligaciones_generales(doc, tabla_encontrada)
                tablas_procesadas.add(tabla_idx)
                print(f"[INFO] Tabla de obligaciones generales procesada (índice {tabla_idx})")
            elif tabla_idx is None:
                print("[ERROR] No se pudo determinar el índice de la tabla encontrada")
            else:
                print("[ERROR] No se pudo encontrar una tabla disponible para obligaciones generales")
        else:
            print("[ERROR] No se pudo encontrar ninguna tabla para reemplazar")
            print("[ERROR] Asegúrate de que el template tenga una tabla con encabezados: ÍTEM, OBLIGACIÓN, PERIODICIDAD, etc.")
    
    def _crear_tabla_obligaciones_generales(self, doc: Document, tabla_existente) -> None:
        """
        Reemplaza el contenido de la tabla existente con los datos dinámicos de obligaciones generales
        """
        print(f"[INFO] Reemplazando tabla de obligaciones generales con {len(self.obligaciones_generales_raw)} obligaciones")
        
        # Limpiar todas las filas excepto el encabezado (fila 0)
        num_filas_originales = len(tabla_existente.rows)
        while len(tabla_existente.rows) > 1:
            tbl = tabla_existente._tbl
            tbl.remove(tabla_existente.rows[-1]._tr)
        
        print(f"[INFO] Tabla limpiada: {num_filas_originales} filas -> {len(tabla_existente.rows)} fila(s) (encabezado)")
        
        # Obtener número de columnas
        num_cols = len(tabla_existente.columns)
        print(f"[INFO] Tabla tiene {num_cols} columnas")
        
        # Verificar encabezados de la tabla para entender su estructura
        if len(tabla_existente.rows) > 0:
            encabezados = [celda.text.strip().upper() for celda in tabla_existente.rows[0].cells]
            print(f"[INFO] Encabezados de la tabla: {encabezados}")
        
        # Si la tabla tiene menos de 6 columnas, intentar agregar columnas faltantes
        if num_cols < 6:
            print(f"[WARNING] La tabla tiene solo {num_cols} columnas, se necesitan 6. Intentando agregar columnas...")
            # Agregar columnas faltantes
            columnas_faltantes = 6 - num_cols
            for i in range(columnas_faltantes):
                tabla_existente.add_column(Inches(1.5))
                print(f"[INFO] Columna {num_cols + i + 1} agregada")
            num_cols = len(tabla_existente.columns)
            print(f"[INFO] Tabla ahora tiene {num_cols} columnas")
        
        # Actualizar encabezados si es necesario
        if len(tabla_existente.rows) > 0:
            encabezados_esperados = ['ÍTEM', 'OBLIGACIÓN', 'PERIODICIDAD', 'CUMPLIÓ / NO CUMPLIÓ', 'OBSERVACIONES', 'ANEXO']
            primera_fila = tabla_existente.rows[0]
            for i in range(min(num_cols, 6)):
                if i < len(primera_fila.cells):
                    celda = primera_fila.cells[i]
                    texto_actual = celda.text.strip().upper()
                    texto_esperado = encabezados_esperados[i].upper()
                    if texto_actual != texto_esperado and len(texto_actual) < 5:  # Solo actualizar si está vacío o muy corto
                        celda.text = encabezados_esperados[i]
                        # Formatear encabezado
                        for parrafo in celda.paragraphs:
                            for run in parrafo.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)
        
        # Mapeo de columnas
        columnas_esperadas = ['item', 'obligacion', 'periodicidad', 'cumplio', 'observaciones', 'anexo']
        mapeo_columnas = {
            0: 'item',
            1: 'obligacion', 
            2: 'periodicidad',
            3: 'cumplio',
            4: 'observaciones',
            5: 'anexo'
        }
        
        # Agregar filas con los datos
        for obligacion in self.obligaciones_generales_raw:
            fila = tabla_existente.add_row()
            celdas = fila.cells
            
            # Llenar cada celda según el mapeo
            for i in range(min(num_cols, 6)):
                campo = mapeo_columnas.get(i, '')
                valor = obligacion.get(campo, '')
                
                if i < len(celdas):
                    # Limpiar contenido existente de la celda
                    celdas[i].text = ''
                    # Agregar el nuevo texto
                    parrafo = celdas[i].paragraphs[0] if celdas[i].paragraphs else celdas[i].add_paragraph()
                    parrafo.clear()
                    run = parrafo.add_run(str(valor) if valor else '')
                    
                    # Aplicar formato según la columna
                    if i == 0:  # ÍTEM
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    elif i == 1:  # OBLIGACIÓN
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(10)
                    elif i == 2:  # PERIODICIDAD
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    elif i == 3:  # CUMPLIÓ
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    elif i == 4:  # OBSERVACIONES
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(9)
                    elif i == 5:  # ANEXO
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(9)
                    
                    run.font.name = 'Calibri'
        
        print(f"[INFO] Tabla actualizada: {len(tabla_existente.rows)} filas totales (1 encabezado + {len(self.obligaciones_generales_raw)} datos)")
    
    def _reemplazar_tabla_obligaciones_especificas(self, doc: Document, tablas_procesadas: set = None) -> None:
        """
        Busca y reemplaza la tabla de obligaciones específicas con datos dinámicos
        Similar a _reemplazar_tabla_obligaciones_generales pero busca "1.5.2" o "OBLIGACIONES ESPECÍFICAS"
        """
        if tablas_procesadas is None:
            tablas_procesadas = set()
        
        print(f"[DEBUG] _reemplazar_tabla_obligaciones_especificas: {len(self.obligaciones_especificas_raw)} elementos")
        if not self.obligaciones_especificas_raw:
            print("[WARNING] No hay obligaciones específicas para reemplazar en la tabla")
            print("[DEBUG] obligaciones_especificas_raw está vacío")
            return
        
        tabla_encontrada = None
        tabla_idx = None
        
        # Estrategia: Buscar el párrafo "1.5.2" o "OBLIGACIONES ESPECÍFICAS" y luego la siguiente tabla con 6 columnas
        encontro_titulo = False
        elementos = doc.element.body
        
        for i, elemento in enumerate(elementos):
            if hasattr(elemento, 'text') and elemento.text:
                texto = elemento.text.strip().upper()
                if ('1.5.2' in texto or '1.5.2.' in texto) and ('OBLIGACIONES' in texto and 'ESPECÍFICAS' in texto):
                    print(f"[INFO] Título '1.5.2. OBLIGACIONES ESPECÍFICAS' encontrado en elemento {i}")
                    encontro_titulo = True
                    
                    # Contar cuántas tablas hay antes de este título
                    tablas_antes = sum(1 for x in elementos[:i] if hasattr(x, 'tag') and x.tag.endswith('}tbl'))
                    print(f"[INFO] Tablas antes del título: {tablas_antes}")
                    
                    # Tomar la primera tabla que aparece después del título y que no haya sido procesada
                    # Como todas las tablas tienen el mismo formato, simplemente tomamos la siguiente disponible
                    for k, doc_table in enumerate(doc.tables):
                        if k not in tablas_procesadas and k >= tablas_antes:
                            # Verificar que tenga el formato correcto (6 columnas con encabezados de obligaciones)
                            if len(doc_table.columns) == 6 and len(doc_table.rows) > 0:
                                primera_fila = doc_table.rows[0]
                                encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                                tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                                tiene_obligacion = any('OBLIGACION' in h or 'OBLIGACIÓN' in h for h in encabezados)
                                if tiene_item and tiene_obligacion:
                                    tabla_encontrada = doc_table
                                    tabla_idx = k
                                    print(f"[INFO] Tabla de obligaciones específicas encontrada (índice {k})")
                                    break
                    break
        
        # Si no encontramos por título, buscar la primera tabla disponible con formato correcto
        if not tabla_encontrada:
            print("[INFO] No se encontró el título, buscando primera tabla disponible con formato de obligaciones...")
            for k, doc_table in enumerate(doc.tables):
                if k not in tablas_procesadas and len(doc_table.columns) == 6 and len(doc_table.rows) > 0:
                    primera_fila = doc_table.rows[0]
                    encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                    tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                    tiene_obligacion = any('OBLIGACION' in h or 'OBLIGACIÓN' in h for h in encabezados)
                    if tiene_item and tiene_obligacion:
                        tabla_encontrada = doc_table
                        tabla_idx = k
                        print(f"[INFO] Tabla de obligaciones específicas encontrada por formato (índice {k})")
                        break
        
        if tabla_encontrada and tabla_idx is not None:
            # Verificar que esta tabla no haya sido procesada antes
            if tabla_idx in tablas_procesadas:
                print(f"[WARNING] La tabla {tabla_idx} ya fue procesada, buscando otra...")
                tabla_encontrada = None
                tabla_idx = None
            
            if tabla_encontrada and tabla_idx is not None and tabla_idx not in tablas_procesadas:
                self._crear_tabla_obligaciones_especificas(doc, tabla_encontrada)
                tablas_procesadas.add(tabla_idx)
                print(f"[INFO] Tabla de obligaciones específicas procesada (índice {tabla_idx})")
        
        if not tabla_encontrada or (tabla_encontrada and tabla_idx is not None and tabla_idx in tablas_procesadas):
            print("[WARNING] No se encontró la tabla de obligaciones específicas en el template con 6 columnas.")
            print("[INFO] Intentando buscar la primera tabla con 6 columnas que no haya sido procesada...")
            for k, doc_table in enumerate(doc.tables):
                if k not in tablas_procesadas and len(doc_table.columns) == 6:
                    primera_fila = doc_table.rows[0] if doc_table.rows else None
                    if primera_fila:
                        encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                        tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                        tiene_obligacion = any('OBLIGACION' in h or 'OBLIGACIÓN' in h for h in encabezados)
                        if tiene_item and tiene_obligacion:
                            tabla_encontrada = doc_table
                            tabla_idx = k
                            print(f"[INFO] Usando tabla con 6 columnas (índice {k}) para obligaciones específicas")
                            self._crear_tabla_obligaciones_especificas(doc, tabla_encontrada)
                            tablas_procesadas.add(k)
                            break
            
            if not tabla_encontrada:
                print("[ERROR] No se pudo encontrar ninguna tabla para reemplazar las obligaciones específicas.")
    
    def _crear_tabla_obligaciones_especificas(self, doc: Document, tabla_existente) -> None:
        """
        Reemplaza el contenido de la tabla existente con los datos dinámicos de obligaciones específicas
        """
        print(f"[INFO] Reemplazando tabla de obligaciones específicas con {len(self.obligaciones_especificas_raw)} obligaciones")
        
        # Limpiar todas las filas excepto el encabezado (fila 0)
        num_filas_originales = len(tabla_existente.rows)
        while len(tabla_existente.rows) > 1:
            tbl = tabla_existente._tbl
            tbl.remove(tabla_existente.rows[-1]._tr)
        
        print(f"[INFO] Tabla limpiada: {num_filas_originales} filas -> {len(tabla_existente.rows)} fila(s) (encabezado)")
        
        # Obtener número de columnas
        num_cols = len(tabla_existente.columns)
        print(f"[INFO] Tabla tiene {num_cols} columnas")
        
        # Encabezados esperados
        encabezados_esperados = ["ÍTEM", "OBLIGACIÓN", "PERIODICIDAD", "CUMPLIÓ / NO CUMPLIÓ", "OBSERVACIONES", "ANEXO"]
        
        # Actualizar encabezados si no coinciden
        if len(tabla_existente.rows) > 0:
            header_cells = tabla_existente.rows[0].cells
            current_headers = [self._limpiar_texto_celda(c.text) for c in header_cells]
            print(f"[INFO] Encabezados de la tabla: {current_headers}")
            
            if current_headers != encabezados_esperados:
                print("[INFO] Los encabezados de la tabla no coinciden, actualizando...")
                for i, header_text in enumerate(encabezados_esperados):
                    if i < num_cols:
                        self._formatear_celda(header_cells[i], header_text, bold=True, center=True)
                    else:
                        print(f"[WARNING] No hay suficientes columnas para el encabezado: {header_text}")
        
        # Agregar datos
        for obligacion in self.obligaciones_especificas_raw:
            row_cells = tabla_existente.add_row().cells
            self._formatear_celda(row_cells[0], str(obligacion.get("item", "")))
            self._formatear_celda(row_cells[1], obligacion.get("obligacion", ""))
            self._formatear_celda(row_cells[2], obligacion.get("periodicidad", ""), center=True)
            self._formatear_celda(row_cells[3], obligacion.get("cumplio", ""), center=True)
            self._formatear_celda(row_cells[4], obligacion.get("observaciones", ""))
            self._formatear_celda(row_cells[5], obligacion.get("anexo", ""))
        
        print(f"[INFO] Tabla actualizada: {len(tabla_existente.rows)} filas totales (1 encabezado + {len(self.obligaciones_especificas_raw)} datos)")
    
    def _formatear_celda(self, cell, text, bold=False, center=False):
        """Formatea el texto de una celda"""
        cell.text = text
        paragraph = cell.paragraphs[0]
        if center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = bold
            run.font.size = Pt(10)  # Tamaño de fuente consistente
    
    def _limpiar_texto_celda(self, text: str) -> str:
        """Limpia el texto de una celda para comparación (elimina saltos de línea y espacios extra)"""
        return ' '.join(text.replace('\n', ' ').split()).upper()
    
    def _reemplazar_tabla_obligaciones_ambientales(self, doc: Document, tablas_procesadas: set = None) -> None:
        """
        Busca y reemplaza la tabla de obligaciones ambientales con datos dinámicos
        Similar a _reemplazar_tabla_obligaciones_generales pero busca "1.5.3" o "OBLIGACIONES AMBIENTALES"
        """
        if tablas_procesadas is None:
            tablas_procesadas = set()
        
        print(f"[DEBUG] _reemplazar_tabla_obligaciones_ambientales: {len(self.obligaciones_ambientales_raw)} elementos")
        if not self.obligaciones_ambientales_raw:
            print("[WARNING] No hay obligaciones ambientales para reemplazar en la tabla")
            print("[DEBUG] obligaciones_ambientales_raw está vacío")
            return
        
        tabla_encontrada = None
        tabla_idx = None
        
        # Estrategia: Buscar el título "1.5.3" y encontrar la primera tabla que aparece INMEDIATAMENTE después en el XML
        elementos = doc.element.body
        tabla_count_before = 0
        
        for i, elemento in enumerate(elementos):
            # Contar tablas antes del título
            if hasattr(elemento, 'tag') and elemento.tag.endswith('}tbl'):
                tabla_count_before += 1
            
            # Buscar el título
            if hasattr(elemento, 'text') and elemento.text:
                texto = elemento.text.strip().upper()
                if ('1.5.3' in texto or '1.5.3.' in texto) and ('OBLIGACIONES' in texto and 'AMBIENTALES' in texto):
                    print(f"[INFO] Título '1.5.3. OBLIGACIONES AMBIENTALES' encontrado en elemento {i}")
                    print(f"[INFO] Tablas antes del título: {tabla_count_before}")
                    
                    # Buscar la primera tabla que aparece DESPUÉS del título en el XML
                    for j in range(i + 1, len(elementos)):
                        siguiente = elementos[j]
                        if hasattr(siguiente, 'tag') and siguiente.tag.endswith('}tbl'):
                            # Esta es la primera tabla después del título
                            tabla_index = tabla_count_before
                            if tabla_index < len(doc.tables):
                                tabla_candidata = doc.tables[tabla_index]
                                # Verificar que no haya sido procesada y que tenga el formato correcto
                                if tabla_index not in tablas_procesadas and len(tabla_candidata.columns) == 6 and len(tabla_candidata.rows) > 0:
                                    primera_fila = tabla_candidata.rows[0]
                                    encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                                    tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                                    tiene_obligacion = any('OBLIGACION' in h or 'OBLIGACIÓN' in h for h in encabezados)
                                    if tiene_item and tiene_obligacion:
                                        tabla_encontrada = tabla_candidata
                                        tabla_idx = tabla_index
                                        print(f"[INFO] Tabla de obligaciones ambientales encontrada (índice {tabla_index})")
                                        break
                            break
                    break
        
        # Si no encontramos por título, buscar la primera tabla disponible con formato correcto
        if not tabla_encontrada:
            print("[INFO] No se encontró el título, buscando primera tabla disponible con formato de obligaciones...")
            for k, doc_table in enumerate(doc.tables):
                if k not in tablas_procesadas and len(doc_table.columns) == 6 and len(doc_table.rows) > 0:
                    primera_fila = doc_table.rows[0]
                    encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                    tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                    tiene_obligacion = any('OBLIGACION' in h or 'OBLIGACIÓN' in h for h in encabezados)
                    if tiene_item and tiene_obligacion:
                        tabla_encontrada = doc_table
                        tabla_idx = k
                        print(f"[INFO] Tabla de obligaciones ambientales encontrada por formato (índice {k})")
                        break
        
        if tabla_encontrada and tabla_idx is not None:
            # Verificar que esta tabla no haya sido procesada antes
            if tabla_idx not in tablas_procesadas:
                # Verificar que tenemos datos de obligaciones ambientales
                if not self.obligaciones_ambientales_raw:
                    print(f"[ERROR] No hay datos de obligaciones_ambientales_raw para llenar la tabla")
                    return
                print(f"[DEBUG] ANTES de crear tabla: obligaciones_ambientales_raw tiene {len(self.obligaciones_ambientales_raw)} elementos")
                self._crear_tabla_obligaciones_ambientales(doc, tabla_encontrada)
                tablas_procesadas.add(tabla_idx)
                print(f"[INFO] Tabla de obligaciones ambientales procesada (índice {tabla_idx})")
            else:
                print(f"[WARNING] La tabla {tabla_idx} ya fue procesada, buscando otra...")
                tabla_encontrada = None
                tabla_idx = None
        
        if not tabla_encontrada:
            print("[ERROR] No se pudo encontrar ninguna tabla para reemplazar las obligaciones ambientales.")
            print("[INFO] Intentando buscar cualquier tabla con 6 columnas que no haya sido procesada...")
            for k, doc_table in enumerate(doc.tables):
                if k not in tablas_procesadas and len(doc_table.columns) == 6:
                    primera_fila = doc_table.rows[0] if doc_table.rows else None
                    if primera_fila:
                        encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                        tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                        tiene_obligacion = any('OBLIGACION' in h or 'OBLIGACIÓN' in h for h in encabezados)
                        if tiene_item and tiene_obligacion:
                            tabla_encontrada = doc_table
                            self._crear_tabla_obligaciones_ambientales(doc, tabla_encontrada)
                            tablas_procesadas.add(k)
                            print(f"[INFO] Tabla de obligaciones ambientales procesada (índice {k})")
                            break
    
    def _crear_tabla_obligaciones_ambientales(self, doc: Document, tabla_existente) -> None:
        """
        Reemplaza el contenido de la tabla existente con los datos dinámicos de obligaciones ambientales
        """
        print(f"[INFO] Reemplazando tabla de obligaciones ambientales con {len(self.obligaciones_ambientales_raw)} obligaciones")
        print(f"[DEBUG] Datos de obligaciones_ambientales_raw: {self.obligaciones_ambientales_raw[:2] if self.obligaciones_ambientales_raw else 'VACÍO'}")
        
        # Limpiar todas las filas excepto el encabezado (fila 0)
        num_filas_originales = len(tabla_existente.rows)
        while len(tabla_existente.rows) > 1:
            tbl = tabla_existente._tbl
            tbl.remove(tabla_existente.rows[-1]._tr)
        
        print(f"[INFO] Tabla limpiada: {num_filas_originales} filas -> {len(tabla_existente.rows)} fila(s) (encabezado)")
        
        # Obtener número de columnas
        num_cols = len(tabla_existente.columns)
        print(f"[INFO] Tabla tiene {num_cols} columnas")
        
        # Encabezados esperados
        encabezados_esperados = ["ÍTEM", "OBLIGACIÓN", "PERIODICIDAD", "CUMPLIÓ / NO CUMPLIÓ", "OBSERVACIONES", "ANEXO"]
        
        # Actualizar encabezados si no coinciden
        if len(tabla_existente.rows) > 0:
            header_cells = tabla_existente.rows[0].cells
            current_headers = [self._limpiar_texto_celda(c.text) for c in header_cells]
            print(f"[INFO] Encabezados de la tabla: {current_headers}")
            
            if current_headers != encabezados_esperados:
                print("[INFO] Los encabezados de la tabla no coinciden, actualizando...")
                for i, header_text in enumerate(encabezados_esperados):
                    if i < num_cols:
                        self._formatear_celda(header_cells[i], header_text, bold=True, center=True)
                    else:
                        print(f"[WARNING] No hay suficientes columnas para el encabezado: {header_text}")
        
        # Agregar datos
        for obligacion in self.obligaciones_ambientales_raw:
            row_cells = tabla_existente.add_row().cells
            self._formatear_celda(row_cells[0], str(obligacion.get("item", "")))
            self._formatear_celda(row_cells[1], obligacion.get("obligacion", ""))
            self._formatear_celda(row_cells[2], obligacion.get("periodicidad", ""), center=True)
            self._formatear_celda(row_cells[3], obligacion.get("cumplio", ""), center=True)
            self._formatear_celda(row_cells[4], obligacion.get("observaciones", ""))
            self._formatear_celda(row_cells[5], obligacion.get("anexo", ""))
        
        print(f"[INFO] Tabla actualizada: {len(tabla_existente.rows)} filas totales (1 encabezado + {len(self.obligaciones_ambientales_raw)} datos)")
    
    def _reemplazar_tabla_obligaciones_anexos(self, doc: Document, tablas_procesadas: set = None) -> None:
        """
        Busca y reemplaza la tabla de obligaciones anexos con datos dinámicos
        Similar a _reemplazar_tabla_obligaciones_generales pero busca "1.5.4" o "OBLIGACIONES ANEXOS"
        """
        if tablas_procesadas is None:
            tablas_procesadas = set()
        
        if not self.obligaciones_anexos_raw:
            print("[WARNING] No hay obligaciones anexos para reemplazar en la tabla")
            return
        
        tabla_encontrada = None
        tabla_idx = None
        
        # Estrategia: Buscar el párrafo "1.5.4" o "OBLIGACIONES ANEXOS" y luego la siguiente tabla
        elementos = doc.element.body
        
        for i, elemento in enumerate(elementos):
            if hasattr(elemento, 'text') and elemento.text:
                texto = elemento.text.strip().upper()
                if ('1.5.4' in texto or '1.5.4.' in texto) and ('OBLIGACIONES' in texto and 'ANEXOS' in texto):
                    print(f"[INFO] Título '1.5.4. OBLIGACIONES ANEXOS' encontrado en elemento {i}")
                    
                    # Contar cuántas tablas hay antes de este título
                    tablas_antes = sum(1 for x in elementos[:i] if hasattr(x, 'tag') and x.tag.endswith('}tbl'))
                    print(f"[INFO] Tablas antes del título: {tablas_antes}")
                    
                    # Tomar la primera tabla que aparece después del título y que no haya sido procesada
                    # Para anexos, puede tener menos columnas o formato diferente
                    for k, doc_table in enumerate(doc.tables):
                        if k not in tablas_procesadas and k >= tablas_antes:
                            if len(doc_table.columns) >= 3 and len(doc_table.rows) > 0:  # Mínimo 3 columnas
                                primera_fila = doc_table.rows[0]
                                encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                                tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                                tiene_anexo = any('ANEXO' in h for h in encabezados)
                                if tiene_item or tiene_anexo:
                                    tabla_encontrada = doc_table
                                    tabla_idx = k
                                    print(f"[INFO] Tabla de obligaciones anexos encontrada (índice {k})")
                                    break
                    break
        
        # Estrategia 2: Si no encontramos por título, buscar tabla con formato correcto
        if not tabla_encontrada:
            print("[INFO] Intentando buscar tabla de obligaciones anexos por formato...")
            
            for k, doc_table in enumerate(doc.tables):
                if k not in tablas_procesadas and len(doc_table.rows) > 0:
                    primera_fila = doc_table.rows[0]
                    encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                    tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                    tiene_anexo = any('ANEXO' in h for h in encabezados)
                    
                    # Si tiene anexo o item y no ha sido procesada
                    if tiene_item or tiene_anexo:
                        tabla_encontrada = doc_table
                        tabla_idx = k
                        print(f"[INFO] Tabla de obligaciones anexos encontrada por formato (índice {k})")
                        break
        
        if tabla_encontrada and tabla_idx is not None:
            # Verificar que esta tabla no haya sido procesada antes
            if tabla_idx not in tablas_procesadas:
                self._crear_tabla_obligaciones_anexos(doc, tabla_encontrada)
                tablas_procesadas.add(tabla_idx)
                print(f"[INFO] Tabla de obligaciones anexos procesada (índice {tabla_idx})")
            else:
                print(f"[WARNING] La tabla {tabla_idx} ya fue procesada, buscando otra...")
                tabla_encontrada = None
                tabla_idx = None
        
        if not tabla_encontrada:
            print("[ERROR] No se pudo encontrar ninguna tabla para reemplazar las obligaciones anexos.")
            print("[INFO] Intentando buscar cualquier tabla que no haya sido procesada...")
            for k, doc_table in enumerate(doc.tables):
                if k not in tablas_procesadas:
                    primera_fila = doc_table.rows[0] if doc_table.rows else None
                    if primera_fila:
                        encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                        tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                        tiene_anexo = any('ANEXO' in h for h in encabezados)
                        if tiene_item or tiene_anexo:
                            tabla_encontrada = doc_table
                            self._crear_tabla_obligaciones_anexos(doc, tabla_encontrada)
                            tablas_procesadas.add(k)
                            print(f"[INFO] Tabla de obligaciones anexos procesada (índice {k})")
                            break
    
    def _crear_tabla_obligaciones_anexos(self, doc: Document, tabla_existente) -> None:
        """
        Reemplaza el contenido de la tabla existente con los datos dinámicos de obligaciones anexos
        Nota: Para anexos, el formato puede ser diferente (solo archivo_existe y anexo)
        """
        print(f"[INFO] Reemplazando tabla de obligaciones anexos con {len(self.obligaciones_anexos_raw)} elementos")
        
        # Limpiar todas las filas excepto el encabezado (fila 0)
        num_filas_originales = len(tabla_existente.rows)
        while len(tabla_existente.rows) > 1:
            tbl = tabla_existente._tbl
            tbl.remove(tabla_existente.rows[-1]._tr)
        
        print(f"[INFO] Tabla limpiada: {num_filas_originales} filas -> {len(tabla_existente.rows)} fila(s) (encabezado)")
        
        # Obtener número de columnas
        num_cols = len(tabla_existente.columns)
        print(f"[INFO] Tabla tiene {num_cols} columnas")
        
        # Para anexos, el formato puede variar. Intentar adaptarse al formato existente
        if len(tabla_existente.rows) > 0:
            header_cells = tabla_existente.rows[0].cells
            current_headers = [self._limpiar_texto_celda(c.text) for c in header_cells]
            print(f"[INFO] Encabezados de la tabla: {current_headers}")
        
        # Agregar datos
        # Si el formato es simple (archivo_existe, anexo), usar ese formato
        # Si no, usar el formato estándar de obligaciones
        for idx, anexo_data in enumerate(self.obligaciones_anexos_raw, start=1):
            row_cells = tabla_existente.add_row().cells
            
            # Verificar si es el formato simple de anexos (archivo_existe, anexo)
            if "archivo_existe" in anexo_data:
                # Formato simple para anexos
                if num_cols >= 1:
                    self._formatear_celda(row_cells[0], str(idx))
                if num_cols >= 2:
                    estado = "Sí" if anexo_data.get("archivo_existe", False) else "No"
                    self._formatear_celda(row_cells[1], estado, center=True)
                if num_cols >= 3:
                    self._formatear_celda(row_cells[2], anexo_data.get("anexo", ""))
            else:
                # Formato estándar de obligaciones
                if num_cols >= 1:
                    self._formatear_celda(row_cells[0], str(anexo_data.get("item", idx)))
                if num_cols >= 2:
                    self._formatear_celda(row_cells[1], anexo_data.get("obligacion", anexo_data.get("anexo", "")))
                if num_cols >= 3:
                    self._formatear_celda(row_cells[2], anexo_data.get("periodicidad", ""), center=True)
                if num_cols >= 4:
                    self._formatear_celda(row_cells[3], anexo_data.get("cumplio", ""), center=True)
                if num_cols >= 5:
                    self._formatear_celda(row_cells[4], anexo_data.get("observaciones", ""))
                if num_cols >= 6:
                    self._formatear_celda(row_cells[5], anexo_data.get("anexo", ""))
        
        print(f"[INFO] Tabla actualizada: {len(tabla_existente.rows)} filas totales (1 encabezado + {len(self.obligaciones_anexos_raw)} datos)")
    
    def _reemplazar_tabla_comunicados_emitidos(self, doc: Document, tablas_procesadas: set = None) -> None:
        """
        Busca y reemplaza la tabla de comunicados emitidos con datos dinámicos
        Busca "1.6.1" o "EMITIDOS CONTRATO"
        """
        if tablas_procesadas is None:
            tablas_procesadas = set()
        
        if not self.comunicados_emitidos:
            print("[WARNING] No hay comunicados emitidos para reemplazar en la tabla")
            return
        
        tabla_encontrada = None
        tabla_idx = None
        
        # Estrategia: Buscar el título "1.6.1" y encontrar la primera tabla que aparece INMEDIATAMENTE después en el XML
        elementos = doc.element.body
        tabla_count_before = 0
        
        for i, elemento in enumerate(elementos):
            # Contar tablas antes del título
            if hasattr(elemento, 'tag') and elemento.tag.endswith('}tbl'):
                tabla_count_before += 1
            
            # Buscar el título
            if hasattr(elemento, 'text') and elemento.text:
                texto = elemento.text.strip().upper()
                if ('1.6.1' in texto or '1.6.1.' in texto) and ('EMITIDOS' in texto or 'COMUNICADOS' in texto):
                    print(f"[INFO] Título '1.6.1. EMITIDOS' encontrado en elemento {i}")
                    print(f"[INFO] Tablas antes del título: {tabla_count_before}")
                    
                    # Buscar la primera tabla que aparece DESPUÉS del título en el XML
                    for j in range(i + 1, len(elementos)):
                        siguiente = elementos[j]
                        if hasattr(siguiente, 'tag') and siguiente.tag.endswith('}tbl'):
                            # Esta es la primera tabla después del título
                            tabla_index = tabla_count_before
                            if tabla_index < len(doc.tables):
                                tabla_candidata = doc.tables[tabla_index]
                                # Verificar que no haya sido procesada y que tenga el formato correcto de comunicados
                                if tabla_index not in tablas_procesadas and len(tabla_candidata.columns) >= 4 and len(tabla_candidata.rows) > 0:
                                    primera_fila = tabla_candidata.rows[0]
                                    encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                                    tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                                    tiene_fecha = any('FECHA' in h for h in encabezados)
                                    tiene_consecutivo = any('CONSECUTIVO' in h or 'RADICADO' in h for h in encabezados)
                                    tiene_asunto = any('ASUNTO' in h or 'DESCRIPCIÓN' in h or 'DESCRIPCION' in h for h in encabezados)
                                    # Asegurar que NO es una tabla de obligaciones (que tiene 6 columnas)
                                    if tiene_item and tiene_fecha and tiene_consecutivo and tiene_asunto and len(tabla_candidata.columns) < 6:
                                        tabla_encontrada = tabla_candidata
                                        tabla_idx = tabla_index
                                        print(f"[INFO] Tabla de comunicados emitidos encontrada (índice {tabla_index})")
                                        break
                            break
                    break
        
        # Estrategia 2: Si no encontramos por título, buscar tabla con formato correcto
        if not tabla_encontrada:
            print("[INFO] Intentando buscar tabla de comunicados emitidos por formato...")
            for k, doc_table in enumerate(doc.tables):
                if k not in tablas_procesadas and len(doc_table.rows) > 0 and len(doc_table.columns) >= 4:
                    primera_fila = doc_table.rows[0]
                    encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                    tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                    tiene_fecha = any('FECHA' in h for h in encabezados)
                    tiene_consecutivo = any('CONSECUTIVO' in h or 'RADICADO' in h for h in encabezados)
                    tiene_descripcion = any('DESCRIPCIÓN' in h or 'ASUNTO' in h for h in encabezados)
                    
                    if tiene_item and tiene_fecha and tiene_consecutivo and tiene_descripcion:
                        tabla_encontrada = doc_table
                        tabla_idx = k
                        print(f"[INFO] Tabla de comunicados emitidos encontrada por formato (índice {k})")
                        break
        
        if tabla_encontrada and tabla_idx is not None:
            # Verificar que esta tabla no haya sido procesada antes
            if tabla_idx not in tablas_procesadas:
                # Verificar que tenemos datos de comunicados emitidos
                if not self.comunicados_emitidos:
                    print(f"[ERROR] No hay datos de comunicados_emitidos para llenar la tabla")
                    return
                print(f"[DEBUG] ANTES de crear tabla: comunicados_emitidos tiene {len(self.comunicados_emitidos)} elementos")
                print(f"[DEBUG] Verificando que NO esté usando obligaciones_anexos_raw (tiene {len(self.obligaciones_anexos_raw)} elementos)")
                self._crear_tabla_comunicados_emitidos(doc, tabla_encontrada)
                tablas_procesadas.add(tabla_idx)
                print(f"[INFO] Tabla de comunicados emitidos procesada (índice {tabla_idx})")
            else:
                print(f"[WARNING] La tabla {tabla_idx} ya fue procesada, buscando otra...")
                tabla_encontrada = None
                tabla_idx = None
        
        if not tabla_encontrada:
            print("[ERROR] No se pudo encontrar ninguna tabla para reemplazar los comunicados emitidos.")
            print("[INFO] Intentando buscar cualquier tabla con formato de comunicados que no haya sido procesada...")
            for k, doc_table in enumerate(doc.tables):
                if k not in tablas_procesadas and len(doc_table.columns) >= 4:
                    primera_fila = doc_table.rows[0] if doc_table.rows else None
                    if primera_fila:
                        encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                        tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                        tiene_fecha = any('FECHA' in h for h in encabezados)
                        tiene_consecutivo = any('CONSECUTIVO' in h or 'RADICADO' in h for h in encabezados)
                        if tiene_item and tiene_fecha and tiene_consecutivo:
                            tabla_encontrada = doc_table
                            self._crear_tabla_comunicados_emitidos(doc, tabla_encontrada)
                            tablas_procesadas.add(k)
                            print(f"[INFO] Tabla de comunicados emitidos procesada (índice {k})")
                            break
    
    def _crear_tabla_comunicados_emitidos(self, doc: Document, tabla_existente) -> None:
        """
        Reemplaza el contenido de la tabla existente con los datos dinámicos de comunicados emitidos
        """
        print(f"[INFO] Reemplazando tabla de comunicados emitidos con {len(self.comunicados_emitidos)} comunicados")
        print(f"[DEBUG] Datos de comunicados_emitidos: {self.comunicados_emitidos[:2] if self.comunicados_emitidos else 'VACÍO'}")
        print(f"[DEBUG] Verificando que NO esté usando obligaciones_anexos_raw: {len(self.obligaciones_anexos_raw) if hasattr(self, 'obligaciones_anexos_raw') else 'N/A'}")
        
        # Limpiar todas las filas excepto el encabezado (fila 0)
        num_filas_originales = len(tabla_existente.rows)
        while len(tabla_existente.rows) > 1:
            tbl = tabla_existente._tbl
            tbl.remove(tabla_existente.rows[-1]._tr)
        
        print(f"[INFO] Tabla limpiada: {num_filas_originales} filas -> {len(tabla_existente.rows)} fila(s) (encabezado)")
        
        # Obtener número de columnas
        num_cols = len(tabla_existente.columns)
        print(f"[INFO] Tabla tiene {num_cols} columnas")
        
        # Agregar datos
        for comunicado in self.comunicados_emitidos:
            row_cells = tabla_existente.add_row().cells
            if num_cols >= 1:
                self._formatear_celda(row_cells[0], str(comunicado.get("item", "")), center=True)
            if num_cols >= 2:
                self._formatear_celda(row_cells[1], comunicado.get("fecha", ""), center=True)
            if num_cols >= 3:
                self._formatear_celda(row_cells[2], comunicado.get("numero", comunicado.get("radicado", "")), center=True)
            if num_cols >= 4:
                self._formatear_celda(row_cells[3], comunicado.get("asunto", ""))
        
        print(f"[INFO] Tabla actualizada: {len(tabla_existente.rows)} filas totales (1 encabezado + {len(self.comunicados_emitidos)} datos)")
    
    def _reemplazar_tabla_comunicados_recibidos(self, doc: Document, tablas_procesadas: set = None) -> None:
        """
        Busca y reemplaza la tabla de comunicados recibidos con datos dinámicos
        Busca "1.6.2" o "RECIBIDOS CONTRATO"
        """
        if tablas_procesadas is None:
            tablas_procesadas = set()
        
        if not self.comunicados_recibidos:
            print("[WARNING] No hay comunicados recibidos para reemplazar en la tabla")
            return
        
        tabla_encontrada = None
        tabla_idx = None
        
        # Estrategia: Buscar el párrafo "1.6.2" o "RECIBIDOS" y luego la siguiente tabla
        elementos = doc.element.body
        
        for i, elemento in enumerate(elementos):
            if hasattr(elemento, 'text') and elemento.text:
                texto = elemento.text.strip().upper()
                if ('1.6.2' in texto or '1.6.2.' in texto) and ('RECIBIDOS' in texto or 'COMUNICADOS' in texto):
                    print(f"[INFO] Título '1.6.2. RECIBIDOS' encontrado en elemento {i}")
                    
                    # Contar cuántas tablas hay antes de este título
                    tablas_antes = sum(1 for x in elementos[:i] if hasattr(x, 'tag') and x.tag.endswith('}tbl'))
                    print(f"[INFO] Tablas antes del título: {tablas_antes}")
                    
                    # Buscar la siguiente tabla después de este párrafo
                    current_table_count = 0
                    for k, doc_table in enumerate(doc.tables):
                        if k not in tablas_procesadas and current_table_count >= tablas_antes:
                            # Esta es una tabla que aparece después del título
                            if len(doc_table.columns) >= 4:  # Buscamos una tabla con al menos 4 columnas
                                if len(doc_table.rows) > 0:
                                    primera_fila = doc_table.rows[0]
                                    encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                                    tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                                    tiene_fecha = any('FECHA' in h for h in encabezados)
                                    tiene_consecutivo = any('CONSECUTIVO' in h or 'RADICADO' in h for h in encabezados)
                                    if tiene_item or (tiene_fecha and tiene_consecutivo):
                                        tabla_encontrada = doc_table
                                        tabla_idx = k
                                        print(f"[INFO] Tabla de comunicados recibidos encontrada (índice {k})")
                                        break
                        if hasattr(doc_table._element, 'tag') and doc_table._element.tag.endswith('}tbl'):
                            current_table_count += 1
                    break
        
        # Estrategia 2: Si no encontramos por título, buscar tabla con formato correcto
        if not tabla_encontrada:
            print("[INFO] Intentando buscar tabla de comunicados recibidos por formato...")
            for k, doc_table in enumerate(doc.tables):
                if k not in tablas_procesadas and len(doc_table.rows) > 0 and len(doc_table.columns) >= 4:
                    primera_fila = doc_table.rows[0]
                    encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                    tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                    tiene_fecha = any('FECHA' in h for h in encabezados)
                    tiene_consecutivo = any('CONSECUTIVO' in h or 'RADICADO' in h for h in encabezados)
                    tiene_descripcion = any('DESCRIPCIÓN' in h or 'ASUNTO' in h for h in encabezados)
                    
                    if tiene_item and tiene_fecha and tiene_consecutivo and tiene_descripcion:
                        tabla_encontrada = doc_table
                        tabla_idx = k
                        print(f"[INFO] Tabla de comunicados recibidos encontrada por formato (índice {k})")
                        break
        
        if tabla_encontrada and tabla_idx is not None:
            # Verificar que esta tabla no haya sido procesada antes
            if tabla_idx not in tablas_procesadas:
                self._crear_tabla_comunicados_recibidos(doc, tabla_encontrada)
                tablas_procesadas.add(tabla_idx)
                print(f"[INFO] Tabla de comunicados recibidos procesada (índice {tabla_idx})")
            else:
                print(f"[WARNING] La tabla {tabla_idx} ya fue procesada, buscando otra...")
                tabla_encontrada = None
                tabla_idx = None
        
        if not tabla_encontrada:
            print("[ERROR] No se pudo encontrar ninguna tabla para reemplazar los comunicados recibidos.")
            print("[INFO] Intentando buscar cualquier tabla con formato de comunicados que no haya sido procesada...")
            for k, doc_table in enumerate(doc.tables):
                if k not in tablas_procesadas and len(doc_table.columns) >= 4:
                    primera_fila = doc_table.rows[0] if doc_table.rows else None
                    if primera_fila:
                        encabezados = [celda.text.strip().upper() for celda in primera_fila.cells]
                        tiene_item = any('ITEM' in h or 'ÍTEM' in h for h in encabezados)
                        tiene_fecha = any('FECHA' in h for h in encabezados)
                        tiene_consecutivo = any('CONSECUTIVO' in h or 'RADICADO' in h for h in encabezados)
                        if tiene_item and tiene_fecha and tiene_consecutivo:
                            tabla_encontrada = doc_table
                            self._crear_tabla_comunicados_recibidos(doc, tabla_encontrada)
                            tablas_procesadas.add(k)
                            print(f"[INFO] Tabla de comunicados recibidos procesada (índice {k})")
                            break
    
    def _crear_tabla_comunicados_recibidos(self, doc: Document, tabla_existente) -> None:
        """
        Reemplaza el contenido de la tabla existente con los datos dinámicos de comunicados recibidos
        """
        print(f"[INFO] Reemplazando tabla de comunicados recibidos con {len(self.comunicados_recibidos)} comunicados")
        
        # Limpiar todas las filas excepto el encabezado (fila 0)
        num_filas_originales = len(tabla_existente.rows)
        while len(tabla_existente.rows) > 1:
            tbl = tabla_existente._tbl
            tbl.remove(tabla_existente.rows[-1]._tr)
        
        print(f"[INFO] Tabla limpiada: {num_filas_originales} filas -> {len(tabla_existente.rows)} fila(s) (encabezado)")
        
        # Obtener número de columnas
        num_cols = len(tabla_existente.columns)
        print(f"[INFO] Tabla tiene {num_cols} columnas")
        
        # Agregar datos
        for comunicado in self.comunicados_recibidos:
            row_cells = tabla_existente.add_row().cells
            if num_cols >= 1:
                self._formatear_celda(row_cells[0], str(comunicado.get("item", "")), center=True)
            if num_cols >= 2:
                self._formatear_celda(row_cells[1], comunicado.get("fecha", ""), center=True)
            if num_cols >= 3:
                self._formatear_celda(row_cells[2], comunicado.get("numero", comunicado.get("radicado", "")), center=True)
            if num_cols >= 4:
                self._formatear_celda(row_cells[3], comunicado.get("asunto", ""))
        
        print(f"[INFO] Tabla actualizada: {len(tabla_existente.rows)} filas totales (1 encabezado + {len(self.comunicados_recibidos)} datos)")


