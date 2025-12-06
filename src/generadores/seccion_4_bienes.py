"""
Generador Sección 4: Informe de Bienes y Servicios
Estructurado siguiendo el patrón de seccion_2_mesa_servicio.py

Subsecciones:
- 4.1 Gestión de Inventario
- 4.2 Entradas Almacén SDSCJ
- 4.3 Entrega Equipos No Operativos Almacén SDSCJ
- 4.4 Gestiones de Inclusión a la Bolsa
"""
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Any, Optional
from datetime import datetime
import calendar
import json
import config
import os
from urllib.parse import urlparse

from docxtpl import DocxTemplate
from docx import Document
from pathlib import Path
from docxtpl import InlineImage
from docx.shared import Mm
import base64
from PIL import Image
from io import BytesIO
import logging
import re
import time

# Importar utilidades
from src.utils.fecha_utils import formatear_fecha_simple, formatear_fechas_en_tabla
from src.utils.tabla_utils import (
    set_cell_shading, centrar_celda_vertical, enable_autofit,
    detectar_columnas_desde_datos, preparar_datos_tabla, count_rows,
    reemplazar_placeholder_con_tabla, reemplazar_multiples_placeholders_con_tablas
)
from src.utils.imagen_utils import base64_to_inline_image, procesar_imagen
from src.utils.documento_utils import convertir_url_sharepoint_a_ruta_relativa
from src.utils.formato_moneda import numero_a_letras, formato_moneda_cop

logger = logging.getLogger(__name__)


class GeneradorSeccion4:
    @property
    def nombre_seccion(self) -> str:
        return "4. INFORME DE BIENES Y SERVICIOS"
    
    @property
    def template_file(self) -> str:
        return "seccion_4_bienes_servicios.docx"
    
    def __init__(self):
        self.datos = {}
        self.doc = None
        self.anio = None
        self.mes = None
    
    def _seccion_4(self, data: Dict[str, Any]):
        """
        Procesa la sección 4 principal (solo datos generales si aplica)
        """
        content_data = {
            "anio": data.get("anio"),
            "mes": data.get("mes"),
            "user_id": data.get("user_id"),
            "name_file": data.get("name_file"),
            "section_id": data.get("section_id"),
            "level": data.get("level"),
            "content": {
                "image": data.get("content", {}).get("image", ""),
            },
        }
        return content_data
    
    async def _seccion_4_1(self, data: Dict[str, Any]):
        """
        Procesa la sección 4.1: Gestión de Inventario
        """
        content = data.get("content", {})
        anio = data.get("anio")
        mes = data.get("mes")
        
        content_data = {
            "anio": anio,
            "mes": mes,
            "user_id": data.get("user_id"),
            "name_file": data.get("name_file"),
            "section_id": data.get("section_id"),
            "level": data.get("level"),
            "content": {
                "texto": content.get("texto", ""),
                "tabla": content.get("tabla", []),
            },
        }
        
        return content_data
    
    async def _seccion_4_2(self, data: Dict[str, Any]):
        """
        Procesa la sección 4.2: Entradas Almacén SDSCJ
        """
        content = data.get("content", {})
        anio = data.get("anio")
        mes = data.get("mes")
        
        # Obtener datos del comunicado
        comunicado = content.get("comunicado", {})
        elementos = content.get("elementos", [])
        
        # Calcular total
        valor_total = 0
        try:
            for elemento in elementos:
                valor = elemento.get("valor_total", 0) or elemento.get("valor", 0)
                if isinstance(valor, (int, float)):
                    valor_total += float(valor)
                elif isinstance(valor, str):
                    valor_limpio = valor.replace(",", "").replace(" ", "").replace("$", "").strip()
                    if valor_limpio:
                        valor_total += float(valor_limpio)
        except (ValueError, TypeError) as e:
            logger.warning(f"Error al calcular total en sección 4.2: {e}")
            valor_total = 0
        
        # Formatear valores
        valor_letras = ""
        if valor_total > 0:
            valor_letras = numero_a_letras(valor_total, incluir_moneda=True)
        
        content_data = {
            "anio": anio,
            "mes": mes,
            "user_id": data.get("user_id"),
            "name_file": data.get("name_file"),
            "section_id": data.get("section_id"),
            "level": data.get("level"),
            "content": {
                "comunicado": comunicado.get("numero", "") if isinstance(comunicado, dict) else str(comunicado) if comunicado else "",
                "fecha": comunicado.get("fecha", "") if isinstance(comunicado, dict) else "",
                "elementos": elementos,
                "total": valor_total,
                "valor_letras": valor_letras,
                "anexos": content.get("anexos", []),
            },
        }
        
        return content_data
    
    async def _seccion_4_3(self, data: Dict[str, Any]):
        """
        Procesa la sección 4.3: Entrega Equipos No Operativos Almacén SDSCJ
        """
        content = data.get("content", {})
        anio = data.get("anio")
        mes = data.get("mes")
        
        # Obtener datos del comunicado
        comunicado = content.get("comunicado", {})
        equipos = content.get("equipos", [])
        
        # Calcular total
        valor_total = 0
        try:
            for equipo in equipos:
                valor = equipo.get("valor", 0) or equipo.get("valor_total", 0)
                if isinstance(valor, (int, float)):
                    valor_total += float(valor)
                elif isinstance(valor, str):
                    valor_limpio = valor.replace(",", "").replace(" ", "").replace("$", "").strip()
                    if valor_limpio:
                        valor_total += float(valor_limpio)
        except (ValueError, TypeError) as e:
            logger.warning(f"Error al calcular total en sección 4.3: {e}")
            valor_total = 0
        
        # Formatear valores
        valor_letras = ""
        if valor_total > 0:
            valor_letras = numero_a_letras(valor_total, incluir_moneda=True)
        
        content_data = {
            "anio": anio,
            "mes": mes,
            "user_id": data.get("user_id"),
            "name_file": data.get("name_file"),
            "section_id": data.get("section_id"),
            "level": data.get("level"),
            "content": {
                "comunicado": comunicado.get("numero", "") if isinstance(comunicado, dict) else str(comunicado) if comunicado else "",
                "fecha": comunicado.get("fecha", "") if isinstance(comunicado, dict) else "",
                "equipos": equipos,
                "total": valor_total,
                "valor_letras": valor_letras,
                "anexos": content.get("anexos", []),
            },
        }
        
        return content_data
    
    async def _seccion_4_4(self, data: Dict[str, Any]):
        """
        Procesa la sección 4.4: Gestiones de Inclusión a la Bolsa
        """
        content = data.get("content", {})
        anio = data.get("anio")
        mes = data.get("mes")
        
        # Obtener datos del comunicado/solicitud
        comunicado = content.get("comunicado", {}) or content.get("solicitud", {})
        items = content.get("items", []) or content.get("elementos", [])
        
        # Calcular total
        valor_total = 0
        try:
            for item in items:
                valor = item.get("valor_total", 0) or item.get("valor", 0)
                if isinstance(valor, (int, float)):
                    valor_total += float(valor)
                elif isinstance(valor, str):
                    valor_limpio = valor.replace(",", "").replace(" ", "").replace("$", "").strip()
                    if valor_limpio:
                        valor_total += float(valor_limpio)
        except (ValueError, TypeError) as e:
            logger.warning(f"Error al calcular total en sección 4.4: {e}")
            valor_total = 0
        
        # Formatear valores
        valor_letras = ""
        if valor_total > 0:
            valor_letras = numero_a_letras(valor_total, incluir_moneda=True)
        
        content_data = {
            "anio": anio,
            "mes": mes,
            "user_id": data.get("user_id"),
            "name_file": data.get("name_file"),
            "section_id": data.get("section_id"),
            "level": data.get("level"),
            "content": {
                "comunicado": comunicado.get("numero", "") if isinstance(comunicado, dict) else str(comunicado) if comunicado else "",
                "fecha": comunicado.get("fecha", "") if isinstance(comunicado, dict) else "",
                "items": items,
                "total": valor_total,
                "valor_letras": valor_letras,
                "anexos": content.get("anexos", []),
            },
        }
        
        return content_data
    
    async def preload_seccion_4(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Precarga todas las secciones del módulo 4 llamando a los métodos correspondientes.
        Retorna un diccionario con los datos procesados de cada sección.
        
        Args:
            data: Diccionario con los datos del documento (anio, mes, user_id, name_file, etc.)
            
        Returns:
            Diccionario con los datos procesados de cada sección:
            {
                "4": {...datos procesados...},
                "4.1": {...datos procesados...},
                "4.2": {...datos procesados...},
                "4.3": {...datos procesados...},
                "4.4": {...datos procesados...},
            }
        """
        resultados = {}
        
        # Preparar datos base para cada sección
        base_data = {
            "anio": data.get("anio"),
            "mes": data.get("mes"),
            "user_id": data.get("user_id", 1),
            "name_file": data.get("name_file", f"INFORME_MENSUAL_{data.get('mes')}_{data.get('anio')}_V1.docx")
        }
        
        # Obtener el documento completo para pasar los datos de cada sección
        document = data.get("document", {})
        index = document.get("index", [])
        
        # Precargar todas las secciones y capturar los datos procesados
        try:
            # Sección 4 (principal)
            section_4_data = base_data.copy()
            section_4_item = next((item for item in index if item.get("id") == "4"), {})
            section_4_data.update({
                "section_id": "4",
                "level": section_4_item.get("level", 1),
                "title": section_4_item.get("title", "4. INFORME DE BIENES Y SERVICIOS"),
                "content": section_4_item.get("content", {})
            })
            resultados["4"] = self._seccion_4(section_4_data)
        except Exception as e:
            logger.error(f"Error al precargar sección 4: {e}", exc_info=True)
        
        try:
            # Sección 4.1
            section_4_1_data = base_data.copy()
            section_4_1_item = next((item for item in index if item.get("id") == "4.1"), {})
            section_4_1_data.update({
                "section_id": "4.1",
                "level": section_4_1_item.get("level", 2),
                "title": section_4_1_item.get("title", "4.1 GESTIÓN DE INVENTARIO"),
                "content": section_4_1_item.get("content", {})
            })
            resultados["4.1"] = await self._seccion_4_1(section_4_1_data)
        except Exception as e:
            logger.error(f"Error al precargar sección 4.1: {e}", exc_info=True)
        
        try:
            # Sección 4.2
            section_4_2_data = base_data.copy()
            section_4_2_item = next((item for item in index if item.get("id") == "4.2"), {})
            section_4_2_data.update({
                "section_id": "4.2",
                "level": section_4_2_item.get("level", 2),
                "title": section_4_2_item.get("title", "4.2 ENTRADAS ALMACÉN SDSCJ"),
                "content": section_4_2_item.get("content", {})
            })
            resultados["4.2"] = await self._seccion_4_2(section_4_2_data)
        except Exception as e:
            logger.error(f"Error al precargar sección 4.2: {e}", exc_info=True)
        
        try:
            # Sección 4.3
            section_4_3_data = base_data.copy()
            section_4_3_item = next((item for item in index if item.get("id") == "4.3"), {})
            section_4_3_data.update({
                "section_id": "4.3",
                "level": section_4_3_item.get("level", 2),
                "title": section_4_3_item.get("title", "4.3 ENTREGA EQUIPOS NO OPERATIVOS ALMACÉN SDSCJ"),
                "content": section_4_3_item.get("content", {})
            })
            resultados["4.3"] = await self._seccion_4_3(section_4_3_data)
        except Exception as e:
            logger.error(f"Error al precargar sección 4.3: {e}", exc_info=True)
        
        try:
            # Sección 4.4
            section_4_4_data = base_data.copy()
            section_4_4_item = next((item for item in index if item.get("id") == "4.4"), {})
            section_4_4_data.update({
                "section_id": "4.4",
                "level": section_4_4_item.get("level", 2),
                "title": section_4_4_item.get("title", "4.4 GESTIONES DE INCLUSIÓN A LA BOLSA"),
                "content": section_4_4_item.get("content", {})
            })
            resultados["4.4"] = await self._seccion_4_4(section_4_4_data)
        except Exception as e:
            logger.error(f"Error al precargar sección 4.4: {e}", exc_info=True)
        
        return resultados
    
    def _construir_contexto(self, document: Dict[str, Any], template: DocxTemplate) -> Dict[str, Any]:
        """
        Construye el contexto para renderizar el template Jinja2.
        Similar a _construir_contexto de seccion_2_mesa_servicio.py
        """
        tiempo_inicio = time.time()
        mes = document.get("mes")
        anio = document.get("anio")
        
        # Asegurar que mes y anio sean válidos
        mes_nombre = config.MESES.get(mes, "") if mes else ""
        
        index = document.get("index", [])
        if not isinstance(index, list):
            index = []
        
        # Inicializar contexto base
        contexto = {
            "mes": mes_nombre,
            "anio": anio,
            "mes_numero": mes,
        }
        
        tiempo_procesamiento_imagenes = 0
        tiempo_procesamiento_tablas = 0
        
        # Sección 4 (index[0]) - INFORME DE BIENES Y SERVICIOS
        if len(index) > 0:
            tiempo_img = time.time()
            content_4 = index[0].get("content", {})
            imagen_4 = procesar_imagen(template, content_4.get("image", ""))
            tiempo_procesamiento_imagenes += time.time() - tiempo_img
            if imagen_4:
                contexto["image_4"] = imagen_4
        
        # Sección 4.1 (index[1]) - GESTIÓN DE INVENTARIO
        if len(index) > 1:
            content_41 = index[1].get("content", {})
            # Obtener texto (ya viene combinado con ruta desde el servicio)
            texto_41 = content_41.get("texto", "")
            contexto["texto_41"] = texto_41
            logger.debug(f"Sección 4.1 - texto_41: {texto_41[:100] if texto_41 else 'vacío'}...")
            
            # Preparar datos de tabla en formato lista de listas
            tiempo_tabla = time.time()
            table_41_data = content_41.get("tabla", [])
            logger.debug(f"Sección 4.1 - Datos de tabla: {len(table_41_data) if table_41_data else 0} registros")
            if table_41_data:
                # Detectar columnas automáticamente desde los datos
                headers_41, campos_41 = detectar_columnas_desde_datos(table_41_data)
                logger.debug(f"Sección 4.1 - Headers detectados: {headers_41}")
                logger.debug(f"Sección 4.1 - Campos detectados: {campos_41}")
                contexto["table_41"] = preparar_datos_tabla(
                    table_41_data, headers_41, campos_41,
                    agregar_totales=False, columna_total_texto=0
                )
                logger.debug(f"Sección 4.1 - Tabla preparada: {len(contexto['table_41'])} filas")
            else:
                contexto["table_41"] = []
                logger.debug("Sección 4.1 - No hay datos de tabla")
            tiempo_procesamiento_tablas += time.time() - tiempo_tabla
        
        # Sección 4.2 (index[2]) - ENTRADAS ALMACÉN SDSCJ
        if len(index) > 2:
            content_42 = index[2].get("content", {})
            logger.info(f"Sección 4.2 - Content obtenido: {list(content_42.keys())}")
            
            # Obtener texto de la sección 4.2 - este va en comunicado_42 según el usuario
            texto_42 = content_42.get("texto", "")
            contexto["comunicado_42"] = texto_42  # El usuario quiere que comunicado_42 tenga el texto
            logger.info(f"Sección 4.2 - comunicado_42 (texto): '{texto_42[:100] if texto_42 else 'vacío'}...' (longitud: {len(texto_42)})")
            
            # Obtener datos del comunicado (número y fecha) si existen
            comunicado = content_42.get("comunicado", {})
            contexto["fecha_42"] = comunicado.get("fecha", "") if isinstance(comunicado, dict) else ""
            logger.info(f"Sección 4.2 - fecha_42: '{contexto['fecha_42']}'")
            
            # Preparar datos de tabla en formato lista de listas
            tiempo_tabla = time.time()
            elementos_42 = content_42.get("elementos", [])
            logger.info(f"Sección 4.2 - elementos_42: {len(elementos_42)} elementos")
            if elementos_42:
                logger.info(f"Sección 4.2 - Procesando {len(elementos_42)} elementos")
                
                # Verificar si los elementos tienen valores monetarios
                tiene_valores = any(
                    elem.get("valor_unitario") or elem.get("valorUnitario") or 
                    elem.get("valor_total") or elem.get("valorTotal")
                    for elem in elementos_42
                )
                
                # Preparar datos (SIN "No." y SIN "UND" según solicitud del usuario)
                datos_preparados = []
                for elemento in elementos_42:
                    dato = elemento.copy()
                    # Asegurar que existan todos los campos necesarios (sin "No." y sin "unidad")
                    if "descripcion" not in dato:
                        dato["descripcion"] = dato.get("descripcion", "")
                    if "cantidad" not in dato:
                        dato["cantidad"] = dato.get("cantidad", 0)
                    
                    # Manejar valores monetarios (pueden no existir)
                    if tiene_valores:
                        # Mantener valores numéricos para calcular totales
                        if "valor_unitario" not in dato:
                            valor_unit = dato.get("valor_unitario", dato.get("valorUnitario", 0))
                            try:
                                dato["valor_unitario"] = float(valor_unit) if valor_unit else 0
                            except (ValueError, TypeError):
                                dato["valor_unitario"] = 0
                        else:
                            try:
                                dato["valor_unitario"] = float(dato["valor_unitario"]) if dato["valor_unitario"] else 0
                            except (ValueError, TypeError):
                                dato["valor_unitario"] = 0
                        
                        if "valor_total" not in dato:
                            # Calcular si no existe
                            cantidad = float(dato.get("cantidad", 0) or 0)
                            valor_unit = float(dato.get("valor_unitario", 0) or 0)
                            dato["valor_total"] = cantidad * valor_unit
                        else:
                            try:
                                dato["valor_total"] = float(dato["valor_total"]) if dato["valor_total"] else 0
                            except (ValueError, TypeError):
                                dato["valor_total"] = 0
                    else:
                        # Si no hay valores monetarios, usar 0
                        dato["valor_unitario"] = 0
                        dato["valor_total"] = 0
                    
                    datos_preparados.append(dato)
                
                # Determinar columnas según si hay valores monetarios
                # Eliminar columnas "No." y "UND" según solicitud del usuario
                if tiene_valores:
                    headers_42_final = ["DESCRIPCIÓN", "CANT.", "VALOR UNIT.", "VALOR TOTAL"]
                    campos_42_final = ["descripcion", "cantidad", "valor_unitario", "valor_total"]
                else:
                    headers_42_final = ["DESCRIPCIÓN", "CANT."]
                    campos_42_final = ["descripcion", "cantidad"]
                try:
                    table_42_raw = preparar_datos_tabla(
                        datos_preparados, headers_42_final, campos_42_final,
                        agregar_totales=True, columna_total_texto=0  # Cambiar a 0 porque "DESCRIPCIÓN" es ahora la primera columna
                    )
                except ZeroDivisionError as e:
                    logger.error(f"Error de división por cero al preparar tabla 4.2: {e}")
                    logger.error(f"Datos preparados: {datos_preparados[:2] if datos_preparados else 'vacío'}")
                    # Crear tabla sin totales como fallback
                    table_42_raw = preparar_datos_tabla(
                        datos_preparados, headers_42_final, campos_42_final,
                        agregar_totales=False, columna_total_texto=0  # Cambiar a 0 porque "DESCRIPCIÓN" es ahora la primera columna
                    )
                except Exception as e:
                    logger.error(f"Error al preparar tabla 4.2: {e}", exc_info=True)
                    raise
                
                # Formatear valores monetarios DESPUÉS de calcular totales
                contexto["table_42"] = []
                for row_idx, fila in enumerate(table_42_raw):
                    fila_formateada = []
                    for col_idx, valor in enumerate(fila):
                        # Formatear columnas monetarias solo si existen (índices 2 y 3: valor_unitario y valor_total, sin "No." y "UND")
                        if tiene_valores and col_idx in [2, 3] and row_idx > 0:  # No formatear encabezados
                            try:
                                # Si es un string formateado, no hacer nada
                                if isinstance(valor, str) and valor.startswith("$"):
                                    fila_formateada.append(valor)
                                # Si es un número, formatearlo
                                elif valor and str(valor).strip():
                                    num_valor = float(valor)
                                    if num_valor > 0:
                                        fila_formateada.append(formato_moneda_cop(num_valor))
                                    else:
                                        fila_formateada.append("-")
                                else:
                                    fila_formateada.append("-")
                            except (ValueError, TypeError):
                                fila_formateada.append(str(valor) if valor else "-")
                        else:
                            fila_formateada.append(str(valor) if valor is not None else "")
                    contexto["table_42"].append(fila_formateada)
                logger.info(f"Sección 4.2 - Tabla preparada: {len(contexto['table_42'])} filas")
            else:
                contexto["table_42"] = []
                logger.info("Sección 4.2 - No hay elementos para la tabla")
            tiempo_procesamiento_tablas += time.time() - tiempo_tabla
            
            # Calcular total y valor en letras
            valor_total_42 = 0
            try:
                for elemento in elementos_42:
                    valor = elemento.get("valor_total", 0) or elemento.get("valor", 0)
                    if isinstance(valor, (int, float)):
                        valor_total_42 += float(valor)
                    elif isinstance(valor, str):
                        valor_limpio = valor.replace(",", "").replace(" ", "").replace("$", "").strip()
                        if valor_limpio:
                            valor_total_42 += float(valor_limpio)
            except:
                valor_total_42 = 0
            
            if valor_total_42 > 0:
                contexto["valor_letras_42"] = numero_a_letras(valor_total_42, incluir_moneda=True)
            else:
                contexto["valor_letras_42"] = ""
        
        # Anexos
            anexos_42 = content_42.get("anexos", [])
            if isinstance(anexos_42, list):
                contexto["anexos_42"] = "\n".join([f"• {anexo}" for anexo in anexos_42 if anexo])
            elif anexos_42:
                contexto["anexos_42"] = str(anexos_42)
            else:
                contexto["anexos_42"] = ""
        
        # Sección 4.3 (index[3]) - ENTREGA EQUIPOS NO OPERATIVOS
        if len(index) > 3:
            content_43 = index[3].get("content", {})
            comunicado = content_43.get("comunicado", {})
            contexto["comunicado_43"] = comunicado.get("numero", "") if isinstance(comunicado, dict) else str(comunicado) if comunicado else ""
            contexto["fecha_43"] = comunicado.get("fecha", "") if isinstance(comunicado, dict) else ""
            
            # Preparar datos de tabla en formato lista de listas
            tiempo_tabla = time.time()
            equipos_43 = content_43.get("equipos", [])
            if equipos_43:
                logger.info(f"Sección 4.3 - Procesando {len(equipos_43)} equipos")
                
                # Verificar si los equipos tienen serial y valor
                tiene_serial = any(eq.get("serial") for eq in equipos_43)
                tiene_valor = any(eq.get("valor") for eq in equipos_43)
                
                # Determinar columnas según los datos disponibles
                if tiene_serial and tiene_valor:
                    headers_43 = ["No.", "DESCRIPCIÓN", "SERIAL", "CANT.", "MOTIVO", "VALOR"]
                    campos_43 = ["descripcion", "serial", "cantidad", "motivo", "valor"]
                elif tiene_serial:
                    headers_43 = ["No.", "DESCRIPCIÓN", "SERIAL", "CANT.", "MOTIVO"]
                    campos_43 = ["descripcion", "serial", "cantidad", "motivo"]
                elif tiene_valor:
                    headers_43 = ["No.", "DESCRIPCIÓN", "CANT.", "MOTIVO", "VALOR"]
                    campos_43 = ["descripcion", "cantidad", "motivo", "valor"]
                else:
                    headers_43 = ["No.", "DESCRIPCIÓN", "CANT.", "MOTIVO"]
                    campos_43 = ["descripcion", "cantidad", "motivo"]
                
                # Preparar datos con numeración (MANTENER VALORES NUMÉRICOS para calcular totales)
                datos_preparados = []
                for idx, equipo in enumerate(equipos_43, 1):
                    dato = equipo.copy()
                    dato["No."] = idx
                    # Asegurar que existan todos los campos
                    if "descripcion" not in dato:
                        dato["descripcion"] = dato.get("descripcion", "")
                    if tiene_serial and "serial" not in dato:
                        dato["serial"] = dato.get("serial", "-")
                    if "cantidad" not in dato:
                        dato["cantidad"] = dato.get("cantidad", 1)
                    if "motivo" not in dato:
                        dato["motivo"] = dato.get("motivo", "")
                    # Mantener valor numérico para calcular totales solo si existe
                    if tiene_valor:
                        if "valor" not in dato:
                            dato["valor"] = 0
                        else:
                            try:
                                dato["valor"] = float(dato["valor"]) if dato["valor"] else 0
                            except (ValueError, TypeError):
                                dato["valor"] = 0
                    datos_preparados.append(dato)
                
                # Usar las columnas dinámicas determinadas anteriormente
                headers_43_final = headers_43
                campos_43_final = campos_43
                try:
                    table_43_raw = preparar_datos_tabla(
                        datos_preparados, headers_43_final, campos_43_final,
                        agregar_totales=True, columna_total_texto=1
                    )
                except ZeroDivisionError as e:
                    logger.error(f"Error de división por cero al preparar tabla 4.3: {e}")
                    logger.error(f"Datos preparados: {datos_preparados[:2] if datos_preparados else 'vacío'}")
                    # Crear tabla sin totales como fallback
                    table_43_raw = preparar_datos_tabla(
                        datos_preparados, headers_43_final, campos_43_final,
                        agregar_totales=False, columna_total_texto=1
                    )
                except Exception as e:
                    logger.error(f"Error al preparar tabla 4.3: {e}", exc_info=True)
                    raise
                
                # Formatear valores monetarios DESPUÉS de calcular totales
                contexto["table_43"] = []
                # Determinar índice de columna de valor (si existe)
                indice_valor = None
                if tiene_valor:
                    try:
                        indice_valor = campos_43_final.index("valor")
                    except ValueError:
                        indice_valor = None
                
                for row_idx, fila in enumerate(table_43_raw):
                    fila_formateada = []
                    for col_idx, valor in enumerate(fila):
                        # Formatear columna monetaria solo si existe
                        if tiene_valor and indice_valor is not None and col_idx == indice_valor and row_idx > 0:  # No formatear encabezados
                            try:
                                if isinstance(valor, str) and valor.startswith("$"):
                                    fila_formateada.append(valor)
                                elif valor and str(valor).strip():
                                    num_valor = float(valor)
                                    if num_valor > 0:
                                        fila_formateada.append(formato_moneda_cop(num_valor))
                                    else:
                                        fila_formateada.append("-")
                                else:
                                    fila_formateada.append("-")
                            except (ValueError, TypeError):
                                fila_formateada.append(str(valor) if valor else "-")
                        else:
                            fila_formateada.append(str(valor) if valor is not None else "")
                    contexto["table_43"].append(fila_formateada)
            else:
                contexto["table_43"] = []
            tiempo_procesamiento_tablas += time.time() - tiempo_tabla
            
            # Calcular total y valor en letras
            valor_total_43 = 0
            try:
                for equipo in equipos_43:
                    valor = equipo.get("valor", 0) or equipo.get("valor_total", 0)
                    if isinstance(valor, (int, float)):
                        valor_total_43 += float(valor)
                    elif isinstance(valor, str):
                        valor_limpio = valor.replace(",", "").replace(" ", "").replace("$", "").strip()
                        if valor_limpio:
                            valor_total_43 += float(valor_limpio)
            except:
                valor_total_43 = 0
            
            if valor_total_43 > 0:
                contexto["valor_letras_43"] = numero_a_letras(valor_total_43, incluir_moneda=True)
            else:
                contexto["valor_letras_43"] = ""
            
            # Anexos
            anexos_43 = content_43.get("anexos", [])
            if isinstance(anexos_43, list):
                contexto["anexos_43"] = "\n".join([f"• {anexo}" for anexo in anexos_43 if anexo])
            elif anexos_43:
                contexto["anexos_43"] = str(anexos_43)
            else:
                contexto["anexos_43"] = ""
        
        # Sección 4.4 (index[4]) - GESTIONES DE INCLUSIÓN A LA BOLSA
        if len(index) > 4:
            content_44 = index[4].get("content", {})
            comunicado = content_44.get("comunicado", {}) or content_44.get("solicitud", {})
            contexto["comunicado_44"] = comunicado.get("numero", "") if isinstance(comunicado, dict) else str(comunicado) if comunicado else ""
            contexto["fecha_44"] = comunicado.get("fecha", "") if isinstance(comunicado, dict) else ""
            
            # Preparar datos de tabla en formato lista de listas
            tiempo_tabla = time.time()
            items_44 = content_44.get("items", []) or content_44.get("elementos", [])
            if items_44:
                # Columnas esperadas para inclusión a bolsa
                headers_44 = ["No.", "DESCRIPCIÓN", "CANT.", "UND", "VALOR UNIT.", "VALOR TOTAL", "JUSTIFICACIÓN"]
                campos_44 = ["descripcion", "cantidad", "unidad", "valor_unitario", "valor_total", "justificacion"]
                
                # Preparar datos con numeración (MANTENER VALORES NUMÉRICOS para calcular totales)
                datos_preparados = []
                for idx, item in enumerate(items_44, 1):
                    dato = item.copy()
                    dato["No."] = idx
                    # Asegurar que existan todos los campos
                    if "descripcion" not in dato:
                        dato["descripcion"] = dato.get("descripcion", "")
                    if "cantidad" not in dato:
                        dato["cantidad"] = dato.get("cantidad", 0)
                    if "unidad" not in dato:
                        dato["unidad"] = dato.get("unidad", "UN")
                    # Mantener valores numéricos para calcular totales
                    if "valor_unitario" not in dato:
                        valor_unit = dato.get("valor_unitario", 0)
                        try:
                            dato["valor_unitario"] = float(valor_unit) if valor_unit else 0
                        except (ValueError, TypeError):
                            dato["valor_unitario"] = 0
                    else:
                        try:
                            dato["valor_unitario"] = float(dato["valor_unitario"]) if dato["valor_unitario"] else 0
                        except (ValueError, TypeError):
                            dato["valor_unitario"] = 0
                    
                    if "valor_total" not in dato:
                        # Calcular si no existe
                        cantidad = float(dato.get("cantidad", 0) or 0)
                        valor_unit = float(dato.get("valor_unitario", 0) or 0)
                        dato["valor_total"] = cantidad * valor_unit
                    else:
                        try:
                            dato["valor_total"] = float(dato["valor_total"]) if dato["valor_total"] else 0
                        except (ValueError, TypeError):
                            dato["valor_total"] = 0
                    if "justificacion" not in dato:
                        dato["justificacion"] = dato.get("justificacion", "")
                    datos_preparados.append(dato)
                
                # Preparar tabla con totales (usando valores numéricos)
                headers_44_final = ["No.", "DESCRIPCIÓN", "CANT.", "UND", "VALOR UNIT.", "VALOR TOTAL", "JUSTIFICACIÓN"]
                campos_44_final = ["No.", "descripcion", "cantidad", "unidad", "valor_unitario", "valor_total", "justificacion"]
                try:
                    table_44_raw = preparar_datos_tabla(
                        datos_preparados, headers_44_final, campos_44_final,
                        agregar_totales=True, columna_total_texto=1
                    )
                except ZeroDivisionError as e:
                    logger.error(f"Error de división por cero al preparar tabla 4.4: {e}")
                    logger.error(f"Datos preparados: {datos_preparados[:2] if datos_preparados else 'vacío'}")
                    # Crear tabla sin totales como fallback
                    table_44_raw = preparar_datos_tabla(
                        datos_preparados, headers_44_final, campos_44_final,
                        agregar_totales=False, columna_total_texto=1
                    )
                except Exception as e:
                    logger.error(f"Error al preparar tabla 4.4: {e}", exc_info=True)
                    raise
                
                # Formatear valores monetarios DESPUÉS de calcular totales
                contexto["table_44"] = []
                for row_idx, fila in enumerate(table_44_raw):
                    fila_formateada = []
                    for col_idx, valor in enumerate(fila):
                        # Formatear columnas monetarias (índices 4 y 5: valor_unitario y valor_total)
                        if col_idx in [4, 5] and row_idx > 0:  # No formatear encabezados
                            try:
                                if isinstance(valor, str) and valor.startswith("$"):
                                    fila_formateada.append(valor)
                                elif valor and str(valor).strip():
                                    num_valor = float(valor)
                                    if num_valor > 0:
                                        fila_formateada.append(formato_moneda_cop(num_valor))
                                    else:
                                        fila_formateada.append("-")
                                else:
                                    fila_formateada.append("-")
                            except (ValueError, TypeError):
                                fila_formateada.append(str(valor) if valor else "-")
                        else:
                            fila_formateada.append(str(valor) if valor is not None else "")
                    contexto["table_44"].append(fila_formateada)
            else:
                contexto["table_44"] = []
            tiempo_procesamiento_tablas += time.time() - tiempo_tabla
            
            # Calcular total y valor en letras
            valor_total_44 = 0
            try:
                for item in items_44:
                    valor = item.get("valor_total", 0) or item.get("valor", 0)
                    if isinstance(valor, (int, float)):
                        valor_total_44 += float(valor)
                    elif isinstance(valor, str):
                        valor_limpio = valor.replace(",", "").replace(" ", "").replace("$", "").strip()
                        if valor_limpio:
                            valor_total_44 += float(valor_limpio)
            except:
                valor_total_44 = 0
            
            if valor_total_44 > 0:
                contexto["valor_letras_44"] = numero_a_letras(valor_total_44, incluir_moneda=True)
            else:
                contexto["valor_letras_44"] = ""
            
            # Anexos
            anexos_44 = content_44.get("anexos", [])
            if isinstance(anexos_44, list):
                contexto["anexos_44"] = "\n".join([f"• {anexo}" for anexo in anexos_44 if anexo])
            elif anexos_44:
                contexto["anexos_44"] = str(anexos_44)
            else:
                contexto["anexos_44"] = ""
        
        tiempo_total_contexto = time.time() - tiempo_inicio
        logger.info(f"  📊 Detalle construcción contexto:")
        if tiempo_total_contexto > 0:
            logger.info(f"     - Procesamiento imágenes: {tiempo_procesamiento_imagenes:.2f}s ({tiempo_procesamiento_imagenes/tiempo_total_contexto*100:.1f}%)")
            logger.info(f"     - Procesamiento tablas: {tiempo_procesamiento_tablas:.2f}s ({tiempo_procesamiento_tablas/tiempo_total_contexto*100:.1f}%)")
        else:
            logger.info(f"     - Procesamiento imágenes: {tiempo_procesamiento_imagenes:.2f}s")
            logger.info(f"     - Procesamiento tablas: {tiempo_procesamiento_tablas:.2f}s")
        logger.info(f"     - Otros: {tiempo_total_contexto - tiempo_procesamiento_imagenes - tiempo_procesamiento_tablas:.2f}s")
        
        return contexto
    
    async def generar(self, document: Dict[str, Any], output_path: Optional[Path] = None):
        """
        Genera el documento completo de la Sección 4.
        Similar a generar() de seccion_2_mesa_servicio.py
        """
        tiempo_inicio_total = time.time()
        logger.info("=" * 80)
        logger.info("INICIO GENERACIÓN DE DOCUMENTO SECCIÓN 4")
        logger.info("=" * 80)
        
        try:
            tiempo_inicio = time.time()
            template_path = config.TEMPLATES_DIR / self.template_file
            template = DocxTemplate(str(template_path))
            tiempo_carga_template = time.time() - tiempo_inicio
            logger.info(f"⏱️  Tiempo carga template: {tiempo_carga_template:.2f}s")
            
            tiempo_inicio = time.time()
            try:
                contexto = self._construir_contexto(document, template)
            except ZeroDivisionError as e:
                logger.error(f"Error de división por cero al construir contexto: {e}", exc_info=True)
                raise ValueError(f"Error de división por cero al procesar datos: {e}")
            tiempo_construccion_contexto = time.time() - tiempo_inicio
            logger.info(f"⏱️  Tiempo construcción contexto: {tiempo_construccion_contexto:.2f}s")
        except ZeroDivisionError as e:
            logger.error(f"Error de división por cero en generar(): {e}", exc_info=True)
            raise ValueError(f"Error de división por cero: {e}")
        
        if output_path:
            # Asegurar que el output_path tenga extensión .docx
            if not str(output_path).endswith('.docx'):
                output_path = Path(str(output_path) + '.docx')
        else:
            # Crear ruta por defecto
            mes = document.get("mes")
            anio = document.get("anio")
            output_path = Path(f"/tmp/seccion_4_{anio}_{mes}.docx")
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        tiempo_inicio = time.time()
        # Preparar placeholders especiales para tablas
        contexto_tablas = {}
        # Solo copiar las claves necesarias en lugar de copiar todo el contexto
        for key in ["mes", "anio", "mes_numero", "texto_41", "texto_42", "comunicado_42", "fecha_42",
                   "valor_letras_42", "anexos_42", "comunicado_43", "fecha_43",
                   "valor_letras_43", "anexos_43", "comunicado_44", "fecha_44",
                   "valor_letras_44", "anexos_44"]:
            if key in contexto:
                contexto_tablas[key] = contexto[key]
        
        # Agregar imágenes solo si existen
        for img_key in ["image_4"]:
            if img_key in contexto:
                contexto_tablas[img_key] = contexto[img_key]
        
        # Agregar placeholders para tablas (solo para Jinja2, las tablas se insertan después)
        placeholders_tablas = {}
        tabla_mapping = {
            "table_41": "[[TABLE_41]]",
            "table_42": "[[TABLE_42]]",
            "table_43": "[[TABLE_43]]",
            "table_44": "[[TABLE_44]]",
        }
        
        for tabla_key, placeholder in tabla_mapping.items():
            tabla_data = contexto.get(tabla_key)
            if tabla_data:
                logger.info(f"   ✓ Tabla {tabla_key}: {len(tabla_data)} filas (incluyendo encabezado)")
                contexto_tablas[f"{tabla_key}_placeholder"] = placeholder
                placeholders_tablas[placeholder] = tabla_data
            else:
                logger.debug(f"   ✗ Tabla {tabla_key}: sin datos")
        
        tiempo_preparacion_contexto = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo preparación contexto tablas: {tiempo_preparacion_contexto:.2f}s")
        logger.info(f"   - Número de tablas a procesar: {len(placeholders_tablas)}")
        if placeholders_tablas:
            logger.info(f"   - Placeholders de tablas: {list(placeholders_tablas.keys())}")
        
        tiempo_inicio = time.time()
        # Renderizar el template con las variables básicas
        template.render(contexto_tablas)
        tiempo_renderizado = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo renderizado template Jinja2: {tiempo_renderizado:.2f}s")
        
        # Log de qué placeholders se renderizaron
        for key in ["table_42_placeholder", "table_43_placeholder", "table_44_placeholder"]:
            if key in contexto_tablas:
                logger.info(f"  ✓ {key} renderizado como: '{contexto_tablas[key]}'")
        
        tiempo_inicio = time.time()
        # Guardar temporalmente para poder trabajar con el documento renderizado
        temp_path = str(output_path).replace('.docx', '_temp.docx')
        template.save(temp_path)
        tiempo_guardado_temp = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo guardado template temporal: {tiempo_guardado_temp:.2f}s")
        
        tiempo_inicio = time.time()
        # Abrir el documento renderizado para insertar tablas programáticamente
        doc = Document(temp_path)
        tiempo_apertura_doc = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo apertura documento: {tiempo_apertura_doc:.2f}s")
        
        # Verificar qué placeholders quedaron en el documento después del renderizado
        logger.info(f"  🔍 Verificando placeholders en documento renderizado...")
        logger.info(f"  📄 Total de párrafos en documento: {len(doc.paragraphs)}")
        logger.info(f"  📊 Total de tablas en documento: {len(doc.tables)}")
        
        # Mostrar los primeros 20 párrafos para depuración
        logger.info(f"  📝 Primeros 20 párrafos del documento:")
        for i, paragraph in enumerate(doc.paragraphs[:20]):
            texto = paragraph.text.strip()
            if texto:  # Solo mostrar párrafos con texto
                logger.info(f"     Párrafo {i}: '{texto[:100]}...'")
        
        textos_encontrados = []
        # Buscar en TODOS los párrafos
        for i, paragraph in enumerate(doc.paragraphs):
            texto = paragraph.text
            # Buscar cualquier referencia a TABLE_42, TABLE_43, TABLE_44
            if any(ph in texto for ph in ["TABLE_42", "TABLE_43", "TABLE_44", "tabla_42", "tabla_43", "tabla_44", "[[TABLE", "]]"]):
                textos_encontrados.append(f"Párrafo {i}: '{texto[:150]}...'")
        
        # Buscar también en tablas existentes
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    texto_celda = cell.text
                    if any(ph in texto_celda for ph in ["TABLE_42", "TABLE_43", "TABLE_44", "tabla_42", "tabla_43", "tabla_44", "[[TABLE", "]]"]):
                        textos_encontrados.append(f"Tabla {table_idx}, Celda ({row_idx}, {cell_idx}): '{texto_celda[:150]}...'")
        
        if textos_encontrados:
            logger.info(f"  📝 Textos relacionados con tablas encontrados ({len(textos_encontrados)}):")
            for texto in textos_encontrados[:10]:  # Primeros 10
                logger.info(f"     {texto}")
        else:
            logger.warning(f"  ⚠️ No se encontraron textos relacionados con tablas en el documento completo")
            logger.warning(f"  ⚠️ Esto significa que el placeholder {{ tabla_42_placeholder }} no se renderizó correctamente")
            logger.warning(f"  ⚠️ Verifica que el template Word tenga exactamente: {{ tabla_42_placeholder }}")
            logger.warning(f"  ⚠️ El placeholder debe estar en un párrafo normal, NO dentro de una tabla existente")
        
        tiempo_inicio = time.time()
        # Reemplazar todos los placeholders de tablas en una sola pasada (MUCHO MÁS RÁPIDO)
        if placeholders_tablas:
            try:
                reemplazar_multiples_placeholders_con_tablas(doc, placeholders_tablas)
            except ZeroDivisionError as e:
                logger.error(f"Error de división por cero al procesar tablas: {e}", exc_info=True)
                raise ValueError(f"Error de división por cero al procesar tablas: {e}")
        tiempo_procesamiento_tablas = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo procesamiento tablas: {tiempo_procesamiento_tablas:.2f}s")
        
        tiempo_inicio = time.time()
        # Guardar el documento final
        doc.save(str(output_path))
        tiempo_guardado_final = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo guardado documento final: {tiempo_guardado_final:.2f}s")
        
        # Eliminar archivo temporal
        try:
            Path(temp_path).unlink()
        except:
            pass
        
        tiempo_total = time.time() - tiempo_inicio_total
        logger.info("=" * 80)
        logger.info("RESUMEN DE TIEMPOS:")
        if tiempo_total > 0:
            logger.info(f"  • Carga template: {tiempo_carga_template:.2f}s ({tiempo_carga_template/tiempo_total*100:.1f}%)")
            logger.info(f"  • Construcción contexto: {tiempo_construccion_contexto:.2f}s ({tiempo_construccion_contexto/tiempo_total*100:.1f}%)")
            logger.info(f"  • Preparación contexto tablas: {tiempo_preparacion_contexto:.2f}s ({tiempo_preparacion_contexto/tiempo_total*100:.1f}%)")
            logger.info(f"  • Renderizado Jinja2: {tiempo_renderizado:.2f}s ({tiempo_renderizado/tiempo_total*100:.1f}%)")
            logger.info(f"  • Guardado template temporal: {tiempo_guardado_temp:.2f}s ({tiempo_guardado_temp/tiempo_total*100:.1f}%)")
            logger.info(f"  • Apertura documento: {tiempo_apertura_doc:.2f}s ({tiempo_apertura_doc/tiempo_total*100:.1f}%)")
            logger.info(f"  • Procesamiento tablas: {tiempo_procesamiento_tablas:.2f}s ({tiempo_procesamiento_tablas/tiempo_total*100:.1f}%)")
            logger.info(f"  • Guardado documento final: {tiempo_guardado_final:.2f}s ({tiempo_guardado_final/tiempo_total*100:.1f}%)")
        else:
            logger.info(f"  • Carga template: {tiempo_carga_template:.2f}s")
            logger.info(f"  • Construcción contexto: {tiempo_construccion_contexto:.2f}s")
            logger.info(f"  • Preparación contexto tablas: {tiempo_preparacion_contexto:.2f}s")
            logger.info(f"  • Renderizado Jinja2: {tiempo_renderizado:.2f}s")
            logger.info(f"  • Guardado template temporal: {tiempo_guardado_temp:.2f}s")
            logger.info(f"  • Apertura documento: {tiempo_apertura_doc:.2f}s")
            logger.info(f"  • Procesamiento tablas: {tiempo_procesamiento_tablas:.2f}s")
            logger.info(f"  • Guardado documento final: {tiempo_guardado_final:.2f}s")
        logger.info(f"  {'='*76}")
        logger.info(f"  ⏱️  TIEMPO TOTAL: {tiempo_total:.2f}s ({tiempo_total/60:.2f} minutos)" if tiempo_total > 0 else f"  ⏱️  TIEMPO TOTAL: {tiempo_total:.2f}s")
        logger.info("=" * 80)
        
        logger.info(f"[OK] Documento generado en: {output_path} (Tiempo total: {tiempo_total:.2f}s)")
        
        return template
    
    async def generar_preview(self, document: Dict[str, Any]) -> bytes:
        """
        Genera el documento de la sección 4 y retorna los bytes para preview.
        Similar a generar() pero retorna bytes en lugar de guardar en disco.
        
        Args:
            document: Diccionario con los datos del documento
            
        Returns:
            bytes: Contenido del archivo .docx generado
        """
        from io import BytesIO
        import tempfile
        
        tiempo_inicio_total = time.time()
        logger.info("=" * 80)
        logger.info("INICIO GENERACIÓN DE PREVIEW DOCUMENTO SECCIÓN 4")
        logger.info("=" * 80)
        
        tiempo_inicio = time.time()
        template_path = config.TEMPLATES_DIR / self.template_file
        template = DocxTemplate(str(template_path))
        tiempo_carga_template = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo carga template: {tiempo_carga_template:.2f}s")
        
        tiempo_inicio = time.time()
        contexto = self._construir_contexto(document, template)
        tiempo_construccion_contexto = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo construcción contexto: {tiempo_construccion_contexto:.2f}s")
        
        tiempo_inicio = time.time()
        # Preparar placeholders especiales para tablas
        contexto_tablas = {}
        # Solo copiar las claves necesarias
        for key in ["mes", "anio", "mes_numero", "texto_41", "texto_42", "comunicado_42", "fecha_42",
                   "valor_letras_42", "anexos_42", "comunicado_43", "fecha_43",
                   "valor_letras_43", "anexos_43", "comunicado_44", "fecha_44",
                   "valor_letras_44", "anexos_44"]:
            if key in contexto:
                contexto_tablas[key] = contexto[key]
        
        # Agregar imágenes solo si existen
        for img_key in ["image_4"]:
            if img_key in contexto:
                contexto_tablas[img_key] = contexto[img_key]
        
        # Agregar placeholders para tablas
        placeholders_tablas = {}
        tabla_mapping = {
            "table_41": "[[TABLE_41]]",
            "table_42": "[[TABLE_42]]",
            "table_43": "[[TABLE_43]]",
            "table_44": "[[TABLE_44]]",
        }
        
        for tabla_key, placeholder in tabla_mapping.items():
            if contexto.get(tabla_key):
                contexto_tablas[f"{tabla_key}_placeholder"] = placeholder
                placeholders_tablas[placeholder] = contexto[tabla_key]
        
        tiempo_preparacion_contexto = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo preparación contexto tablas: {tiempo_preparacion_contexto:.2f}s")
        logger.info(f"   - Número de tablas a procesar: {len(placeholders_tablas)}")
        
        tiempo_inicio = time.time()
        # Renderizar el template con las variables básicas
        template.render(contexto_tablas)
        tiempo_renderizado = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo renderizado template Jinja2: {tiempo_renderizado:.2f}s")
        
        tiempo_inicio = time.time()
        # Guardar temporalmente en memoria usando BytesIO
        temp_buffer = BytesIO()
        template.save(temp_buffer)
        temp_buffer.seek(0)
        tiempo_guardado_temp = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo guardado template temporal: {tiempo_guardado_temp:.2f}s")
        
        tiempo_inicio = time.time()
        # Abrir el documento renderizado para insertar tablas programáticamente
        doc = Document(temp_buffer)
        tiempo_apertura_doc = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo apertura documento: {tiempo_apertura_doc:.2f}s")
        
        tiempo_inicio = time.time()
        # Reemplazar todos los placeholders de tablas en una sola pasada
        if placeholders_tablas:
            reemplazar_multiples_placeholders_con_tablas(doc, placeholders_tablas)
        tiempo_procesamiento_tablas = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo procesamiento tablas: {tiempo_procesamiento_tablas:.2f}s")
        
        tiempo_inicio = time.time()
        # Guardar el documento final en BytesIO
        final_buffer = BytesIO()
        doc.save(final_buffer)
        final_buffer.seek(0)
        file_bytes = final_buffer.read()
        tiempo_guardado_final = time.time() - tiempo_inicio
        logger.info(f"⏱️  Tiempo guardado documento final: {tiempo_guardado_final:.2f}s")
        
        tiempo_total = time.time() - tiempo_inicio_total
        logger.info("=" * 80)
        logger.info(f"⏱️  TIEMPO TOTAL PREVIEW: {tiempo_total:.2f}s ({tiempo_total/60:.2f} minutos)")
        logger.info("=" * 80)
        
        return file_bytes
