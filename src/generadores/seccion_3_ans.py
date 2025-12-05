"""
Generador Sección 3: Informes de Medición de Niveles de Servicio (ANS)
Tipo: 🟨 GENERACIÓN PROGRAMÁTICA (python-docx)

SECCIÓN CRÍTICA - Impacto contractual y financiero
- Umbral contractual: 98.9% de disponibilidad
- Si no cumple: calcular penalidades y explicar causas

Subsecciones:
- 3.1 Penalidad de ANS
- 3.2 Consolidado ANS (histórico y gráficos)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Dict, List, Any, Optional
import re
import matplotlib
matplotlib.use('Agg')  # Usar backend sin GUI
import matplotlib.pyplot as plt
import numpy as np
from docxtpl import DocxTemplate, InlineImage
import config
from src.utils.formato_moneda import formato_moneda_cop, formato_cantidad


class GeneradorSeccion3():
    """Genera la Sección 3: Informes de Medición de Niveles de Servicio (ANS)"""
    
    # Umbrales contractuales
    UMBRAL_ANS = 98.9           # Porcentaje mínimo requerido
    UMBRAL_AMARILLO = 97.9      # Zona de alerta
    
    # Colores corporativos
    COLOR_AZUL_OSCURO = RGBColor(31, 78, 121)   # Encabezados principales
    COLOR_AZUL_MEDIO = RGBColor(46, 117, 182)   # Subsecciones
    COLOR_GRIS = RGBColor(64, 64, 64)           # Subtítulos
    COLOR_VERDE = RGBColor(0, 176, 80)          # ANS cumplido
    COLOR_ROJO = RGBColor(192, 0, 0)            # ANS no cumplido
    COLOR_AMARILLO = RGBColor(255, 192, 0)      # Zona de alerta
    
    # Colores de fondo para semáforo
    COLOR_VERDE_CLARO = RGBColor(198, 239, 206)
    COLOR_AMARILLO_CLARO = RGBColor(255, 242, 204)
    COLOR_ROJO_CLARO = RGBColor(255, 199, 206)
    
    @property
    def nombre_seccion(self) -> str:
        return "3. INFORMES DE MEDICIÓN DE NIVELES DE SERVICIO (ANS)"
    
    @property
    def template_file(self) -> str:
        return "Seccion 3.docx"  # Template Word con placeholder {{grafico_lineas}}
    
    def __init__(self, anio: int, mes: int):
        self.anio = anio
        self.mes = mes
        self.datos: Dict[str, Any] = {}
        self.periodo = config.get_periodo_texto(anio, mes)
        self.contrato = config.CONTRATO
    
    def _centrar_celda_vertical(self, cell):
        """Centra verticalmente el contenido de una celda"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
    
    def _cargar_contexto_base(self) -> Dict[str, Any]:
        """Carga el contexto base común a todas las secciones"""
        return {
            "contrato_numero": self.contrato["numero"],
            "entidad": self.contrato["entidad"],
            "entidad_corto": self.contrato["entidad_corto"],
            "periodo": self.periodo,
            "mes": config.MESES[self.mes],
            "anio": self.anio,
            "mes_numero": self.mes,
        }
    
    def _set_cell_shading(self, cell, hex_color: str):
        """
        Aplica color de fondo a una celda de tabla
        
        Args:
            cell: Celda de la tabla
            hex_color: Color en formato hexadecimal (ej: "1F4E79" para azul oscuro)
        """
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), hex_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def procesar(self, document: Dict[str, Any]) -> Dict[str, Any]:
        
        if document:
            for item in document.get("index", []):
                if item.get("id") == "3":
                    self.datos = item.get("content", {})
                    return self.datos
        else:
            raise ValueError("No se encontraron datos para la sección 3")
    
  
    def _generar_grafico_lineas_ans(self, output_dir: Path) -> Optional[str]:
        """
        Genera el gráfico de líneas con los valores ANS de la sección 3.1
        Retorna la ruta del archivo JPG generado
        """
        try:
            # Obtener datos para el gráfico
            categorias = [
                "ANS CALIDAD DE LOS REPORTES ENTREGADOS",
                "ANS DISPONIBILIDAD DEL SISTEMA",
                "ANS OPORTUNIDAD DE LOS REPORTES ENTREGADOS",
                "ANS OPORTUNIDAD EN LAS ACTIVIDADES DE MANTENIMIENTO PREVENTIVO",
                "ANS RTO",
                "ANS TIEMPO DE RESTAURACIÓN DE SERVICIOS EN EL DATA CENTER"
            ]
            
            # Calcular valores consolidados desde los datos
            def safe_float(value, default=0):
                """Convierte valor a float de forma segura"""
                if value is None:
                    return default
                try:
                    return float(value)
                except (ValueError, TypeError):
                    return default
            
            valor_calidad = (
                safe_float(self.datos.get('valor_calidad_informes_1')) +
                safe_float(self.datos.get('valor_calidad_informes_2')) +
                safe_float(self.datos.get('valor_calidad_informes_3')) +
                safe_float(self.datos.get('valor_calidad_informes_4'))
            )
            
            valor_disponibilidad = safe_float(self.datos.get('valor_disponibilidad_sistema'))
            valor_oportunidad_informes = safe_float(self.datos.get('valor_oportunidad_informes'))
            valor_oportunidad_actividades = safe_float(self.datos.get('valor_oportunidad_actividades'))
            
            # Consolidar RTO (sumar los 4 valores)
            valor_rto = (
                safe_float(self.datos.get('valor_rto_1')) +
                safe_float(self.datos.get('valor_rto_2')) +
                safe_float(self.datos.get('valor_rto_3')) +
                safe_float(self.datos.get('valor_rto_4'))
            )
            
            valor_tiempo_restauracion = safe_float(self.datos.get('valor_tiempo_restauracion'))
            
            valores = [
                valor_calidad,
                valor_disponibilidad,
                valor_oportunidad_informes,
                valor_oportunidad_actividades,
                valor_rto,
                valor_tiempo_restauracion
            ]
            
            # Crear figura
            plt.figure(figsize=(18, 8))
            
            # Dibujar la línea
            plt.plot(categorias, valores, marker='o', linewidth=2, color='#2E75B6', markersize=8)
            
            # Títulos
            mes_nombre = config.MESES[self.datos.get('mes')]
            anio = self.datos.get('anio', self.anio)
            plt.title(f"ANS MES {mes_nombre.upper()} {anio}", fontsize=18, fontweight='bold')
            plt.xlabel("Categorías", fontsize=12)
            plt.ylabel("Valor", fontsize=12)
            
            # Rotación del texto del eje X
            plt.xticks(rotation=45, ha='right')
            
            # Grid
            plt.grid(True, linestyle='--', alpha=0.5)
            
            # Etiquetas sobre cada punto
            for x, y in zip(categorias, valores):
                if abs(y) > 0.01:  # Solo mostrar si el valor es significativo
                    plt.text(x, y, f"${y:,.0f}", fontsize=10, ha='center', 
                            va='bottom' if y >= 0 else 'top')
                else:
                    plt.text(x, y, "$-", fontsize=10, ha='center', va='center')
            
            plt.tight_layout()
            
            # Guardar como JPG
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / f"grafico_lineas_ans_{self.mes}_{self.anio}.jpg"
            plt.savefig(output_path, dpi=150, bbox_inches='tight', format='jpg')
            plt.close()
            
            return str(output_path)
        except Exception as e:
            print(f"[WARNING] Error al generar gráfico de líneas: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _generar_grafico_pastel_ans(self, datos_actuales: Dict[str, Any], 
                                     datos_historicos: List[Dict[str, Any]], 
                                     output_dir: Path) -> Optional[str]:
        """
        Genera el gráfico de pastel (pie chart) con el consolidado de valores ANS por mes
        Retorna la ruta del archivo JPG generado
        """
        try:
            def safe_float(value, default=0.0):
                """Convierte valor a float de forma segura"""
                if value is None:
                    return default
                try:
                    return float(value)
                except (ValueError, TypeError):
                    return default
            
            # Calcular valor total del mes actual
            total_valor_actual = (
                safe_float(datos_actuales.get('valor_calidad_informes_1', 0)) +
                safe_float(datos_actuales.get('valor_calidad_informes_2', 0)) +
                safe_float(datos_actuales.get('valor_calidad_informes_3', 0)) +
                safe_float(datos_actuales.get('valor_calidad_informes_4', 0)) +
                safe_float(datos_actuales.get('valor_disponibilidad_sistema', 0)) +
                safe_float(datos_actuales.get('valor_oportunidad_informes', 0)) +
                safe_float(datos_actuales.get('valor_oportunidad_actividades', 0)) +
                safe_float(datos_actuales.get('valor_rto_1', 0)) +
                safe_float(datos_actuales.get('valor_rto_2', 0)) +
                safe_float(datos_actuales.get('valor_rto_3', 0)) +
                safe_float(datos_actuales.get('valor_rto_4', 0)) +
                safe_float(datos_actuales.get('valor_tiempo_restauracion', 0))
            )
            
            # Preparar datos para el gráfico
            meses = []
            valores = []
            
            # Agregar meses históricos (orden ascendente: más antiguo primero)
            for historico in datos_historicos:
                mes_nombre = historico.get("mes_nombre", "")
                anio = historico.get("anio", "")
                valor = historico.get("total_valor", 0)
                if valor > 0:
                    meses.append(f"{mes_nombre}-{anio}")
                    valores.append(valor)
            
            # Agregar mes actual (al final)
            mes_nombre_actual = config.MESES[self.mes].lower()[:3]
            if total_valor_actual > 0:
                meses.append(f"{mes_nombre_actual}-{self.anio}")
                valores.append(total_valor_actual)
            
            # Si no hay datos, retornar None
            if not valores or sum(valores) == 0:
                return None
            
            # Crear figura con espacio para la leyenda
            fig, ax = plt.subplots(figsize=(14, 10))
            
            # Colores para el gráfico (usar colores distintos y visibles)
            # Usar una paleta de colores más variada
            colores = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', 
                      '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf',
                      '#aec7e8', '#ffbb78']
            
            # Asegurar que tengamos suficientes colores
            if len(meses) > len(colores):
                import matplotlib.cm as cm
                colores = [cm.Set3(i) for i in range(len(meses))]
            else:
                colores = colores[:len(meses)]
            
            # Crear el gráfico de pastel
            wedges, texts, autotexts = ax.pie(
                valores,
                labels=None,  # No usar labels automáticos
                autopct='%1.0f%%',
                startangle=90,
                colors=colores,
                textprops={'fontsize': 10, 'fontweight': 'bold', 'color': 'white'},
                pctdistance=0.85
            )
            
            # Calcular porcentajes para cada segmento
            total = sum(valores)
            porcentajes = [(v / total * 100) for v in valores]
            
            # Actualizar los textos para incluir el mes y el porcentaje
            for i, (autotext, mes, pct) in enumerate(zip(autotexts, meses, porcentajes)):
                if pct > 0:  # Solo mostrar si el porcentaje es mayor a 0
                    autotext.set_text(f'{mes}\n{pct:.0f}%')
                else:
                    autotext.set_text('')
            
            # Título
            ax.set_title("CONSOLIDADO VALOR ANS", fontsize=18, fontweight='bold', pad=30)
            
            # Función auxiliar para ordenar cronológicamente
            def ordenar_mes_anio(mes_anio_str):
                """Convierte 'mes-año' a un número para ordenar cronológicamente"""
                try:
                    mes_str, anio_str = mes_anio_str.split('-')
                    # Buscar el número del mes en config.MESES
                    mes_num = None
                    for i, nombre_mes in enumerate(config.MESES.values(), 1):
                        if nombre_mes.lower()[:3] == mes_str:
                            mes_num = i
                            break
                    if mes_num is None:
                        return 0
                    anio_num = int(anio_str)
                    return anio_num * 100 + mes_num  # Año*100 + mes para ordenar
                except:
                    return 0
            
            # Crear leyenda en la parte superior con los meses y sus colores
            # Ordenar por fecha para mostrar en orden cronológico
            meses_datos = list(zip(meses, valores, colores))
            meses_ordenados = sorted(meses_datos, key=lambda x: ordenar_mes_anio(x[0]))
            
            # Crear etiquetas para la leyenda con formato: "mes-año"
            legend_labels = [f"{mes}" for mes, _, _ in meses_ordenados]
            legend_colors = [color for _, _, color in meses_ordenados]
            
            # Crear la leyenda en la parte superior
            legend_elements = [plt.Rectangle((0,0),1,1, facecolor=color, edgecolor='black', linewidth=1) 
                             for color in legend_colors]
            
            ax.legend(legend_elements, legend_labels, 
                     loc='upper center', 
                     bbox_to_anchor=(0.5, 1.15),
                     ncol=min(5, len(meses)),
                     fontsize=9,
                     frameon=True,
                     fancybox=True,
                     shadow=True)
            
            # Ajustar el layout
            plt.tight_layout()
            
            # Guardar como JPG
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / f"grafico_pastel_ans_{self.mes}_{self.anio}.jpg"
            plt.savefig(output_path, dpi=150, bbox_inches='tight', format='jpg')
            plt.close()
            
            return str(output_path)
        except Exception as e:
            print(f"[WARNING] Error al generar gráfico de pastel: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _generar_graficos_barras_ans(self, datos_actuales: Dict[str, Any], 
                                      datos_historicos: List[Dict[str, Any]], 
                                      output_dir: Path) -> Dict[str, Optional[str]]:
        """
        Genera 6 gráficos de barras (waterfall) para cada tipo de ANS
        Retorna un diccionario con las rutas de los archivos JPG generados
        """
        def safe_float(value, default=0.0):
            """Convierte valor a float de forma segura"""
            if value is None:
                return default
            try:
                return float(value)
            except (ValueError, TypeError):
                return default
        
        # Definir los 6 tipos de ANS
        tipos_ans = [
            {
                "nombre": "ANS CALIDAD DE LOS REPORTES ENTREGADOS",
                "key_cant": ["cant_calidad_informes_1", "cant_calidad_informes_2", 
                            "cant_calidad_informes_3", "cant_calidad_informes_4"],
                "key_valor": ["valor_calidad_informes_1", "valor_calidad_informes_2",
                             "valor_calidad_informes_3", "valor_calidad_informes_4"],
                "key_hist_cant": "calidad_reportes",
                "key_hist_valor": "calidad_reportes"
            },
            {
                "nombre": "ANS DISPONIBILIDAD DEL SISTEMA",
                "key_cant": ["cant_disponibilidad_sistema"],
                "key_valor": ["valor_disponibilidad_sistema"],
                "key_hist_cant": "disponibilidad",
                "key_hist_valor": "disponibilidad"
            },
            {
                "nombre": "ANS OPORTUNIDAD DE LOS REPORTES ENTREGADOS",
                "key_cant": ["cant_oportunidad_informes"],
                "key_valor": ["valor_oportunidad_informes"],
                "key_hist_cant": "oportunidad_reportes",
                "key_hist_valor": "oportunidad_reportes"
            },
            {
                "nombre": "ANS OPORTUNIDAD EN LAS ACTIVIDADES DE MANTENIMIENTO PREVENTIVO",
                "key_cant": ["cant_oportunidad_actividades"],
                "key_valor": ["valor_oportunidad_actividades"],
                "key_hist_cant": "oportunidad_actividades",
                "key_hist_valor": "oportunidad_actividades"
            },
            {
                "nombre": "ANS RTO",
                "key_cant": ["cant_rto_1", "cant_rto_2", "cant_rto_3", "cant_rto_4"],
                "key_valor": ["valor_rto_1", "valor_rto_2", "valor_rto_3", "valor_rto_4"],
                "key_hist_cant": "rto",
                "key_hist_valor": "rto"
            },
            {
                "nombre": "ANS TIEMPO DE RESTAURACIÓN DE SERVICIOS EN EL DATA CENTER",
                "key_cant": ["cant_tiempo_restauracion"],
                "key_valor": ["valor_tiempo_restauracion"],
                "key_hist_cant": "tiempo_restauracion",
                "key_hist_valor": "tiempo_restauracion"
            }
        ]
        
        rutas_graficos = {}
        
        print(f"[DEBUG] Iniciando generación de gráficos de barras. Datos históricos: {len(datos_historicos)} meses")
        
        for tipo_ans in tipos_ans:
            try:
                print(f"[DEBUG] Procesando gráfico para: {tipo_ans['nombre']}")
                
                # Preparar datos para el gráfico
                meses = []
                cantidades = []
                valores = []
                
                # Agregar meses históricos
                for historico in datos_historicos:
                    mes_nombre = historico.get("mes_nombre", "")
                    anio = historico.get("anio", "")
                    indicadores = historico.get("indicadores", {})
                    
                    # Obtener cantidad y valor del histórico
                    indicador = indicadores.get(tipo_ans["key_hist_cant"], {})
                    cant = indicador.get("cant", 0) if isinstance(indicador, dict) else 0
                    valor = indicador.get("valor", 0) if isinstance(indicador, dict) else 0
                    
                    # Agregar siempre, incluso si es 0 (para mostrar todos los meses)
                    meses.append(f"{mes_nombre}-{anio}")
                    cantidades.append(cant)
                    valores.append(valor)
                
                # Agregar mes actual
                mes_nombre_actual = config.MESES[self.mes].lower()[:3]
                
                # Calcular cantidad y valor del mes actual (sumar si hay múltiples keys)
                cant_actual = sum(safe_float(datos_actuales.get(key, 0)) for key in tipo_ans["key_cant"])
                valor_actual = sum(safe_float(datos_actuales.get(key, 0)) for key in tipo_ans["key_valor"])
                
                # Agregar siempre el mes actual, incluso si es 0
                meses.append(f"{mes_nombre_actual}-{self.anio}")
                cantidades.append(cant_actual)
                valores.append(valor_actual)
                
                print(f"[DEBUG] {tipo_ans['nombre']}: {len(meses)} meses, valores: {valores}")
                
                # Si no hay meses, continuar con el siguiente tipo
                if not meses:
                    print(f"[WARNING] No hay meses para {tipo_ans['nombre']}, saltando...")
                    rutas_graficos[tipo_ans["nombre"]] = None
                    continue
                
                # Asegurar que tenemos valores (llenar con 0 si faltan)
                while len(valores) < len(meses):
                    valores.append(0.0)
                while len(cantidades) < len(meses):
                    cantidades.append(0.0)
                
                # Crear figura con espacio para la tabla
                fig = plt.figure(figsize=(16, 10))
                gs = fig.add_gridspec(2, 1, height_ratios=[3, 1], hspace=0.3)
                ax = fig.add_subplot(gs[0])
                ax_table = fig.add_subplot(gs[1])
                ax_table.axis('off')
                
                # Crear gráfico de barras (valores negativos hacia abajo)
                x_pos = np.arange(len(meses))
                bars = ax.bar(x_pos, valores, color='#2E75B6', width=0.6)
                
                # Configurar eje Y
                max_valor = max(abs(v) for v in valores) if valores else 1
                min_valor = min(valores) if valores else 0
                
                # Calcular límites del eje Y correctamente
                if min_valor < 0:
                    y_min = min_valor * 1.15
                else:
                    y_min = -max_valor * 0.1  # Un poco de espacio negativo para visualización
                
                y_max = max_valor * 1.15 if max_valor > 0 else 1
                
                ax.set_ylim(y_min, y_max)
                ax.axhline(y=0, color='black', linewidth=1)
                
                # Formatear eje Y con formato de moneda (usar puntos como separadores de miles)
                def formatear_moneda_y(x, p):
                    if abs(x) >= 1000:
                        return f'${x/1000:,.0f}K'.replace(',', '.')
                    else:
                        return f'${x:,.0f}'.replace(',', '.')
                
                ax.yaxis.set_major_formatter(plt.FuncFormatter(formatear_moneda_y))
                
                # Ajustar número de ticks en el eje Y
                ax.yaxis.set_major_locator(plt.MaxNLocator(8))
                
                # Etiquetas en el eje X (solo el mes, sin texto adicional)
                ax.set_xticks(x_pos)
                ax.set_xticklabels(meses, rotation=45, ha='right', fontsize=9)
                ax.set_xlabel('')  # Eliminar etiqueta del eje X para evitar duplicación
                
                # Agregar etiquetas en las barras
                for i, (bar, cant, valor) in enumerate(zip(bars, cantidades, valores)):
                    if valor != 0:
                        # Formatear cantidad (mostrar hasta 4 decimales, eliminar ceros innecesarios)
                        cant_str = f'{cant:.4f}'.rstrip('0').rstrip('.')
                        if '.' in cant_str:
                            cant_str = cant_str.replace('.', ',')
                        
                        # Etiqueta con la cantidad (arriba de la barra)
                        if valor > 0:
                            y_cant = valor + (y_max - valor) * 0.05
                        else:
                            y_cant = abs(valor) * 0.05
                        
                        ax.text(bar.get_x() + bar.get_width()/2, y_cant,
                               cant_str,
                               ha='center', va='bottom', fontsize=9, fontweight='bold', color='black')
                        
                        # Etiqueta con el valor monetario (dentro de la barra)
                        if valor != 0:
                            # Para valores positivos, poner dentro de la barra
                            if valor > 0:
                                y_valor = valor / 2  # Mitad de la barra
                                va = 'center'
                            else:
                                # Para valores negativos, poner al final de la barra
                                y_valor = valor - abs(valor) * 0.1
                                va = 'top'
                            
                            ax.text(bar.get_x() + bar.get_width()/2, y_valor,
                                   formato_moneda_cop(valor),
                                   ha='center', va=va, 
                                   fontsize=8, fontweight='bold', color='white')
                
                # Título
                ax.set_title(tipo_ans["nombre"], fontsize=14, fontweight='bold', pad=20)
                
                # Grid
                ax.grid(True, axis='y', linestyle='--', alpha=0.3)
                ax.set_facecolor('#f5f5f5')
                
                # Crear tabla debajo del gráfico
                # Headers (solo el mes, sin el nombre completo para evitar repetición)
                headers = [mes for mes in meses]
                
                # Fila de cantidades (formatear con coma como separador decimal)
                cantidades_str = []
                for cant in cantidades:
                    if cant > 0:
                        cant_str = f'{cant:.4f}'.rstrip('0').rstrip('.')
                        if '.' in cant_str:
                            cant_str = cant_str.replace('.', ',')
                        cantidades_str.append(cant_str)
                    else:
                        cantidades_str.append('')
                
                # Fila de valores
                valores_str = [formato_moneda_cop(valor) if valor != 0 else '' for valor in valores]
                
                # Crear tabla: cellText solo contiene las filas de datos (sin headers)
                tabla_data = [cantidades_str, valores_str]
                
                # Crear tabla con headers como colLabels
                tabla = ax_table.table(cellText=tabla_data,
                                      rowLabels=['CANTIDAD', 'VALOR'],
                                      colLabels=headers,
                                      cellLoc='center',
                                      loc='center',
                                      bbox=[0, 0, 1, 1])
                tabla.auto_set_font_size(False)
                tabla.set_fontsize(7)
                tabla.scale(1, 2.5)
                
                # Ajustar ancho de columnas automáticamente
                for i in range(len(headers)):
                    tabla.auto_set_column_width(i)
                
                # Estilo de la tabla
                # Headers (fila 0 son los colLabels) - ya son solo los meses, no necesitan truncamiento
                for i in range(len(headers)):
                    tabla[(0, i)].set_facecolor('#1F4E79')
                    tabla[(0, i)].set_text_props(weight='bold', color='white', fontsize=7)
                
                # Filas de datos (fila 1 = CANTIDAD, fila 2 = VALOR)
                for i in range(1, 3):
                    for j in range(len(headers)):
                        if i % 2 == 0:
                            tabla[(i, j)].set_facecolor('#f2f2f2')
                        else:
                            tabla[(i, j)].set_facecolor('#ffffff')
                        tabla[(i, j)].set_text_props(fontsize=7)
                
                # Guardar gráfico
                output_dir.mkdir(parents=True, exist_ok=True)
                nombre_archivo = tipo_ans["nombre"].lower().replace(" ", "_").replace("ans_", "")
                output_path = output_dir / f"grafico_barras_{nombre_archivo}_{self.mes}_{self.anio}.jpg"
                plt.savefig(output_path, dpi=150, bbox_inches='tight', format='jpg')
                plt.close()
                
                rutas_graficos[tipo_ans["nombre"]] = str(output_path)
                print(f"[OK] Gráfico de barras generado: {output_path}")
                
            except Exception as e:
                print(f"[WARNING] Error al generar gráfico de barras para {tipo_ans['nombre']}: {e}")
                import traceback
                traceback.print_exc()
                rutas_graficos[tipo_ans["nombre"]] = None
        
        return rutas_graficos
    
    def _obtener_datos_historicos_ans(self) -> List[Dict[str, Any]]:
        """
        Obtiene datos históricos de ANS desde GLPI (por ahora retorna datos dummy)
        
        Returns:
            Lista de diccionarios con datos históricos por mes
        """
        # TODO: Implementar obtención real desde GLPI
        # Por ahora retornamos datos dummy
        datos_dummy = []
        
        # Generar datos para los últimos 6 meses (excluyendo el mes actual)
        # Orden descendente: mes más reciente primero
        meses_anteriores = []
        mes_actual = self.mes
        anio_actual = self.anio
        
        for i in range(1, 7):  # Últimos 6 meses
            mes = mes_actual - i
            anio = anio_actual
            if mes <= 0:
                mes += 12
                anio -= 1
            
            mes_nombre = config.MESES[mes].lower()[:3]  # Primeras 3 letras del mes
            datos_dummy.append({
                "mes": mes,
                "anio": anio,
                "mes_nombre": mes_nombre,
                "total_cant": 150 + (i * 10),  # Valores dummy
                "total_valor": 5000000 + (i * 100000),
                "indicadores": {
                    "calidad_reportes": {"cant": 25 + i, "valor": 800000 + (i * 10000)},
                    "disponibilidad": {"cant": 30 + i, "valor": 1000000 + (i * 20000)},
                    "oportunidad_reportes": {"cant": 20 + i, "valor": 600000 + (i * 8000)},
                    "oportunidad_actividades": {"cant": 15 + i, "valor": 500000 + (i * 6000)},
                    "rto": {"cant": 35 + i, "valor": 1200000 + (i * 25000)},
                    "tiempo_restauracion": {"cant": 25 + i, "valor": 900000 + (i * 12000)}
                }
            })
        
        # Retornar en orden ascendente (del más antiguo al más reciente)
        return datos_dummy
  
    def _preparar_datos_tabla_indicadores_ans(self, datos_actuales: Dict[str, Any], 
                                               datos_historicos: List[Dict[str, Any]]) -> List[List[str]]:
        """
        Prepara los datos de la tabla de indicadores ANS
        
        Args:
            datos_actuales: Datos del mes actual desde MongoDB
            datos_historicos: Lista de datos históricos (meses anteriores)
        
        Returns:
            Lista de listas con los datos de la tabla (primera fila son headers)
        """
        def safe_float(value, default=0.0):
            """Convierte valor a float de forma segura"""
            if value is None:
                return default
            try:
                return float(value)
            except (ValueError, TypeError):
                return default
        
        # Calcular totales del mes actual
        total_cant_actual = 0
        total_valor_actual = 0.0
        
        # ANS CALIDAD DE LOS REPORTES ENTREGADOS (suma de 4 valores)
        cant_calidad = (
            safe_float(datos_actuales.get('cant_calidad_informes_1', 0)) +
            safe_float(datos_actuales.get('cant_calidad_informes_2', 0)) +
            safe_float(datos_actuales.get('cant_calidad_informes_3', 0)) +
            safe_float(datos_actuales.get('cant_calidad_informes_4', 0))
        )
        valor_calidad = (
            safe_float(datos_actuales.get('valor_calidad_informes_1', 0)) +
            safe_float(datos_actuales.get('valor_calidad_informes_2', 0)) +
            safe_float(datos_actuales.get('valor_calidad_informes_3', 0)) +
            safe_float(datos_actuales.get('valor_calidad_informes_4', 0))
        )
        total_cant_actual += cant_calidad
        total_valor_actual += valor_calidad
        
        # ANS DISPONIBILIDAD DEL SISTEMA
        cant_disponibilidad = safe_float(datos_actuales.get('cant_disponibilidad_sistema', 0))
        valor_disponibilidad = safe_float(datos_actuales.get('valor_disponibilidad_sistema', 0))
        total_cant_actual += cant_disponibilidad
        total_valor_actual += valor_disponibilidad
        
        # ANS OPORTUNIDAD DE LOS REPORTES ENTREGADOS
        cant_oportunidad_informes = safe_float(datos_actuales.get('cant_oportunidad_informes', 0))
        valor_oportunidad_informes = safe_float(datos_actuales.get('valor_oportunidad_informes', 0))
        total_cant_actual += cant_oportunidad_informes
        total_valor_actual += valor_oportunidad_informes
        
        # ANS OPORTUNIDAD EN LAS ACTIVIDADES DE MANTENIMIENTO PREVENTIVO
        cant_oportunidad_actividades = safe_float(datos_actuales.get('cant_oportunidad_actividades', 0))
        valor_oportunidad_actividades = safe_float(datos_actuales.get('valor_oportunidad_actividades', 0))
        total_cant_actual += cant_oportunidad_actividades
        total_valor_actual += valor_oportunidad_actividades
        
        # ANS RTO (suma de 4 valores)
        cant_rto = (
            safe_float(datos_actuales.get('cant_rto_1', 0)) +
            safe_float(datos_actuales.get('cant_rto_2', 0)) +
            safe_float(datos_actuales.get('cant_rto_3', 0)) +
            safe_float(datos_actuales.get('cant_rto_4', 0))
        )
        valor_rto = (
            safe_float(datos_actuales.get('valor_rto_1', 0)) +
            safe_float(datos_actuales.get('valor_rto_2', 0)) +
            safe_float(datos_actuales.get('valor_rto_3', 0)) +
            safe_float(datos_actuales.get('valor_rto_4', 0))
        )
        total_cant_actual += cant_rto
        total_valor_actual += valor_rto
        
        # ANS TIEMPO DE RESTAURACIÓN DE SERVICIOS EN EL DATA CENTER
        cant_tiempo_restauracion = safe_float(datos_actuales.get('cant_tiempo_restauracion', 0))
        valor_tiempo_restauracion = safe_float(datos_actuales.get('valor_tiempo_restauracion', 0))
        total_cant_actual += cant_tiempo_restauracion
        total_valor_actual += valor_tiempo_restauracion
        
        # Preparar datos de la tabla
        table_data = []
        
        # Headers
        table_data.append(["TIPO INDICADOR", "CANTIDAD", "VALOR"])
        
        # Inicializar totales generales
        total_cant_general = total_cant_actual
        total_valor_general = total_valor_actual
        
        # PRIMERO: Agregar meses anteriores desde datos históricos (orden ascendente: más antiguo primero)
        for historico in datos_historicos:
            mes_nombre_hist = historico.get("mes_nombre", "")
            anio_hist = historico.get("anio", "")
            total_cant_hist = historico.get("total_cant", 0)
            total_valor_hist = historico.get("total_valor", 0)
            indicadores_hist = historico.get("indicadores", {})
            
            # Fila de totales del mes histórico
            fila_total_hist = [
                f"{mes_nombre_hist}-{anio_hist}",
                formato_cantidad(total_cant_hist),
                formato_moneda_cop(total_valor_hist)
            ]
            table_data.append(fila_total_hist)
            
            # Filas de indicadores del mes histórico
            indicadores_hist_list = [
                ("ANS CALIDAD DE LOS REPORTES ENTREGADOS", 
                 indicadores_hist.get("calidad_reportes", {}).get("cant", 0),
                 indicadores_hist.get("calidad_reportes", {}).get("valor", 0)),
                ("ANS DISPONIBILIDAD DEL SISTEMA",
                 indicadores_hist.get("disponibilidad", {}).get("cant", 0),
                 indicadores_hist.get("disponibilidad", {}).get("valor", 0)),
                ("ANS OPORTUNIDAD DE LOS REPORTES ENTREGADOS",
                 indicadores_hist.get("oportunidad_reportes", {}).get("cant", 0),
                 indicadores_hist.get("oportunidad_reportes", {}).get("valor", 0)),
                ("ANS OPORTUNIDAD EN LAS ACTIVIDADES DE MANTENIMIENTO PREVENTIVO",
                 indicadores_hist.get("oportunidad_actividades", {}).get("cant", 0),
                 indicadores_hist.get("oportunidad_actividades", {}).get("valor", 0)),
                ("ANS RTO",
                 indicadores_hist.get("rto", {}).get("cant", 0),
                 indicadores_hist.get("rto", {}).get("valor", 0)),
                ("ANS TIEMPO DE RESTAURACIÓN DE SERVICIOS EN EL DATA CENTER",
                 indicadores_hist.get("tiempo_restauracion", {}).get("cant", 0),
                 indicadores_hist.get("tiempo_restauracion", {}).get("valor", 0))
            ]
            
            for nombre_hist, cant_hist, valor_hist in indicadores_hist_list:
                table_data.append([
                    nombre_hist,
                    formato_cantidad(cant_hist) if cant_hist > 0 else "",
                    formato_moneda_cop(valor_hist) if valor_hist > 0 else ""
                ])
            
            # Acumular para Total GENERAL
            total_cant_general += total_cant_hist
            total_valor_general += total_valor_hist
        
        # SEGUNDO: Agregar mes actual (al final, antes del Total GENERAL)
        mes_nombre = config.MESES[self.mes].lower()[:3]
        fila_total = [
            f"{mes_nombre}-{self.anio}",
            formato_cantidad(total_cant_actual),
            formato_moneda_cop(total_valor_actual)
        ]
        table_data.append(fila_total)
        
        # Filas de indicadores del mes actual
        indicadores = [
            ("ANS CALIDAD DE LOS REPORTES ENTREGADOS", cant_calidad, valor_calidad),
            ("ANS DISPONIBILIDAD DEL SISTEMA", cant_disponibilidad, valor_disponibilidad),
            ("ANS OPORTUNIDAD DE LOS REPORTES ENTREGADOS", cant_oportunidad_informes, valor_oportunidad_informes),
            ("ANS OPORTUNIDAD EN LAS ACTIVIDADES DE MANTENIMIENTO PREVENTIVO", cant_oportunidad_actividades, valor_oportunidad_actividades),
            ("ANS RTO", cant_rto, valor_rto),
            ("ANS TIEMPO DE RESTAURACIÓN DE SERVICIOS EN EL DATA CENTER", cant_tiempo_restauracion, valor_tiempo_restauracion)
        ]
        
        for nombre, cant, valor in indicadores:
            table_data.append([
                nombre,
                formato_cantidad(cant) if cant > 0 else "",
                formato_moneda_cop(valor) if valor > 0 else ""
            ])
        
        # TERCERO: Fila de Total GENERAL al final (suma de todos los meses)
        fila_total_general = [
            "Total GENERAL",
            formato_cantidad(total_cant_general),
            formato_moneda_cop(total_valor_general)
        ]
        table_data.append(fila_total_general)
        
        return table_data
    
    def _es_fila_totales(self, primer_valor: str) -> bool:
        """
        Determina si una fila es una fila de totales basándose en el primer valor
        
        Args:
            primer_valor: Primer valor de la fila (columna TIPO INDICADOR)
        
        Returns:
            True si es una fila de totales, False en caso contrario
        """
        if not primer_valor:
            return False
        
        primer_valor_upper = primer_valor.upper()
        
        # Verificar si es "Total GENERAL"
        if primer_valor_upper == "TOTAL GENERAL":
            return True
        
        # Verificar si tiene formato "mes-año" (ej: "dic-25", "nov-25")
        # Patrón: 3 letras, guion, 2 dígitos
        patron_mes_anio = r'^[a-z]{3}-\d{2,4}$'
        if re.match(patron_mes_anio, primer_valor.lower()):
            return True
        
        return False
    
    def _reemplazar_placeholder_con_tabla(self, doc: Document, placeholder: str, table_data: List[List[str]]):
        """
        Busca un placeholder en el documento y lo reemplaza con una tabla creada programáticamente.
        
        Args:
            doc: Documento Word (después de renderizar)
            placeholder: Texto del placeholder a buscar (ej: "[[TABLE_ANS_INDICADORES]]")
            table_data: Lista de listas con los datos de la tabla (primera fila son headers)
        """
        if not table_data or len(table_data) == 0:
            return
        
        # Estilos de tabla comunes en Word (en orden de preferencia)
        estilos_tabla = ['Table Grid', 'Light Shading', 'Light List', 'Medium Shading 1', 'Light Grid']
        
        # Buscar el placeholder en todos los párrafos
        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder in paragraph.text:
                # Crear la tabla
                tabla = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                
                # Intentar aplicar un estilo de tabla disponible
                estilo_aplicado = False
                for estilo in estilos_tabla:
                    try:
                        tabla.style = estilo
                        estilo_aplicado = True
                        break
                    except:
                        continue
                
                # Si ningún estilo funcionó, usar el estilo por defecto
                if not estilo_aplicado:
                    try:
                        tabla.style = 'Table Grid'
                    except:
                        pass
                
                # Llenar la tabla
                total_rows = len(table_data)
                for row_idx, fila in enumerate(table_data):
                    for col_idx, valor in enumerate(fila):
                        celda = tabla.rows[row_idx].cells[col_idx]
                        celda.text = str(valor) if valor is not None else ""
                        
                        # Formatear headers (primera fila) - Azul oscuro con texto blanco
                        if row_idx == 0:
                            self._set_cell_shading(celda, "1F4E79")
                            self._centrar_celda_vertical(celda)
                            for paragraph_celda in celda.paragraphs:
                                paragraph_celda.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph_celda.runs:
                                    run.bold = True
                                    run.font.size = Pt(8)
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                        elif self._es_fila_totales(fila[0]):  # Fila de totales (mes-año o Total GENERAL)
                            # Fila de totales - Fondo azul claro con texto en negrita
                            self._set_cell_shading(celda, "D9E1F2")
                            self._centrar_celda_vertical(celda)
                            for paragraph_celda in celda.paragraphs:
                                paragraph_celda.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph_celda.runs:
                                    run.bold = True
                                    run.font.size = Pt(8)
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                        else:
                            # Filas de datos - Fondo gris claro alternado
                            if row_idx % 2 == 0:
                                self._set_cell_shading(celda, "F2F2F2")
                            else:
                                self._set_cell_shading(celda, "FFFFFF")
                            self._centrar_celda_vertical(celda)
                            for paragraph_celda in celda.paragraphs:
                                paragraph_celda.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
                                for run in paragraph_celda.runs:
                                    run.font.size = Pt(8)
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                
                # Insertar la tabla después del párrafo que contiene el placeholder
                parent = paragraph._element.getparent()
                para_idx = parent.index(paragraph._element)
                parent.insert(para_idx + 1, tabla._element)
                
                # Eliminar el placeholder del párrafo
                paragraph.text = paragraph.text.replace(placeholder, "").strip()
                
                # Si el párrafo quedó vacío, eliminarlo
                if not paragraph.text.strip():
                    parent.remove(paragraph._element)
                
                break

  
    
    def generar(self, document: Optional[Dict[str, Any]] = None, output_path: Optional[Path] = None) -> DocxTemplate:
      
      

        datos_seccion = self.procesar(document)
        
        # Cargar template
        template_path = config.TEMPLATES_DIR / self.template_file
        if not template_path.exists():
            raise FileNotFoundError(f"Template no encontrado: {template_path}")
        
        template = DocxTemplate(str(template_path))
        
        # Generar gráficos
        output_dir = output_path.parent if output_path else config.OUTPUT_DIR / "seccion_3"
        grafico_lineas_path = self._generar_grafico_lineas_ans(output_dir)
        
        # Obtener datos históricos para los gráficos
        datos_historicos = self._obtener_datos_historicos_ans()
        grafico_pastel_path = self._generar_grafico_pastel_ans(datos_seccion, datos_historicos, output_dir)
        
        # Generar gráficos de barras para cada tipo de ANS
        graficos_barras = self._generar_graficos_barras_ans(datos_seccion, datos_historicos, output_dir)
        
        # Preparar contexto para el template
        contexto = self._cargar_contexto_base()
        
        print(f"datos_seccion: {datos_seccion}")
        contexto.update(
            **datos_seccion,
        )
        
        # Agregar gráfico de líneas como InlineImage si existe
        if grafico_lineas_path and Path(grafico_lineas_path).exists():
            try:
                contexto["grafico_lineas"] = InlineImage(template, grafico_lineas_path, width=Mm(150))
            except Exception as e:
                print(f"[WARNING] Error al crear InlineImage del gráfico de líneas: {e}")
                contexto["grafico_lineas"] = None
        else:
            contexto["grafico_lineas"] = None
        
        # Agregar gráfico de pastel como InlineImage si existe
        if grafico_pastel_path and Path(grafico_pastel_path).exists():
            try:
                contexto["grafico_pastel"] = InlineImage(template, grafico_pastel_path, width=Mm(150))
            except Exception as e:
                print(f"[WARNING] Error al crear InlineImage del gráfico de pastel: {e}")
                contexto["grafico_pastel"] = None
        else:
            contexto["grafico_pastel"] = None
        
        # Agregar gráficos de barras como InlineImage si existen
        nombres_ans = [
            "ANS CALIDAD DE LOS REPORTES ENTREGADOS",
            "ANS DISPONIBILIDAD DEL SISTEMA",
            "ANS OPORTUNIDAD DE LOS REPORTES ENTREGADOS",
            "ANS OPORTUNIDAD EN LAS ACTIVIDADES DE MANTENIMIENTO PREVENTIVO",
            "ANS RTO",
            "ANS TIEMPO DE RESTAURACIÓN DE SERVICIOS EN EL DATA CENTER"
        ]
        
        print(f"[DEBUG] Agregando gráficos de barras al contexto. Total generados: {len([k for k, v in graficos_barras.items() if v])}")
        
        for nombre_ans in nombres_ans:
            key_contexto = nombre_ans.lower().replace(" ", "_").replace("ans_", "grafico_barras_")
            if nombre_ans in graficos_barras and graficos_barras[nombre_ans] and Path(graficos_barras[nombre_ans]).exists():
                try:
                    contexto[key_contexto] = InlineImage(template, graficos_barras[nombre_ans], width=Mm(150))
                    print(f"[OK] Gráfico de barras agregado al contexto: {key_contexto}")
                except Exception as e:
                    print(f"[WARNING] Error al crear InlineImage del gráfico de barras {nombre_ans}: {e}")
                    import traceback
                    traceback.print_exc()
                    contexto[key_contexto] = None
            else:
                print(f"[WARNING] Gráfico de barras no encontrado para {nombre_ans}: {graficos_barras.get(nombre_ans, 'No en diccionario')}")
                contexto[key_contexto] = None
        
        # Preparar placeholder para tabla de indicadores
        contexto["table_ans_indicadores_placeholder"] = "[[TABLE_ANS_INDICADORES]]"
        
        # Renderizar template
        template.render(contexto)
        
        # Guardar temporalmente para poder trabajar con el documento renderizado
        if output_path:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            temp_path = str(output_path).replace('.docx', '_temp.docx')
            template.save(temp_path)
            
            # Abrir el documento renderizado para insertar tabla programáticamente
            doc = Document(temp_path)
            
            # Los datos históricos ya se obtuvieron arriba para el gráfico de pastel
            
            # Preparar datos de la tabla
            table_data = self._preparar_datos_tabla_indicadores_ans(datos_seccion, datos_historicos)
            
            # Reemplazar placeholder de tabla con tabla creada programáticamente
            self._reemplazar_placeholder_con_tabla(doc, "[[TABLE_ANS_INDICADORES]]", table_data)
            
            # Guardar el documento final
            doc.save(str(output_path))
            
            # Eliminar archivo temporal
            try:
                Path(temp_path).unlink()
            except:
                pass
            
            print(f"[OK] {self.nombre_seccion} guardada en: {output_path}")
        else:
            # Si no hay output_path, solo renderizar
            pass
        
        return template
    
    def guardar(self, output_path: Path) -> None:
        """
        Genera y guarda la sección usando template
        """
        self.generar(output_path=output_path)
