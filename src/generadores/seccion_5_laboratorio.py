"""
Generador Secci√≥n 5: Laboratorio
Tipo: üü© EXTRACCI√ìN (datos desde MongoDB)
"""
from io import BytesIO
from pathlib import Path
from typing import Dict, Any, List, Optional
from .base import GeneradorSeccion
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate
import config
import logging

logger = logging.getLogger(__name__)


class GeneradorSeccion5(GeneradorSeccion):
    """Genera la secci√≥n 5: Laboratorio"""
    
    @property
    def nombre_seccion(self) -> str:
        return "5. LABORATORIO"
    
    @property
    def template_file(self) -> str:
        return "seccion_5_laboratorio.docx"
    
    def __init__(self, anio: int, mes: int, cargar_desde_mongodb: bool = True):
        super().__init__(anio, mes)
        self.cargar_desde_mongodb = cargar_desde_mongodb
        self.datos_laboratorio_raw: List[Dict] = []
        self.reporte_laboratorio: List[Dict] = []
        self.reintegrados_inventario: List[Dict] = []
    
    def cargar_datos(self) -> None:
        """Carga datos de laboratorio desde MongoDB"""
        # Los datos se cargan desde MongoDB de forma as√≠ncrona
        # Este m√©todo se llama desde el servicio despu√©s de cargar los datos
        pass
    
    def _calcular_cantidades_por_estado(self) -> List[Dict[str, Any]]:
        """
        Calcula las cantidades seg√∫n los estados de los registros de laboratorio
        
        Returns:
            Lista de diccionarios con el reporte de laboratorio:
            [
                {"reporte": "REINTEGRADOS AL INVENTARIO", "cantidad": 10},
                {"reporte": "NO OPERATIVOS", "cantidad": 5},
                ...
            ]
        """
        if not self.datos_laboratorio_raw:
            logger.warning("No hay datos de laboratorio para calcular cantidades")
            return []
        
        # Estados para REINTEGRADOS AL INVENTARIO
        estados_reintegrados = ["OPERATIVO", "REPARADO", "REPARADO CAMPO"]
        
        # Contar registros por estado
        cantidad_reintegrados = 0
        cantidad_no_operativos = 0
        cantidad_garantia = 0
        cantidad_pendiente_parte = 0
        
        for registro in self.datos_laboratorio_raw:
            estado = str(registro.get("estado", "")).strip().upper()
            
            # REINTEGRADOS AL INVENTARIO
            if estado in estados_reintegrados:
                cantidad_reintegrados += 1
            
            # NO OPERATIVOS
            elif estado == "IRREPARABLE":
                cantidad_no_operativos += 1
            
            # ESTADO DE GARANT√çA
            elif estado == "ESTADO DE GARANT√çA":
                cantidad_garantia += 1
            
            # PENDIENTE POR PARTE
            elif estado == "PENDIENTE POR PARTE":
                cantidad_pendiente_parte += 1
        
        # Calcular total
        total = cantidad_reintegrados + cantidad_no_operativos + cantidad_garantia + cantidad_pendiente_parte
        
        # Construir lista de reporte
        reporte = [
            {
                "reporte": "REINTEGRADOS AL INVENTARIO",
                "cantidad": cantidad_reintegrados
            },
            {
                "reporte": "NO OPERATIVOS",
                "cantidad": cantidad_no_operativos
            },
            {
                "reporte": "ESTADO DE GARANT√çA",
                "cantidad": cantidad_garantia
            },
            {
                "reporte": "PENDIENTE POR PARTE",
                "cantidad": cantidad_pendiente_parte
            },
            {
                "reporte": "TOTAL",
                "cantidad": total
            }
        ]
        
        logger.info(f"Cantidades calculadas: Reintegrados={cantidad_reintegrados}, "
                   f"No Operativos={cantidad_no_operativos}, "
                   f"Garant√≠a={cantidad_garantia}, "
                   f"Pendiente={cantidad_pendiente_parte}, "
                   f"Total={total}")
        
        return reporte
    
    def _filtrar_reintegrados_inventario(self) -> List[Dict[str, Any]]:
        """
        Filtra los registros que corresponden a REINTEGRADOS AL INVENTARIO
        
        Returns:
            Lista de registros con ESTADO = "OPERATIVO", "REPARADO" o "REPARADO CAMPO"
        """
        if not self.datos_laboratorio_raw:
            logger.warning("No hay datos de laboratorio para filtrar")
            return []
        
        # Estados para REINTEGRADOS AL INVENTARIO
        estados_reintegrados = ["OPERATIVO", "REPARADO", "REPARADO CAMPO"]
        
        reintegrados = []
        for registro in self.datos_laboratorio_raw:
            estado = str(registro.get("estado", "")).strip().upper()
            if estado in estados_reintegrados:
                reintegrados.append(registro)
        
        logger.info(f"Registros REINTEGRADOS AL INVENTARIO encontrados: {len(reintegrados)}")
        return reintegrados
    
    def procesar(self) -> Dict[str, Any]:
        """Procesa los datos y retorna el contexto para el template"""
        # Calcular cantidades por estado
        self.reporte_laboratorio = self._calcular_cantidades_por_estado()
        
        # Filtrar registros REINTEGRADOS AL INVENTARIO
        self.reintegrados_inventario = self._filtrar_reintegrados_inventario()
        
        # Agregar marcadores de tabla al contexto
        contexto = {
            "TABLA_MARKER_REPORTE_LABORATORIO": "[[TABLA_REPORTE_LABORATORIO]]",
            "TABLA_MARKER_CONCEPTO_TECNICO": "[[TABLA_CONCEPTO_TECNICO]]",
            "reporte_laboratorio": self.reporte_laboratorio,
            "reintegrados_inventario": self.reintegrados_inventario
        }
        
        return contexto
    
    def generar(self):
        """
        Genera la secci√≥n completa con tablas din√°micas
        """
        # Cargar contexto base
        self.contexto = self.cargar_contexto_base()
        
        # Cargar datos espec√≠ficos
        self.cargar_datos()
        
        # Procesar y agregar al contexto
        datos_seccion = self.procesar()
        self.contexto.update(datos_seccion)
        
        # Renderizar template con Jinja2
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template no encontrado: {self.template_path}")
        
        doc_template = DocxTemplate(self.template_path)
        doc_template.render(self.contexto)
        
        # Guardar el DocxTemplate a BytesIO para evitar problemas de XML
        buffer = BytesIO()
        doc_template.save(buffer)
        buffer.seek(0)
        
        # Abrir como Document de python-docx para manipulaci√≥n de tablas
        doc = Document(buffer)
        
        # Reemplazar tabla de reporte de laboratorio
        self._reemplazar_tabla_por_marcador(
            doc,
            "TABLA_REPORTE_LABORATORIO",
            self.reporte_laboratorio,
            self._crear_tabla_reporte_laboratorio
        )
        
        # Reemplazar tabla de concepto t√©cnico (REINTEGRADOS AL INVENTARIO)
        self._reemplazar_tabla_por_marcador(
            doc,
            "TABLA_CONCEPTO_TECNICO",
            self.reintegrados_inventario,
            self._crear_tabla_concepto_tecnico
        )
        
        return doc
    
    def _reemplazar_tabla_por_marcador(self, doc: Document, marcador: str, datos: list, metodo_creacion) -> None:
        """
        Busca una tabla en el documento usando un marcador √∫nico y la reemplaza con datos.
        
        Args:
            doc: Documento de python-docx
            marcador: Nombre del marcador (ej: "TABLA_REPORTE_LABORATORIO")
            datos: Lista de datos para llenar la tabla
            metodo_creacion: M√©todo que crea/llena la tabla
        """
        # El marcador se renderiza como [[TABLA_XXX]] despu√©s de procesar Jinja2
        marcador_renderizado = f"[[{marcador}]]"
        
        # Variaciones del marcador que pueden aparecer
        marcador_variaciones = [
            marcador_renderizado,  # Despu√©s de procesar Jinja2: [[TABLA_XXX]]
            f"{{{{ TABLA_MARKER_{marcador.replace('TABLA_', '')} }}}}",  # En el template original
            f"TABLA_MARKER_{marcador.replace('TABLA_', '')}",  # Variable de Jinja2 sin renderizar
            marcador.upper(),  # Nombre del marcador en may√∫sculas
        ]
        
        # Normalizar el marcador para b√∫squeda
        marcador_busqueda = marcador.upper().replace("TABLA_", "")
        logger.info(f"Buscando tabla con marcador: {marcador_renderizado} (variaciones: {marcador_variaciones})")
        
        tabla_encontrada = None
        tabla_idx = None
        celda_con_marcador = None
        
        # Buscar el marcador en todas las tablas del documento
        for idx, tabla in enumerate(doc.tables):
            # Buscar en todas las celdas de la tabla
            for fila_idx, fila in enumerate(tabla.rows):
                for celda_idx, celda in enumerate(fila.cells):
                    # Obtener todo el texto de la celda (incluyendo p√°rrafos)
                    texto_celda = ""
                    for parrafo in celda.paragraphs:
                        texto_celda += parrafo.text
                    
                    texto_celda_upper = texto_celda.upper()
                    
                    # Verificar si contiene el marcador renderizado [[TABLA_XXX]]
                    encontro_marcador = marcador_renderizado.upper() in texto_celda_upper
                    
                    # Si no se encuentra, buscar por variaciones
                    if not encontro_marcador:
                        for variacion in marcador_variaciones:
                            if variacion.upper() in texto_celda_upper:
                                encontro_marcador = True
                                break
                    
                    # Tambi√©n buscar por el nombre del marcador sin prefijos
                    if not encontro_marcador and marcador_busqueda in texto_celda_upper:
                        encontro_marcador = True
                    
                    if encontro_marcador:
                        tabla_encontrada = tabla
                        tabla_idx = idx
                        celda_con_marcador = (fila_idx, celda_idx)
                        logger.info(f"Marcador encontrado en tabla {idx}, fila {fila_idx}, celda {celda_idx}")
                        logger.debug(f"Texto de la celda: '{texto_celda[:100]}...'")
                        break
                
                if tabla_encontrada:
                    break
            
            if tabla_encontrada:
                break
        
        if tabla_encontrada and tabla_idx is not None:
            # Limpiar el marcador si est√° en el encabezado (fila 0)
            if celda_con_marcador:
                fila_idx, celda_idx = celda_con_marcador
                if fila_idx == 0:
                    # Si el marcador est√° en el encabezado, limpiarlo
                    celda = tabla_encontrada.rows[0].cells[celda_idx]
                    for parrafo in celda.paragraphs:
                        parrafo.clear()
            
            # Llamar al m√©todo de creaci√≥n para reemplazar la tabla
            logger.info(f"Reemplazando tabla {tabla_idx} con {len(datos)} elementos")
            try:
                metodo_creacion(doc, tabla_encontrada)
                logger.info(f"Tabla '{marcador}' procesada correctamente")
            except Exception as e:
                logger.error(f"Error al procesar tabla '{marcador}': {str(e)}")
                import traceback
                traceback.print_exc()
        else:
            logger.warning(f"No se encontr√≥ tabla con marcador '{marcador}'")
            logger.info(f"Aseg√∫rate de agregar '{{{{ TABLA_MARKER_{marcador.replace('TABLA_', '')} }}}}' en la primera celda de datos de la tabla en el template")
            logger.debug(f"Marcador buscado: {marcador_variaciones}")
    
    def _crear_tabla_reporte_laboratorio(self, doc: Document, tabla_existente) -> None:
        """
        Reemplaza el contenido de la tabla existente con los datos del reporte de laboratorio
        
        Args:
            doc: Documento de python-docx
            tabla_existente: Tabla existente en el documento
        """
        logger.info(f"Reemplazando tabla de reporte de laboratorio con {len(self.reporte_laboratorio)} filas")
        
        # Limpiar todas las filas excepto el encabezado (fila 0)
        num_filas_originales = len(tabla_existente.rows)
        while len(tabla_existente.rows) > 1:
            tbl = tabla_existente._tbl
            tbl.remove(tabla_existente.rows[-1]._tr)
        
        logger.info(f"Tabla limpiada: {num_filas_originales} filas -> {len(tabla_existente.rows)} fila(s) (encabezado)")
        
        # Obtener n√∫mero de columnas
        num_cols = len(tabla_existente.columns)
        logger.info(f"Tabla tiene {num_cols} columnas")
        
        # Verificar/actualizar encabezados
        if len(tabla_existente.rows) > 0:
            encabezados_esperados = ["REPORTE LABORATORIO", "CANTIDAD"]
            primera_fila = tabla_existente.rows[0]
            
            # Actualizar encabezados si no coinciden
            for i, header_text in enumerate(encabezados_esperados):
                if i < num_cols and i < len(primera_fila.cells):
                    celda = primera_fila.cells[i]
                    texto_actual = celda.text.strip().upper()
                    if texto_actual != header_text.upper():
                        celda.text = header_text
                        # Formatear encabezado
                        for parrafo in celda.paragraphs:
                            for run in parrafo.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)
                            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar filas con los datos
        for item in self.reporte_laboratorio:
            fila = tabla_existente.add_row()
            celdas = fila.cells
            
            # Columna 1: REPORTE LABORATORIO
            if len(celdas) > 0:
                celda_reporte = celdas[0]
                celda_reporte.text = item.get("reporte", "")
                # Formatear celda
                for parrafo in celda_reporte.paragraphs:
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in parrafo.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'
            
            # Columna 2: CANTIDAD
            if len(celdas) > 1:
                celda_cantidad = celdas[1]
                cantidad = item.get("cantidad", 0)
                celda_cantidad.text = str(cantidad)
                # Formatear celda
                for parrafo in celda_cantidad.paragraphs:
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in parrafo.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'
                        # Si es TOTAL, poner en negrita
                        if item.get("reporte", "").upper() == "TOTAL":
                            run.font.bold = True
        
        logger.info(f"Tabla actualizada: {len(tabla_existente.rows)} filas totales (1 encabezado + {len(self.reporte_laboratorio)} datos)")
    
    def _crear_tabla_concepto_tecnico(self, doc: Document, tabla_existente) -> None:
        """
        Reemplaza el contenido de la tabla existente con los datos de REINTEGRADOS AL INVENTARIO
        
        Args:
            doc: Documento de python-docx
            tabla_existente: Tabla existente en el documento
        """
        logger.info(f"Reemplazando tabla de concepto t√©cnico con {len(self.reintegrados_inventario)} registros")
        
        # Limpiar todas las filas excepto el encabezado (fila 0)
        num_filas_originales = len(tabla_existente.rows)
        while len(tabla_existente.rows) > 1:
            tbl = tabla_existente._tbl
            tbl.remove(tabla_existente.rows[-1]._tr)
        
        logger.info(f"Tabla limpiada: {num_filas_originales} filas -> {len(tabla_existente.rows)} fila(s) (encabezado)")
        
        # Obtener n√∫mero de columnas
        num_cols = len(tabla_existente.columns)
        logger.info(f"Tabla tiene {num_cols} columnas")
        
        # Verificar/actualizar encabezados
        if len(tabla_existente.rows) > 0:
            encabezados_esperados = ["ID", "FECHA", "PUNTO", "EQUIPO", "SERIAL", "ESTADO"]
            primera_fila = tabla_existente.rows[0]
            
            # Si la tabla tiene menos de 6 columnas, intentar agregar columnas faltantes
            if num_cols < 6:
                logger.warning(f"La tabla tiene solo {num_cols} columnas, se necesitan 6. Intentando agregar columnas...")
                from docx.shared import Inches
                columnas_faltantes = 6 - num_cols
                for i in range(columnas_faltantes):
                    tabla_existente.add_column(Inches(1.5))
                    logger.info(f"Columna {num_cols + i + 1} agregada")
                num_cols = len(tabla_existente.columns)
                logger.info(f"Tabla ahora tiene {num_cols} columnas")
            
            # Actualizar encabezados si no coinciden
            for i, header_text in enumerate(encabezados_esperados):
                if i < num_cols and i < len(primera_fila.cells):
                    celda = primera_fila.cells[i]
                    texto_actual = celda.text.strip().upper()
                    if texto_actual != header_text.upper():
                        celda.text = header_text
                        # Formatear encabezado
                        for parrafo in celda.paragraphs:
                            for run in parrafo.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)
                            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Mapeo de campos de los registros a las columnas
        mapeo_columnas = {
            0: 'id',
            1: 'fecha',
            2: 'punto',
            3: 'equipo',
            4: 'serial',
            5: 'estado'
        }
        
        # Agregar filas con los datos
        for registro in self.reintegrados_inventario:
            fila = tabla_existente.add_row()
            celdas = fila.cells
            
            # Llenar cada celda seg√∫n el mapeo
            for i in range(min(num_cols, 6)):
                campo = mapeo_columnas.get(i, '')
                valor = registro.get(campo, '')
                
                if i < len(celdas):
                    # Limpiar contenido existente de la celda
                    celdas[i].text = ''
                    # Agregar el nuevo texto
                    parrafo = celdas[i].paragraphs[0] if celdas[i].paragraphs else celdas[i].add_paragraph()
                    parrafo.clear()
                    run = parrafo.add_run(str(valor) if valor else '')
                    
                    # Aplicar formato seg√∫n la columna
                    if i == 0:  # ID
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    elif i == 1:  # FECHA
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    elif i == 2:  # PUNTO
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(10)
                    elif i == 3:  # EQUIPO
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(10)
                    elif i == 4:  # SERIAL
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(10)
                    elif i == 5:  # ESTADO
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.size = Pt(10)
                    
                    run.font.name = 'Calibri'
        
        logger.info(f"Tabla actualizada: {len(tabla_existente.rows)} filas totales (1 encabezado + {len(self.reintegrados_inventario)} datos)")
    
    def guardar(self, output_path: Path) -> None:
        """
        Genera y guarda la secci√≥n, asegurando que los cambios en las tablas se guarden correctamente
        """
        doc = self.generar()
        
        # Guardar el documento
        doc.save(str(output_path))
        logger.info(f"{self.nombre_seccion} guardada en: {output_path}")
