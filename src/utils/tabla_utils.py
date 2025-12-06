"""
Utilidades para crear tablas en documentos Word
"""
from typing import List, Dict, Any, Tuple
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import logging
import time
from concurrent.futures import ThreadPoolExecutor, as_completed

logger = logging.getLogger(__name__)

def crear_tabla_desde_dict(doc: Document, datos: List[Dict[str, Any]], encabezados: List[str]) -> None:
    """
    Crea una tabla en un documento Word a partir de una lista de diccionarios
    
    Args:
        doc: Documento Word
        datos: Lista de diccionarios con los datos
        encabezados: Lista de nombres de columnas
    """
    if not datos:
        return
    
    # Crear tabla
    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = 'Light Grid Accent 1'
    
    # Agregar encabezados
    header_cells = tabla.rows[0].cells
    for i, encabezado in enumerate(encabezados):
        header_cells[i].text = encabezado
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Agregar datos
    for fila_datos in datos:
        row_cells = tabla.add_row().cells
        for i, encabezado in enumerate(encabezados):
            valor = fila_datos.get(encabezado, "")
            row_cells[i].text = str(valor) if valor is not None else ""
    
    return tabla

def crear_tabla_desde_lista(doc: Document, datos: List[List[Any]], encabezados: List[str] = None) -> None:
    """
    Crea una tabla en un documento Word a partir de una lista de listas
    
    Args:
        doc: Documento Word
        datos: Lista de listas con los datos (cada lista es una fila)
        encabezados: Lista opcional de nombres de columnas
    """
    if not datos:
        return
    
    num_cols = len(datos[0]) if datos else len(encabezados) if encabezados else 0
    if num_cols == 0:
        return
    
    # Crear tabla
    num_rows = len(datos) + (1 if encabezados else 0)
    tabla = doc.add_table(rows=num_rows, cols=num_cols)
    tabla.style = 'Light Grid Accent 1'
    
    # Agregar encabezados si existen
    if encabezados:
        header_cells = tabla.rows[0].cells
        for i, encabezado in enumerate(encabezados):
            header_cells[i].text = str(encabezado)
            header_cells[i].paragraphs[0].runs[0].bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        inicio_datos = 1
    else:
        inicio_datos = 0
    
    # Agregar datos
    for idx, fila_datos in enumerate(datos):
        row_cells = tabla.rows[inicio_datos + idx].cells
        for i, valor in enumerate(fila_datos):
            row_cells[i].text = str(valor) if valor is not None else ""
    
    return tabla


def set_cell_shading(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), hex_color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def centrar_celda_vertical(cell):
    """Centra verticalmente el contenido de una celda"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


def enable_autofit(table):
    """Activa el autoajuste para tablas en Word"""
    tbl = table._element
    if tbl.tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    else:
        tblPr = tbl.tblPr
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is not None:
        tblPr.remove(tblLayout)
    new_tblLayout = OxmlElement('w:tblLayout')
    new_tblLayout.set(qn('w:type'), 'autofit')
    tblPr.append(new_tblLayout)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblPr.remove(tblW)


def habilitar_encabezado_repetido(table, num_filas: int = 1):
    """
    Habilita que el encabezado de la tabla se repita en cada página cuando hay un salto de página.
    
    Args:
        table: Tabla de Word
        num_filas: Número de filas a marcar como encabezado (por defecto 1, la primera fila)
    """
    if num_filas < 1 or num_filas > len(table.rows):
        num_filas = 1
    
    for i in range(num_filas):
        row = table.rows[i]
        tr = row._element  # Elemento XML de la fila
        
        # Obtener o crear trPr (propiedades de la fila)
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        
        # Buscar si ya existe tblHeader
        tblHeader = trPr.find(qn('w:tblHeader'))
        if tblHeader is None:
            # Crear el elemento tblHeader
            tblHeader = OxmlElement('w:tblHeader')
            trPr.append(tblHeader)
        
        # Establecer el valor a 'true' (aunque el elemento solo necesita existir)
        tblHeader.set(qn('w:val'), 'true')


def detectar_columnas_desde_datos(datos: List[Dict[str, Any]]) -> tuple[List[str], List[str]]:
    """Detecta automáticamente los headers y campos desde los datos"""
    if not datos:
        return [], []
    campos = list(datos[0].keys())
    headers = [campo.replace("_", " ").upper() for campo in campos]
    return headers, campos


def preparar_datos_tabla(datos: List[Dict[str, Any]], headers: List[str], campos: List[str], 
                          agregar_totales: bool = False, columna_total_texto: int = 0) -> List[List[str]]:
    """Prepara los datos de una tabla en formato lista de listas"""
    table_data = [headers]
    for item in datos:
        fila = []
        for campo in campos:
            valor = item.get(campo, "")
            fila.append(str(valor) if valor is not None else "")
        table_data.append(fila)
    if agregar_totales and datos:
        fila_total = []
        for idx, campo in enumerate(campos):
            if idx == columna_total_texto:
                fila_total.append("TOTAL")
            else:
                try:
                    total = 0
                    for item in datos:
                        valor = item.get(campo, "")
                        if valor is not None and valor != "":
                            try:
                                if isinstance(valor, str):
                                    valor_limpio = valor.replace(",", "").replace(" ", "").replace("$", "").strip()
                                    if valor_limpio and valor_limpio != "-":
                                        total += float(valor_limpio)
                                elif isinstance(valor, (int, float)):
                                    total += float(valor)
                            except (ValueError, TypeError, ZeroDivisionError) as e:
                                # Ignorar valores inválidos, continuar con el siguiente
                                logger.debug(f"Valor inválido en campo {campo}: {valor}, error: {e}")
                                continue
                    # Evitar división por cero
                    if total == int(total):
                        fila_total.append(str(int(total)))
                    else:
                        fila_total.append(str(total))
                except (ValueError, TypeError, ZeroDivisionError) as e:
                    logger.warning(f"Error al calcular total para campo {campo}: {e}")
                    fila_total.append("")
        table_data.append(fila_total)
    return table_data


def count_rows(table_data: List[List[str]]) -> int:
    """Cuenta el número de registros en una tabla preparada"""
    if not table_data or len(table_data) == 0:
        return 0
    count = len(table_data) - 1
    if len(table_data) > 1:
        ultima_fila = table_data[-1]
        if ultima_fila and len(ultima_fila) > 0:
            if str(ultima_fila[0]).upper() == "TOTAL":
                count -= 1
    return max(0, count)


def reemplazar_placeholder_con_tabla(doc: Document, placeholder: str, table_data: List[List[str]]):
    """Busca un placeholder y lo reemplaza con una tabla"""
    if not table_data or len(table_data) == 0:
        return
    estilos_tabla = ['Table Grid', 'Light Shading', 'Light List', 'Medium Shading 1', 'Light Grid']
    for i, paragraph in enumerate(doc.paragraphs):
        if placeholder in paragraph.text:
            tabla = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            try:
                tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
            except:
                try:
                    tbl_pr = tabla._tbl.tblPr
                    if tbl_pr is None:
                        from docx.oxml import parse_xml
                        tbl_pr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:jc w:val="center"/></w:tblPr>')
                        tabla._tbl.insert(0, tbl_pr)
                    else:
                        jc = tbl_pr.find(qn('w:jc'))
                        if jc is None:
                            jc = OxmlElement('w:jc')
                            jc.set(qn('w:val'), 'center')
                            tbl_pr.append(jc)
                        else:
                            jc.set(qn('w:val'), 'center')
                except Exception as e:
                    logger.warning(f"No se pudo centrar la tabla: {e}")
            estilo_aplicado = False
            for estilo in estilos_tabla:
                try:
                    tabla.style = estilo
                    estilo_aplicado = True
                    break
                except:
                    continue
            if not estilo_aplicado:
                try:
                    tabla.style = 'Table Grid'
                except:
                    pass
            enable_autofit(tabla)
            # Habilitar que el encabezado se repita en cada página
            habilitar_encabezado_repetido(tabla, num_filas=1)
            total_rows = len(table_data)
            for row_idx, fila in enumerate(table_data):
                for col_idx, valor in enumerate(fila):
                    celda = tabla.rows[row_idx].cells[col_idx]
                    celda.text = str(valor) if valor is not None else ""
                    if row_idx == 0:
                        set_cell_shading(celda, "1F4E79")
                        centrar_celda_vertical(celda)
                        for paragraph_celda in celda.paragraphs:
                            paragraph_celda.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph_celda.runs:
                                run.bold = True
                                run.font.size = Pt(8)
                                run.font.color.rgb = RGBColor(255, 255, 255)
                    elif row_idx == total_rows - 1 and fila[0].upper() == "TOTAL":
                        set_cell_shading(celda, "D9E1F2")
                        centrar_celda_vertical(celda)
                        for paragraph_celda in celda.paragraphs:
                            paragraph_celda.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph_celda.runs:
                                run.bold = True
                                run.font.size = Pt(8)
                                run.font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        if row_idx % 2 == 0:
                            set_cell_shading(celda, "F2F2F2")
                        else:
                            set_cell_shading(celda, "FFFFFF")
                        centrar_celda_vertical(celda)
                        for paragraph_celda in celda.paragraphs:
                            if col_idx == 0:
                                paragraph_celda.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            else:
                                paragraph_celda.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph_celda.runs:
                                run.font.size = Pt(8)
                                run.font.color.rgb = RGBColor(0, 0, 0)
            parent = paragraph._element.getparent()
            para_idx = parent.index(paragraph._element)
            parent.insert(para_idx + 1, tabla._element)
            paragraph.text = paragraph.text.replace(placeholder, "").strip()
            if not paragraph.text.strip():
                parent.remove(paragraph._element)
            break


def _crear_tabla_procesada(table_data: List[List[str]], estilos_tabla: List[str]) -> Tuple[Any, float]:
    """
    Crea y procesa una tabla de forma independiente (para procesamiento paralelo).
    Retorna el elemento XML de la tabla y el tiempo que tomó procesarla.
    
    Args:
        table_data: Datos de la tabla
        estilos_tabla: Lista de estilos a intentar
        
    Returns:
        Tupla (elemento_tabla, tiempo_procesamiento)
    """
    tiempo_inicio = time.time()
    
    # Crear un documento temporal para crear la tabla
    doc_temp = Document()
    total_rows = len(table_data)
    total_cols = len(table_data[0]) if table_data else 0
    
    if total_rows == 0 or total_cols == 0:
        return None, 0.0
    
    tabla = doc_temp.add_table(rows=total_rows, cols=total_cols)
    
    # Centrar tabla
    try:
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    except:
        try:
            tbl_pr = tabla._tbl.tblPr
            if tbl_pr is None:
                from docx.oxml import parse_xml
                tbl_pr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:jc w:val="center"/></w:tblPr>')
                tabla._tbl.insert(0, tbl_pr)
            else:
                jc = tbl_pr.find(qn('w:jc'))
                if jc is None:
                    jc = OxmlElement('w:jc')
                    jc.set(qn('w:val'), 'center')
                    tbl_pr.append(jc)
                else:
                    jc.set(qn('w:val'), 'center')
        except Exception as e:
            logger.warning(f"No se pudo centrar la tabla: {e}")
    
    # Aplicar estilo
    try:
        tabla.style = estilos_tabla[0]
    except:
        try:
            tabla.style = 'Table Grid'
        except:
            pass
    
    enable_autofit(tabla)
    habilitar_encabezado_repetido(tabla, num_filas=1)
    
    # OPTIMIZACIÓN CRÍTICA: Para tablas grandes, aplicar estilos mínimos
    es_fila_total = total_rows > 1 and table_data[total_rows - 1][0].upper() == "TOTAL"
    es_tabla_grande = total_rows > 100
    es_tabla_muy_grande = total_rows > 200  # Tablas con más de 200 filas - formato ultra mínimo
    
    # Llenar todos los datos primero (más rápido usando acceso directo)
    for row_idx, fila in enumerate(table_data):
        for col_idx, valor in enumerate(fila):
            celda = tabla.rows[row_idx].cells[col_idx]
            # Usar acceso directo al texto sin crear runs innecesarios
            celda.paragraphs[0].text = str(valor) if valor is not None else ""
    
    # OPTIMIZACIÓN ULTRA AGRESIVA: Para tablas muy grandes, solo formato en encabezado
    if es_tabla_muy_grande:
        # Solo aplicar formato al encabezado, nada más (SIN cambiar tamaño de fuente para acelerar)
        for col_idx in range(total_cols):
            celda = tabla.rows[0].cells[col_idx]
            parrafo = celda.paragraphs[0]
            run = parrafo.runs[0] if parrafo.runs else None
            
            set_cell_shading(celda, "1F4E79")
            centrar_celda_vertical(celda)
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if run:
                run.bold = True
                # NO cambiar tamaño de fuente para acelerar (usa el tamaño por defecto del template)
                run.font.color.rgb = RGBColor(255, 255, 255)
    else:
        # Aplicar estilos de manera optimizada solo a encabezado y totales
        for row_idx in range(total_rows):
            fila = table_data[row_idx]
            for col_idx in range(total_cols):
                celda = tabla.rows[row_idx].cells[col_idx]
                parrafo = celda.paragraphs[0]
                run = parrafo.runs[0] if parrafo.runs else None
                
                if row_idx == 0:
                    # Fila de encabezado - SIEMPRE formatear
                    set_cell_shading(celda, "1F4E79")
                    centrar_celda_vertical(celda)
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if run:
                        run.bold = True
                        # Solo cambiar tamaño de fuente en encabezado si no es tabla muy grande
                        if not es_tabla_muy_grande:
                            run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                elif row_idx == total_rows - 1 and es_fila_total:
                    # Fila de totales - SIEMPRE formatear
                    set_cell_shading(celda, "D9E1F2")
                    centrar_celda_vertical(celda)
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if run:
                        run.bold = True
                        # Solo cambiar tamaño de fuente en totales si no es tabla muy grande
                        if not es_tabla_muy_grande:
                            run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                elif not es_tabla_grande:
                    # Fila normal en tablas pequeñas - aplicar formato completo
                    set_cell_shading(celda, "F2F2F2" if row_idx % 2 == 0 else "FFFFFF")
                    centrar_celda_vertical(celda)
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                    if run:
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                # Para tablas grandes (100-200 filas), no aplicar formato a filas normales
    
    # Retornar el elemento XML de la tabla (clonado para evitar problemas de referencia)
    tabla_element = tabla._element
    tiempo_procesamiento = time.time() - tiempo_inicio
    
    return tabla_element, tiempo_procesamiento


def reemplazar_multiples_placeholders_con_tablas(doc: Document, placeholders_tablas: Dict[str, List[List[str]]]):
    """
    Reemplaza múltiples placeholders con tablas en una sola pasada del documento.
    OPTIMIZADO: Usa acceso directo a XML para aplicar estilos más rápido.
    
    Args:
        doc: Documento Word
        placeholders_tablas: Diccionario con formato {placeholder: table_data}
    """
    if not placeholders_tablas:
        return
    
    # Crear un conjunto de placeholders para búsqueda rápida
    placeholders_set = set(placeholders_tablas.keys())
    estilos_tabla = ['Table Grid', 'Light Shading', 'Light List', 'Medium Shading 1', 'Light Grid']
    
    # Buscar todos los placeholders en una sola pasada
    tiempo_busqueda = time.time()
    paragraphs_to_process = []
    
    # Log de placeholders que se buscan
    logger.info(f"  🔍 Buscando placeholders: {list(placeholders_set)}")
    
    # Buscar en TODOS los párrafos (no solo los primeros)
    logger.info(f"  📄 Buscando en {len(doc.paragraphs)} párrafos del documento...")
    for i, paragraph in enumerate(doc.paragraphs):
        texto_parrafo = paragraph.text  # Acceder una sola vez
        # Log de los primeros párrafos para depuración
        if i < 10:
            logger.debug(f"  Párrafo {i}: '{texto_parrafo[:150]}...'")
        
        for placeholder in placeholders_set:
            # Buscar el placeholder exacto
            if placeholder in texto_parrafo:
                table_data = placeholders_tablas[placeholder]
                if table_data and len(table_data) > 0:
                    paragraphs_to_process.append((i, paragraph, placeholder, table_data))
                    logger.info(f"  ✓ Placeholder '{placeholder}' encontrado en párrafo {i}: '{texto_parrafo[:100]}...'")
                    break  # Solo procesar un placeholder por párrafo
            # También buscar sin los corchetes dobles (por si Jinja2 los procesó diferente)
            elif placeholder.replace("[[", "").replace("]]", "") in texto_parrafo:
                logger.warning(f"  ⚠️ Placeholder '{placeholder}' encontrado sin corchetes en párrafo {i}")
                # Intentar buscar la variación
                placeholder_variacion = placeholder.replace("[[", "").replace("]]", "")
                if placeholder_variacion in texto_parrafo:
                    table_data = placeholders_tablas[placeholder]
                    if table_data and len(table_data) > 0:
                        paragraphs_to_process.append((i, paragraph, placeholder, table_data))
                        logger.info(f"  ✓ Placeholder procesado como variación en párrafo {i}")
                        break
            # Buscar también sin espacios (por si Word agregó espacios)
            elif placeholder.replace(" ", "") in texto_parrafo.replace(" ", ""):
                logger.warning(f"  ⚠️ Placeholder '{placeholder}' encontrado con espacios diferentes en párrafo {i}")
                table_data = placeholders_tablas[placeholder]
                if table_data and len(table_data) > 0:
                    paragraphs_to_process.append((i, paragraph, placeholder, table_data))
                    logger.info(f"  ✓ Placeholder procesado con espacios en párrafo {i}")
                    break
    
    # Si no se encontraron en párrafos, buscar en tablas existentes y otros lugares
    if len(paragraphs_to_process) == 0:
        logger.warning(f"  ⚠️ No se encontraron placeholders en párrafos. Buscando en tablas, headers, footers...")
        
        # Buscar en tablas existentes
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Buscar en todos los párrafos de la celda
                    for para_idx, para in enumerate(cell.paragraphs):
                        texto_celda = para.text
                        for placeholder in placeholders_set:
                            if placeholder in texto_celda:
                                table_data = placeholders_tablas[placeholder]
                                if table_data and len(table_data) > 0:
                                    # Reemplazar el contenido de la celda con la tabla
                                    logger.info(f"  ✓ Placeholder '{placeholder}' encontrado en tabla {table_idx}, celda ({row_idx}, {cell_idx}), párrafo {para_idx}")
                                    # Limpiar la celda
                                    cell.text = ""
                                    # Insertar tabla después de la fila de la tabla
                                    # Necesitamos insertar la tabla después de la tabla existente
                                    parent = table._element.getparent()
                                    table_idx_in_parent = parent.index(table._element)
                                    
                                    # Crear la tabla
                                    from src.utils.tabla_utils import _crear_tabla_procesada
                                    tabla_element, _ = _crear_tabla_procesada(table_data, estilos_tabla)
                                    if tabla_element:
                                        # Insertar después de la tabla actual
                                        parent.insert(table_idx_in_parent + 1, tabla_element)
                                        paragraphs_to_process.append((None, None, placeholder, table_data))  # Marcar como procesado
                                        logger.info(f"  ✓ Tabla insertada después de tabla {table_idx}")
                                    break
        
        # Buscar en headers y footers
        for section in doc.sections:
            # Headers
            for header in [section.header]:
                for para in header.paragraphs:
                    texto_para = para.text
                    for placeholder in placeholders_set:
                        if placeholder in texto_para:
                            logger.info(f"  ✓ Placeholder '{placeholder}' encontrado en header")
                            # Nota: Los headers/footers requieren un manejo especial
                            break
            
            # Footers
            for footer in [section.footer]:
                for para in footer.paragraphs:
                    texto_para = para.text
                    for placeholder in placeholders_set:
                        if placeholder in texto_para:
                            logger.info(f"  ✓ Placeholder '{placeholder}' encontrado en footer")
                            break
    
    tiempo_busqueda_total = time.time() - tiempo_busqueda
    logger.info(f"  🔍 Búsqueda de placeholders: {tiempo_busqueda_total:.2f}s ({len(paragraphs_to_process)} tablas encontradas)")
    
    if len(paragraphs_to_process) == 0:
        logger.warning(f"  ⚠️ ADVERTENCIA: No se encontraron placeholders en el documento. Verifica que el template tenga:")
        for placeholder in placeholders_set:
            logger.warning(f"     - {placeholder}")
    
    # Procesar tablas en paralelo usando ThreadPoolExecutor
    tiempo_total_procesamiento = time.time()
    
    # Determinar si usar procesamiento paralelo (solo si hay más de 1 tabla)
    usar_paralelo = len(paragraphs_to_process) > 1
    
    if usar_paralelo:
        logger.info(f"  🚀 Procesando {len(paragraphs_to_process)} tablas en paralelo...")
        tablas_procesadas = {}
        
        # Procesar tablas en paralelo
        with ThreadPoolExecutor(max_workers=min(len(paragraphs_to_process), 4)) as executor:
            # Enviar todas las tareas
            future_to_tabla = {
                executor.submit(_crear_tabla_procesada, table_data, estilos_tabla): (tabla_idx, para_idx, paragraph, placeholder, table_data)
                for tabla_idx, (para_idx, paragraph, placeholder, table_data) in enumerate(paragraphs_to_process)
            }
            
            # Recoger resultados conforme se completan
            for future in as_completed(future_to_tabla):
                tabla_idx, para_idx, paragraph, placeholder, table_data = future_to_tabla[future]
                try:
                    tabla_element, tiempo_procesamiento = future.result()
                    tablas_procesadas[tabla_idx] = {
                        'element': tabla_element,
                        'paragraph': paragraph,
                        'placeholder': placeholder,
                        'table_data': table_data,
                        'tiempo': tiempo_procesamiento
                    }
                    if tiempo_procesamiento > 1.0:
                        total_rows = len(table_data)
                        total_cols = len(table_data[0]) if table_data else 0
                        logger.info(f"    ✅ Tabla {tabla_idx + 1}/{len(paragraphs_to_process)} ({placeholder}) procesada: {tiempo_procesamiento:.2f}s ({total_rows}x{total_cols})")
                except Exception as e:
                    logger.error(f"    ❌ Error procesando tabla {tabla_idx + 1} ({placeholder}): {e}")
        
        # Insertar tablas en el documento en orden (secuencial para evitar problemas)
        tiempo_insercion = time.time()
        for tabla_idx in sorted(tablas_procesadas.keys()):
            info = tablas_procesadas[tabla_idx]
            paragraph = info['paragraph']
            placeholder = info['placeholder']
            tabla_element = info['element']
            
            if tabla_element is None:
                continue
            
            # Clonar el elemento para insertarlo en el documento
            from copy import deepcopy
            tabla_element_clonado = deepcopy(tabla_element)
            
            # Insertar tabla en el documento
            parent = paragraph._element.getparent()
            para_idx = parent.index(paragraph._element)
            parent.insert(para_idx + 1, tabla_element_clonado)
            
            # Limpiar placeholder del párrafo
            texto_original = paragraph.text
            nuevo_texto = texto_original.replace(placeholder, "").strip()
            if nuevo_texto != texto_original:
                paragraph.text = nuevo_texto
                if not nuevo_texto:
                    parent.remove(paragraph._element)
        
        tiempo_insercion_total = time.time() - tiempo_insercion
        logger.info(f"  📌 Inserción de tablas en documento: {tiempo_insercion_total:.2f}s")
    else:
        # Procesamiento secuencial (para 1 tabla o si hay problemas)
        logger.info(f"  📝 Procesando {len(paragraphs_to_process)} tabla(s) secuencialmente...")
        for tabla_idx, (para_idx, paragraph, placeholder, table_data) in enumerate(paragraphs_to_process):
            tiempo_tabla = time.time()
            total_rows = len(table_data)
            total_cols = len(table_data[0]) if table_data else 0
            
            if total_rows == 0 or total_cols == 0:
                continue
            
            tabla = doc.add_table(rows=total_rows, cols=total_cols)
            
            # Centrar tabla (optimizado)
            try:
                tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
            except:
                try:
                    tbl_pr = tabla._tbl.tblPr
                    if tbl_pr is None:
                        from docx.oxml import parse_xml
                        tbl_pr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:jc w:val="center"/></w:tblPr>')
                        tabla._tbl.insert(0, tbl_pr)
                    else:
                        jc = tbl_pr.find(qn('w:jc'))
                        if jc is None:
                            jc = OxmlElement('w:jc')
                            jc.set(qn('w:val'), 'center')
                            tbl_pr.append(jc)
                        else:
                            jc.set(qn('w:val'), 'center')
                except Exception as e:
                    logger.warning(f"No se pudo centrar la tabla: {e}")
            
            # Aplicar estilo (solo intentar el primero que funcione)
            try:
                tabla.style = estilos_tabla[0]
            except:
                try:
                    tabla.style = 'Table Grid'
                except:
                    pass
            
            enable_autofit(tabla)
            habilitar_encabezado_repetido(tabla, num_filas=1)
            
            # OPTIMIZACIÓN CRÍTICA: Para tablas grandes, aplicar estilos mínimos
            es_fila_total = total_rows > 1 and table_data[total_rows - 1][0].upper() == "TOTAL"
            es_tabla_grande = total_rows > 100
            es_tabla_muy_grande = total_rows > 200  # Tablas con más de 200 filas - formato ultra mínimo
            
            # Para tablas grandes, solo aplicar formato a encabezado y totales
            if es_tabla_muy_grande:
                logger.info(f"    ⚡ Tabla MUY grande detectada ({total_rows}x{total_cols}), aplicando formato ULTRA mínimo (solo encabezado)")
            elif es_tabla_grande:
                logger.info(f"    📊 Tabla grande detectada ({total_rows}x{total_cols}), aplicando formato optimizado")
            
            # Llenar todos los datos primero (más rápido)
            tiempo_llenado = time.time()
            for row_idx, fila in enumerate(table_data):
                for col_idx, valor in enumerate(fila):
                    celda = tabla.rows[row_idx].cells[col_idx]
                    celda.paragraphs[0].text = str(valor) if valor is not None else ""
            tiempo_llenado_total = time.time() - tiempo_llenado
            
            # OPTIMIZACIÓN ULTRA AGRESIVA: Para tablas muy grandes, solo formato en encabezado
            tiempo_estilos = time.time()
            if es_tabla_muy_grande:
                # Solo aplicar formato al encabezado, nada más (SIN cambiar tamaño de fuente para acelerar)
                for col_idx in range(total_cols):
                    celda = tabla.rows[0].cells[col_idx]
                    parrafo = celda.paragraphs[0]
                    run = parrafo.runs[0] if parrafo.runs else None
                    
                    set_cell_shading(celda, "1F4E79")
                    centrar_celda_vertical(celda)
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if run:
                        run.bold = True
                        # NO cambiar tamaño de fuente para acelerar (usa el tamaño por defecto del template)
                        run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                # Aplicar estilos de manera optimizada solo a encabezado y totales
                for row_idx in range(total_rows):
                    fila = table_data[row_idx]
                    for col_idx in range(total_cols):
                        celda = tabla.rows[row_idx].cells[col_idx]
                        parrafo = celda.paragraphs[0]
                        run = parrafo.runs[0] if parrafo.runs else None
                        
                        if row_idx == 0:
                            # Fila de encabezado - SIEMPRE formatear
                            set_cell_shading(celda, "1F4E79")
                            centrar_celda_vertical(celda)
                            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if run:
                                run.bold = True
                                # Solo cambiar tamaño de fuente en encabezado si no es tabla muy grande
                                if not es_tabla_muy_grande:
                                    run.font.size = Pt(8)
                                run.font.color.rgb = RGBColor(255, 255, 255)
                        elif row_idx == total_rows - 1 and es_fila_total:
                            # Fila de totales - SIEMPRE formatear
                            set_cell_shading(celda, "D9E1F2")
                            centrar_celda_vertical(celda)
                            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if run:
                                run.bold = True
                                # Solo cambiar tamaño de fuente en totales si no es tabla muy grande
                                if not es_tabla_muy_grande:
                                    run.font.size = Pt(8)
                                run.font.color.rgb = RGBColor(0, 0, 0)
                        elif not es_tabla_grande:
                            # Fila normal en tablas pequeñas - aplicar formato completo
                            set_cell_shading(celda, "F2F2F2" if row_idx % 2 == 0 else "FFFFFF")
                            centrar_celda_vertical(celda)
                            parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                            if run:
                                run.font.size = Pt(8)
                                run.font.color.rgb = RGBColor(0, 0, 0)
                        # Para tablas grandes (100-200 filas), no aplicar formato a filas normales
            tiempo_estilos_total = time.time() - tiempo_estilos
            
            if tiempo_llenado_total > 1.0 or tiempo_estilos_total > 1.0:
                logger.info(f"      - Llenado datos: {tiempo_llenado_total:.2f}s")
                logger.info(f"      - Aplicación estilos: {tiempo_estilos_total:.2f}s")
            
            # Insertar tabla en el documento
            parent = paragraph._element.getparent()
            para_idx = parent.index(paragraph._element)
            parent.insert(para_idx + 1, tabla._element)
            
            # Limpiar placeholder del párrafo
            texto_original = paragraph.text
            nuevo_texto = texto_original.replace(placeholder, "").strip()
            if nuevo_texto != texto_original:
                paragraph.text = nuevo_texto
                if not nuevo_texto:
                    parent.remove(paragraph._element)
            
            tiempo_tabla_total = time.time() - tiempo_tabla
            if tiempo_tabla_total > 1.0:  # Solo loggear si tarda más de 1 segundo
                logger.info(f"    ⚠️  Tabla {tabla_idx + 1}/{len(paragraphs_to_process)} ({placeholder}): {tiempo_tabla_total:.2f}s ({total_rows}x{total_cols})")
    
    tiempo_total_procesamiento_total = time.time() - tiempo_total_procesamiento
    logger.info(f"  ⏱️  Tiempo total procesamiento tablas: {tiempo_total_procesamiento_total:.2f}s")